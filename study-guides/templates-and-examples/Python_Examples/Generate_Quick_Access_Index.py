#!/usr/bin/env python3
"""
GENERATE QUICK_ACCESS.MD INDEX

Creates searchable markdown index of all drugs/conditions across study guides.
Solves the "can't find which file has Drug X" problem with alphabetical lookup.

Purpose:
- Scans all study guide files in a directory
- Extracts drug/condition names
- Creates alphabetical QUICK_ACCESS.md index
- Maps each entity to its source file(s)

Usage:
    python Generate_Quick_Access_Index.py <study_guides_directory>

Example:
    python Generate_Quick_Access_Index.py "Pharmacology/Exam 3/Claude Study Tools/"

Output:
    QUICK_ACCESS.md in the specified directory

Features:
- Alphabetical organization with letter section headers
- Links to source files
- Multiple files per entity (if entity appears in multiple guides)
- Auto-updates when new files added
- Fast visual scanning
"""

import sys
import os
import re
from pathlib import Path
from openpyxl import load_workbook
from docx import Document


def extract_drugs_from_excel(file_path):
    """
    Extract drug names from Excel study guide.
    Looks for Drug Name column (typically column B).

    Args:
        file_path: Path to Excel file

    Returns:
        list: Drug names found
    """
    drugs = []
    try:
        wb = load_workbook(file_path, read_only=True, data_only=True)

        # Check each sheet
        for sheet_name in wb.sheetnames:
            if sheet_name in ["Index", "High-Yield & Pearls", "Summary"]:
                continue  # Skip non-drug-list sheets

            sheet = wb[sheet_name]

            # Look for "Drug Name" header in first 3 rows
            drug_col_idx = None
            for row_idx in range(1, min(4, sheet.max_row + 1)):
                for col_idx in range(1, min(10, sheet.max_column + 1)):
                    cell_value = sheet.cell(row_idx, col_idx).value
                    if cell_value and "drug" in str(cell_value).lower():
                        drug_col_idx = col_idx
                        header_row = row_idx
                        break
                if drug_col_idx:
                    break

            # Extract drugs from that column
            if drug_col_idx:
                for row_idx in range(header_row + 1, sheet.max_row + 1):
                    cell_value = sheet.cell(row_idx, drug_col_idx).value
                    if cell_value:
                        drug_name = str(cell_value).strip()
                        # Remove (Brand Name) suffix if present
                        drug_name = re.sub(r'\s*\([^)]+\)\s*$', '', drug_name)
                        if drug_name and len(drug_name) > 1:
                            drugs.append(drug_name)

        wb.close()
    except Exception as e:
        print(f"  ⚠️  Could not extract from {os.path.basename(file_path)}: {e}")

    return drugs


def extract_drugs_from_word(file_path):
    """
    Extract drug/condition names from Word study guide.
    Looks for drug names in tables and headings.

    Args:
        file_path: Path to Word file

    Returns:
        list: Drug/condition names found
    """
    drugs = []
    try:
        doc = Document(file_path)

        # Extract from tables
        for table in doc.tables:
            # Check first row for "Drug" or "Condition" column
            header_row = table.rows[0]
            drug_col_idx = None

            for idx, cell in enumerate(header_row.cells):
                cell_text = cell.text.lower()
                if "drug" in cell_text or "condition" in cell_text or "name" in cell_text:
                    drug_col_idx = idx
                    break

            # Extract from that column
            if drug_col_idx is not None:
                for row in table.rows[1:]:  # Skip header
                    if drug_col_idx < len(row.cells):
                        cell_text = row.cells[drug_col_idx].text.strip()
                        # Remove (Brand Name) suffix if present
                        cell_text = re.sub(r'\s*\([^)]+\)\s*$', '', cell_text)
                        if cell_text and len(cell_text) > 1:
                            drugs.append(cell_text)

        # Extract from headings (Level 2 and 3)
        for para in doc.paragraphs:
            if para.style.name in ['Heading 2', 'Heading 3']:
                heading_text = para.text.strip()
                # Remove common prefixes
                heading_text = re.sub(r'^(Drug|Condition|Learning Objective \d+):\s*', '', heading_text)
                if heading_text and len(heading_text) > 1 and len(heading_text) < 100:
                    drugs.append(heading_text)

    except Exception as e:
        print(f"  ⚠️  Could not extract from {os.path.basename(file_path)}: {e}")

    return drugs


def extract_drugs_from_anki_csv(file_path):
    """
    Extract drug names from Anki CSV file.
    Looks for drug names in questions.

    Args:
        file_path: Path to CSV file

    Returns:
        list: Drug names found
    """
    drugs = set()
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            for line in f:
                # Look for patterns like "What is the mechanism of [Drug]?"
                matches = re.findall(r'(?:mechanism|class|use|effect).*?(?:of|for)\s+([A-Z][a-z]+(?:[a-z]+)?)', line)
                drugs.update(matches)

    except Exception as e:
        print(f"  ⚠️  Could not extract from {os.path.basename(file_path)}: {e}")

    return list(drugs)


def scan_study_guides_directory(directory):
    """
    Scan directory for all study guide files and extract entities.

    Args:
        directory: Path to directory containing study guides

    Returns:
        dict: {entity_name: [list of files containing it]}
    """
    entity_files_map = {}

    print(f"📂 Scanning directory: {directory}")
    print("")

    # Find all study guide files
    excel_files = list(Path(directory).glob("*.xlsx"))
    word_files = list(Path(directory).glob("*.docx"))
    csv_files = list(Path(directory).glob("*Flashcards.csv"))

    total_files = len(excel_files) + len(word_files) + len(csv_files)
    print(f"Found {total_files} study guide files:")
    print(f"  - Excel: {len(excel_files)}")
    print(f"  - Word: {len(word_files)}")
    print(f"  - Anki CSV: {len(csv_files)}")
    print("")

    # Process Excel files
    for excel_file in excel_files:
        print(f"📊 Processing: {excel_file.name}")
        drugs = extract_drugs_from_excel(str(excel_file))
        print(f"  ✓ Found {len(drugs)} entities")

        for drug in drugs:
            if drug not in entity_files_map:
                entity_files_map[drug] = []
            entity_files_map[drug].append(excel_file.name)

    # Process Word files
    for word_file in word_files:
        print(f"📝 Processing: {word_file.name}")
        drugs = extract_drugs_from_word(str(word_file))
        print(f"  ✓ Found {len(drugs)} entities")

        for drug in drugs:
            if drug not in entity_files_map:
                entity_files_map[drug] = []
            entity_files_map[drug].append(word_file.name)

    # Process CSV files
    for csv_file in csv_files:
        print(f"🃏 Processing: {csv_file.name}")
        drugs = extract_drugs_from_anki_csv(str(csv_file))
        print(f"  ✓ Found {len(drugs)} entities")

        for drug in drugs:
            if drug not in entity_files_map:
                entity_files_map[drug] = []
            entity_files_map[drug].append(csv_file.name)

    print("")
    print(f"✓ Extracted {len(entity_files_map)} unique entities total")
    print("")

    return entity_files_map


def generate_quick_access_md(entity_files_map, output_path):
    """
    Generate QUICK_ACCESS.md file with alphabetical index.

    Args:
        entity_files_map: dict {entity_name: [files]}
        output_path: Path to save QUICK_ACCESS.md

    Returns:
        str: Path to created file
    """
    print(f"📝 Generating QUICK_ACCESS.md...")

    # Sort entities alphabetically
    sorted_entities = sorted(entity_files_map.keys(), key=str.lower)

    # Build markdown content
    lines = []
    lines.append("# Quick Access Index - Pharmacology Study Guides")
    lines.append("")
    lines.append("**Purpose:** Fast lookup of which study guide contains which drug/condition")
    lines.append("")
    lines.append("**Last Updated:** " + __import__('datetime').datetime.now().strftime("%Y-%m-%d %H:%M"))
    lines.append("")
    lines.append("---")
    lines.append("")

    # Group by first letter
    current_letter = None
    for entity in sorted_entities:
        first_letter = entity[0].upper()

        # Add letter section header
        if first_letter != current_letter:
            if current_letter is not None:
                lines.append("")  # Blank line before new section

            lines.append(f"## {first_letter}")
            lines.append("")
            current_letter = first_letter

        # Add entity entry
        files = entity_files_map[entity]

        # Format: - **Entity** → file1.xlsx, file2.docx
        if len(files) == 1:
            lines.append(f"- **{entity}** → {files[0]}")
        else:
            file_list = ", ".join(files)
            lines.append(f"- **{entity}** → {file_list}")

    # Footer
    lines.append("")
    lines.append("---")
    lines.append("")
    lines.append("**Total Entities:** " + str(len(entity_files_map)))
    lines.append("")
    lines.append("**How to use:**")
    lines.append("1. Use Ctrl+F (Cmd+F on Mac) to search for drug/condition name")
    lines.append("2. See which file(s) contain that entity")
    lines.append("3. Open the file to access detailed information")
    lines.append("")

    # Write to file
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(lines))

    print(f"✓ Created: {output_path}")
    print(f"  - {len(entity_files_map)} entities indexed")
    print(f"  - {len([l for l in lines if l.startswith('##')])} letter sections")
    print("")

    return output_path


def main():
    """Main entry point for script."""
    if len(sys.argv) != 2:
        print("Usage: python Generate_Quick_Access_Index.py <study_guides_directory>")
        print("")
        print("Example:")
        print('  python Generate_Quick_Access_Index.py "Pharmacology/Exam 3/Claude Study Tools/"')
        sys.exit(1)

    directory = sys.argv[1]

    if not os.path.exists(directory):
        print(f"❌ ERROR: Directory not found: {directory}")
        sys.exit(1)

    if not os.path.isdir(directory):
        print(f"❌ ERROR: Path is not a directory: {directory}")
        sys.exit(1)

    print("═══════════════════════════════════════")
    print("  GENERATE QUICK ACCESS INDEX")
    print("═══════════════════════════════════════")
    print("")

    # Scan directory and extract entities
    entity_files_map = scan_study_guides_directory(directory)

    if not entity_files_map:
        print("⚠️  WARNING: No entities found in study guides")
        print("   Make sure directory contains .xlsx, .docx, or .csv files")
        sys.exit(1)

    # Generate QUICK_ACCESS.md
    output_path = os.path.join(directory, "QUICK_ACCESS.md")
    generate_quick_access_md(entity_files_map, output_path)

    print("═══════════════════════════════════════")
    print("  INDEX GENERATION COMPLETE")
    print("═══════════════════════════════════════")
    print(f"Output: {output_path}")
    print("")
    print("✅ QUICK_ACCESS.md created successfully!")
    print("   Use Ctrl+F to quickly find any drug/condition")


if __name__ == '__main__':
    main()
