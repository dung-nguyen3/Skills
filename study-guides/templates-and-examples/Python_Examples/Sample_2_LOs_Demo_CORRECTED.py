#!/usr/bin/env python3
"""
Sample Word LO Document - CORRECTED Formatting Demo
Demonstrates ALL fixes:
1. NO "Summary:" labels
2. NO "TABLE X:" prefixes
3. Word default indents (0.5" per level)
4. Universal formatting (Calibri, 12pt, consistent spacing)
"""

from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# =============================================================================
# HELPER FUNCTIONS
# =============================================================================

def set_cell_background(cell, color_hex):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    cell._element.get_or_add_tcPr().append(shading_elm)

def set_cell_text(cell, text, bold=False, size=12):
    """Set cell text with formatting - DEFAULT 12pt"""
    cell.text = text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(size)
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.bold = bold

def add_plain_bullets(doc, items, bold_title=None):
    """Add plain bulleted list (Clinical Pearls) - Calibri 12pt with hanging indent"""
    if bold_title:
        para = doc.add_paragraph()
        run = para.add_run(bold_title)
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(12)

    for item in items:
        para = doc.add_paragraph(item, style='List Bullet')
        # Word default: bullet at 0.25", text at 0.5"
        para.paragraph_format.left_indent = Inches(0.5)
        para.paragraph_format.first_line_indent = Inches(-0.25)  # Hanging indent
        # Universal formatting: 1.5 spacing
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = 1.5
        for run in para.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(12)

def add_summary_simple(doc, summary_text):
    """Add simple paragraph summary - NO 'Summary:' label"""
    summary_para = doc.add_paragraph(summary_text)
    # Universal formatting: 1.5 spacing
    summary_para.paragraph_format.space_before = Pt(0)
    summary_para.paragraph_format.space_after = Pt(0)
    summary_para.paragraph_format.line_spacing = 1.5
    summary_para.runs[0].font.name = 'Calibri'
    summary_para.runs[0].font.size = Pt(12)
    return summary_para

def add_summary_structured(doc, intro_text, items):
    """
    Add structured numbered summary - NO 'Summary:' label
    Uses Word multilevel list defaults with hanging indents
    """
    # Intro paragraph (NO label)
    summary_intro = doc.add_paragraph(intro_text)
    summary_intro.paragraph_format.space_before = Pt(0)
    summary_intro.paragraph_format.space_after = Pt(0)
    summary_intro.paragraph_format.line_spacing = 1.5
    summary_intro.runs[0].font.name = 'Calibri'
    summary_intro.runs[0].font.size = Pt(12)

    # Numbered items with manual numbering
    for idx, item in enumerate(items, start=1):
        # Main numbered item
        p = doc.add_paragraph()

        # Add bold number
        num_run = p.add_run(f'{idx}. ')
        num_run.bold = True
        num_run.font.size = Pt(12)
        num_run.font.name = 'Calibri'

        # Add main text
        text_run = p.add_run(item['main'])
        text_run.font.size = Pt(12)
        text_run.font.name = 'Calibri'

        # Word default: number at 0", text at 0.25"
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)  # Hanging indent
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.5

        # Sub-bullets (if any)
        if 'subs' in item and item['subs']:
            for sub in item['subs']:
                sub_p = doc.add_paragraph(sub, style='List Bullet')
                # Word default: bullet at 0.25", text at 0.5"
                sub_p.paragraph_format.left_indent = Inches(0.5)
                sub_p.paragraph_format.first_line_indent = Inches(-0.25)  # Hanging indent
                # Universal formatting: 1.5 spacing
                sub_p.paragraph_format.space_before = Pt(0)
                sub_p.paragraph_format.space_after = Pt(0)
                sub_p.paragraph_format.line_spacing = 1.5
                sub_p.runs[0].font.name = 'Calibri'
                sub_p.runs[0].font.size = Pt(12)

# =============================================================================
# MAIN DOCUMENT CREATION
# =============================================================================

def create_sample_document(output_path):
    """Create sample Word document with CORRECTED formatting"""

    doc = Document()

    # Set margins
    for section in doc.sections:
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)

    # ==========================================================================
    # TITLE PAGE
    # ==========================================================================
    title = doc.add_heading('Formatting Demo - CORRECTED', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(118, 75, 162)

    subtitle = doc.add_paragraph('2 Learning Objectives - All Fixes Applied')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(118, 75, 162)

    doc.add_page_break()

    # ==========================================================================
    # LEARNING OBJECTIVES SECTION
    # ==========================================================================

    # --- Learning Objective 1: Simple Summary (NO "Summary:" label) ---
    obj1 = doc.add_heading('1. Describe the mechanisms of action of beta blockers and their clinical uses', 2)
    obj1.runs[0].font.color.rgb = RGBColor(118, 75, 162)
    obj1.runs[0].font.size = Pt(14)

    # Simple paragraph summary (NO "Summary:" label)
    add_summary_simple(doc,
        'Beta blockers competitively antagonize β-adrenergic receptors, reducing heart rate, '
        'contractility, and blood pressure. They are classified as selective (β₁) or non-selective '
        '(β₁ and β₂). Clinical uses include hypertension, angina, heart failure, and arrhythmias.')

    doc.add_paragraph()

    # NO "TABLE 1:" prefix - descriptive title only
    doc.add_heading('Selective vs Non-Selective Beta Blockers', 3)
    table1 = doc.add_table(rows=4, cols=3)
    table1.style = 'Table Grid'

    headers = ['Feature', 'Selective (β₁)', 'Non-Selective (β₁+β₂)']
    for i, header_text in enumerate(headers):
        set_cell_background(table1.rows[0].cells[i], 'B3E5FC')
        set_cell_text(table1.rows[0].cells[i], header_text, bold=True, size=12)

    row_data = [
        ['Examples', 'Metoprolol, Atenolol, Bisoprolol', 'Propranolol, Nadolol, Timolol'],
        ['Target', 'β₁ receptors (heart)', 'β₁ + β₂ receptors'],
        ['Respiratory Effects', 'Minimal bronchospasm risk', '⚠️ Bronchospasm risk']
    ]

    for row_idx, row_content in enumerate(row_data, start=1):
        cells = table1.rows[row_idx].cells
        set_cell_background(cells[0], 'E1F5FE')
        set_cell_text(cells[0], row_content[0], bold=True, size=12)
        set_cell_background(cells[1], 'FFFFFF')
        set_cell_text(cells[1], row_content[1], size=12)
        set_cell_background(cells[2], 'FFFFFF')
        set_cell_text(cells[2], row_content[2], size=12)

    doc.add_paragraph()

    # Clinical Pearls with proper indent (0.5")
    add_plain_bullets(doc, [
        'Beta blockers reduce mortality in heart failure (bisoprolol, carvedilol, metoprolol)',
        'Never stop beta blockers abruptly - risk of rebound tachycardia',
        'Cardioselective agents lose selectivity at high doses'
    ], bold_title='Clinical Pearls & High-Yield Points:')

    doc.add_page_break()

    # --- Learning Objective 2: Structured Summary (NO "Summary:" label) ---
    obj2 = doc.add_heading('2. Describe the regulation of digestive motility', 2)
    obj2.runs[0].font.color.rgb = RGBColor(118, 75, 162)
    obj2.runs[0].font.size = Pt(14)

    # Structured numbered summary (NO "Summary:" label, proper indents)
    add_summary_structured(doc,
        'Digestive motility is regulated by three main mechanisms:',
        [
            {
                'main': 'Neural regulation',
                'subs': [
                    'Parasympathetic (vagus nerve) increases intestinal contraction',
                    'Sympathetic nervous system decreases motility during stress'
                ]
            },
            {
                'main': 'Hormonal regulation',
                'subs': [
                    'Secretin stimulates bicarbonate release from pancreas',
                    'Cholecystokinin (CCK) stimulates enzyme release and gallbladder contraction'
                ]
            },
            {
                'main': 'Mechanical regulation',
                'subs': [
                    'Distension of intestinal wall triggers stretch reflexes',
                    'Smooth muscle contractions propel contents forward'
                ]
            }
        ]
    )

    doc.add_paragraph()

    # NO "TABLE 2:" prefix - descriptive title only
    doc.add_heading('Neural Regulation Details', 3)
    table2 = doc.add_table(rows=3, cols=3)
    table2.style = 'Table Grid'

    headers = ['Division', 'Neurotransmitters', 'Effects']
    for i, header_text in enumerate(headers):
        set_cell_background(table2.rows[0].cells[i], 'C8E6C9')
        set_cell_text(table2.rows[0].cells[i], header_text, bold=True, size=12)

    # Data rows with bulleted cells (3+ items use bullets with line breaks)
    row_data = [
        ['Parasympathetic', 'Acetylcholine', '• ↑ Motility\n• ↑ Secretion\n• ↑ Blood flow'],
        ['Sympathetic', 'Norepinephrine', '• ↓ Motility\n• ↓ Secretion\n• ↓ Blood flow']
    ]

    for row_idx, row_content in enumerate(row_data, start=1):
        cells = table2.rows[row_idx].cells
        set_cell_background(cells[0], 'E8F5E9')
        set_cell_text(cells[0], row_content[0], bold=True, size=12)
        for col_idx in range(1, 3):
            set_cell_background(cells[col_idx], 'FFFFFF')
            set_cell_text(cells[col_idx], row_content[col_idx], size=12)

    doc.add_paragraph()

    # Clinical Pearls with proper indent (0.5")
    add_plain_bullets(doc, [
        'Vagal stimulation during eating initiates motility ("rest and digest")',
        'Sympathetic activation during stress shuts down digestion ("fight or flight")',
        'Loss of myenteric plexus → intestinal dysmotility (e.g., Hirschsprung disease)'
    ], bold_title='Clinical Pearls & High-Yield Points:')

    doc.add_page_break()

    # ==========================================================================
    # REFERENCE TABLES SECTION (NO "Learning Objective X" HEADINGS)
    # ==========================================================================
    ref_tables_heading = doc.add_heading('Reference Tables', 1)
    ref_tables_heading.runs[0].font.color.rgb = RGBColor(118, 75, 162)

    doc.add_paragraph()

    # NOTE: Tables listed directly WITHOUT "Learning Objective 1", "Learning Objective 2" headings

    # NO "TABLE X:" prefix - descriptive title only
    doc.add_heading('Beta Blocker Comparison', 3)
    table3 = doc.add_table(rows=5, cols=3)
    table3.style = 'Table Grid'

    headers = ['Property', 'Cardioselective', 'Non-Selective']
    for i, header_text in enumerate(headers):
        set_cell_background(table3.rows[0].cells[i], 'FFE0B2')
        set_cell_text(table3.rows[0].cells[i], header_text, bold=True, size=12)

    row_data = [
        ['Examples', '• Metoprolol\n• Atenolol\n• Bisoprolol', '• Propranolol\n• Nadolol\n• Timolol'],
        ['Target', 'β₁ (heart)', 'β₁ + β₂'],
        ['Heart Effects', '↓ HR, ↓ contractility', '↓ HR, ↓ contractility'],
        ['Lung Effects', 'Minimal', 'Bronchospasm risk']
    ]

    for row_idx, row_content in enumerate(row_data, start=1):
        cells = table3.rows[row_idx].cells
        set_cell_background(cells[0], 'FFF3E0')
        set_cell_text(cells[0], row_content[0], bold=True, size=12)
        for col_idx in range(1, 3):
            set_cell_background(cells[col_idx], 'FFFFFF')
            set_cell_text(cells[col_idx], row_content[col_idx], size=12)

    doc.add_paragraph()

    # NO "TABLE X:" prefix - descriptive title only
    doc.add_heading('Autonomic Control of Digestion', 3)
    table4 = doc.add_table(rows=3, cols=4)
    table4.style = 'Table Grid'

    headers = ['System', 'Neurotransmitter', 'Motility', 'Secretion']
    for i, header_text in enumerate(headers):
        set_cell_background(table4.rows[0].cells[i], 'D1C4E9')
        set_cell_text(table4.rows[0].cells[i], header_text, bold=True, size=12)

    row_data = [
        ['Parasympathetic', 'Acetylcholine', '↑ Increased', '↑ Increased'],
        ['Sympathetic', 'Norepinephrine', '↓ Decreased', '↓ Decreased']
    ]

    for row_idx, row_content in enumerate(row_data, start=1):
        cells = table4.rows[row_idx].cells
        set_cell_background(cells[0], 'EDE7F6')
        set_cell_text(cells[0], row_content[0], bold=True, size=12)
        for col_idx in range(1, 4):
            set_cell_background(cells[col_idx], 'FFFFFF')
            set_cell_text(cells[col_idx], row_content[col_idx], size=12)

    # Save document
    doc.save(output_path)
    print(f"✅ CORRECTED sample document created: {output_path}")

# =============================================================================
# MAIN
# =============================================================================

if __name__ == '__main__':
    output_path = '/Users/kimnguyen/Documents/Github/Skills/study-guides/templates-and-examples/Sample_2_LOs_CORRECTED.docx'
    create_sample_document(output_path)
