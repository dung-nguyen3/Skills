#!/usr/bin/env python3
"""
Sample Word LO Document - Formatting Demo
Demonstrates:
1. Simple paragraph summary (LO 1)
2. Structured numbered summary with indented sub-bullets (LO 2)
3. Reference Tables section without "Learning Objective X" headings
"""

from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# =============================================================================
# HELPER FUNCTIONS
# =============================================================================

def set_cell_background(cell, color_hex):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    cell._element.get_or_add_tcPr().append(shading_elm)

def set_cell_text(cell, text, bold=False, size=11):
    """Set cell text with formatting"""
    cell.text = text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(size)
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.bold = bold

def add_colored_box(doc, title, content_list, title_color, bg_color):
    """Add colored box with title and bullet points"""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.left_indent = Pt(12)
    para.paragraph_format.right_indent = Pt(12)

    # Title
    title_run = para.add_run(f'{title}\n')
    title_run.bold = True
    title_run.font.size = Pt(11)
    title_run.font.name = 'Calibri'
    title_run.font.color.rgb = RGBColor(*title_color)

    # Background color
    pPr = para._element.get_or_add_pPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), bg_color)
    pPr.append(shading)

    # Content
    for item in content_list:
        run = para.add_run(f'• {item}\n')
        run.font.size = Pt(11)
        run.font.name = 'Calibri'

def add_summary_simple(doc, summary_text):
    """Add simple paragraph summary with 'Summary:' label"""
    summary_para = doc.add_paragraph()
    summary_run = summary_para.add_run('Summary: ')
    summary_run.bold = True
    summary_run.font.name = 'Calibri'
    summary_run.font.size = Pt(12)

    text_run = summary_para.add_run(summary_text)
    text_run.font.name = 'Calibri'
    text_run.font.size = Pt(12)
    return summary_para

def add_summary_structured(doc, intro_text, items):
    """
    Add structured numbered summary with indented sub-bullets - WITH 'Summary:' label

    Args:
        doc: Document object
        intro_text: Introduction text after 'Summary:' label
        items: List of dicts with structure:
               [
                   {'main': 'Main point 1', 'subs': ['Sub-point A', 'Sub-point B']},
                   {'main': 'Main point 2', 'subs': ['Sub-point C']},
               ]
    """
    # Intro paragraph with 'Summary:' label
    summary_intro = doc.add_paragraph()
    label_run = summary_intro.add_run('Summary: ')
    label_run.bold = True
    label_run.font.name = 'Calibri'
    label_run.font.size = Pt(12)

    intro_run = summary_intro.add_run(intro_text)
    intro_run.font.name = 'Calibri'
    intro_run.font.size = Pt(12)

    # Numbered items with manual numbering
    for idx, item in enumerate(items, start=1):
        # Main numbered item with manual number
        p = doc.add_paragraph()

        # Add bold number
        num_run = p.add_run(f'{idx}. ')
        num_run.bold = True
        num_run.font.size = Pt(12)
        num_run.font.name = 'Calibri'

        # Add main text
        text_run = p.add_run(item['main'])
        text_run.font.size = Pt(12)
        text_run.font.name = 'Calibri'

        # Indent the numbered item
        p.paragraph_format.left_indent = Inches(0.25)

        # Sub-bullets (if any)
        if 'subs' in item and item['subs']:
            for sub in item['subs']:
                sub_p = doc.add_paragraph(sub, style='List Bullet')
                # Indent sub-bullets further below main item (0.5 inches = 36 points)
                sub_p.paragraph_format.left_indent = Inches(0.5)
                sub_p.runs[0].font.name = 'Calibri'
                sub_p.runs[0].font.size = Pt(12)

# =============================================================================
# MAIN DOCUMENT CREATION
# =============================================================================

def create_sample_document(output_path):
    """Create sample Word document with 2 learning objectives"""

    doc = Document()

    # Set margins
    for section in doc.sections:
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)

    # ==========================================================================
    # TITLE PAGE
    # ==========================================================================
    title = doc.add_heading('Formatting Demo', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(118, 75, 162)

    subtitle = doc.add_paragraph('2 Learning Objectives - Demonstrating New Formatting')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(118, 75, 162)

    doc.add_page_break()

    # ==========================================================================
    # LEARNING OBJECTIVES SECTION
    # ==========================================================================

    # --- Learning Objective 1: Simple Summary ---
    obj1 = doc.add_heading('1. Describe the mechanisms of action of beta blockers and their clinical uses', 2)
    obj1.runs[0].font.color.rgb = RGBColor(118, 75, 162)
    obj1.runs[0].font.size = Pt(14)

    # Simple paragraph summary
    add_summary_simple(doc,
        'Beta blockers competitively antagonize β-adrenergic receptors, reducing heart rate, '
        'contractility, and blood pressure. They are classified as selective (β₁) or non-selective '
        '(β₁ and β₂). Clinical uses include hypertension, angina, heart failure, and arrhythmias.')

    doc.add_paragraph()

    # TABLE 1
    doc.add_heading('TABLE 1: Selective vs Non-Selective Beta Blockers', 3)
    table1 = doc.add_table(rows=4, cols=3)
    table1.style = 'Table Grid'

    headers = ['Feature', 'Selective (β₁)', 'Non-Selective (β₁+β₂)']
    for i, header_text in enumerate(headers):
        set_cell_background(table1.rows[0].cells[i], 'B3E5FC')
        set_cell_text(table1.rows[0].cells[i], header_text, bold=True, size=12)

    row_data = [
        ['Examples', 'Metoprolol, Atenolol, Bisoprolol', 'Propranolol, Nadolol, Timolol'],
        ['Target', 'β₁ receptors (heart)', 'β₁ + β₂ receptors'],
        ['Respiratory Effects', 'Minimal bronchospasm risk', '⚠️ Bronchospasm risk']
    ]

    for row_idx, row_content in enumerate(row_data, start=1):
        cells = table1.rows[row_idx].cells
        set_cell_background(cells[0], 'E1F5FE')
        set_cell_text(cells[0], row_content[0], bold=True)
        set_cell_background(cells[1], 'FFFFFF')
        set_cell_text(cells[1], row_content[1])
        set_cell_background(cells[2], 'FFFFFF')
        set_cell_text(cells[2], row_content[2])

    doc.add_paragraph()

    # Clinical Pearls
    add_colored_box(doc, 'Clinical Pearls & High-Yield Points:', [
        'Beta blockers reduce mortality in heart failure (bisoprolol, carvedilol, metoprolol)',
        'Never stop beta blockers abruptly - risk of rebound tachycardia',
        'Cardioselective agents lose selectivity at high doses'
    ], (0, 77, 64), 'E0F2F1')

    doc.add_page_break()

    # --- Learning Objective 2: Structured Summary ---
    obj2 = doc.add_heading('2. Describe the regulation of digestive motility', 2)
    obj2.runs[0].font.color.rgb = RGBColor(118, 75, 162)
    obj2.runs[0].font.size = Pt(14)

    # Structured numbered summary with indented sub-bullets
    add_summary_structured(doc,
        'Digestive motility is regulated by three main mechanisms:',
        [
            {
                'main': 'Neural regulation',
                'subs': [
                    'Parasympathetic (vagus nerve) increases intestinal contraction',
                    'Sympathetic nervous system decreases motility during stress'
                ]
            },
            {
                'main': 'Hormonal regulation',
                'subs': [
                    'Secretin stimulates bicarbonate release from pancreas',
                    'Cholecystokinin (CCK) stimulates enzyme release and gallbladder contraction'
                ]
            },
            {
                'main': 'Mechanical regulation',
                'subs': [
                    'Distension of intestinal wall triggers stretch reflexes',
                    'Smooth muscle contractions propel contents forward'
                ]
            }
        ]
    )

    doc.add_paragraph()

    # TABLE 2: Shows bulleted multi-item cells
    doc.add_heading('TABLE 2: Neural Regulation Details', 3)
    table2 = doc.add_table(rows=3, cols=3)
    table2.style = 'Table Grid'

    headers = ['Division', 'Neurotransmitters', 'Effects']
    for i, header_text in enumerate(headers):
        set_cell_background(table2.rows[0].cells[i], 'C8E6C9')
        set_cell_text(table2.rows[0].cells[i], header_text, bold=True, size=12)

    # Data rows with bulleted cells (3+ items use bullets with line breaks)
    row_data = [
        ['Parasympathetic', 'Acetylcholine', '• ↑ Motility\n• ↑ Secretion\n• ↑ Blood flow'],
        ['Sympathetic', 'Norepinephrine', '• ↓ Motility\n• ↓ Secretion\n• ↓ Blood flow']
    ]

    for row_idx, row_content in enumerate(row_data, start=1):
        cells = table2.rows[row_idx].cells
        set_cell_background(cells[0], 'E8F5E9')
        set_cell_text(cells[0], row_content[0], bold=True)
        for col_idx in range(1, 3):
            set_cell_background(cells[col_idx], 'FFFFFF')
            set_cell_text(cells[col_idx], row_content[col_idx])

    doc.add_paragraph()

    # Clinical Pearls
    add_colored_box(doc, 'Clinical Pearls & High-Yield Points:', [
        'Vagal stimulation during eating initiates motility ("rest and digest")',
        'Sympathetic activation during stress shuts down digestion ("fight or flight")',
        'Loss of myenteric plexus → intestinal dysmotility (e.g., Hirschsprung disease)'
    ], (0, 77, 64), 'E0F2F1')

    doc.add_page_break()

    # ==========================================================================
    # REFERENCE TABLES SECTION (NO "Learning Objective X" HEADINGS)
    # ==========================================================================
    ref_tables_heading = doc.add_heading('Reference Tables', 1)
    ref_tables_heading.runs[0].font.color.rgb = RGBColor(118, 75, 162)

    doc.add_paragraph()

    # NOTE: Tables listed directly WITHOUT "Learning Objective 1", "Learning Objective 2" headings

    # Table from LO 1
    doc.add_heading('Beta Blocker Comparison', 3)
    table3 = doc.add_table(rows=5, cols=3)
    table3.style = 'Table Grid'

    headers = ['Property', 'Cardioselective', 'Non-Selective']
    for i, header_text in enumerate(headers):
        set_cell_background(table3.rows[0].cells[i], 'FFE0B2')
        set_cell_text(table3.rows[0].cells[i], header_text, bold=True, size=12)

    row_data = [
        ['Examples', '• Metoprolol\n• Atenolol\n• Bisoprolol', '• Propranolol\n• Nadolol\n• Timolol'],
        ['Target', 'β₁ (heart)', 'β₁ + β₂'],
        ['Heart Effects', '↓ HR, ↓ contractility', '↓ HR, ↓ contractility'],
        ['Lung Effects', 'Minimal', 'Bronchospasm risk']
    ]

    for row_idx, row_content in enumerate(row_data, start=1):
        cells = table3.rows[row_idx].cells
        set_cell_background(cells[0], 'FFF3E0')
        set_cell_text(cells[0], row_content[0], bold=True)
        for col_idx in range(1, 3):
            set_cell_background(cells[col_idx], 'FFFFFF')
            set_cell_text(cells[col_idx], row_content[col_idx])

    doc.add_paragraph()

    # Table from LO 2
    doc.add_heading('Autonomic Control of Digestion', 3)
    table4 = doc.add_table(rows=3, cols=4)
    table4.style = 'Table Grid'

    headers = ['System', 'Neurotransmitter', 'Motility', 'Secretion']
    for i, header_text in enumerate(headers):
        set_cell_background(table4.rows[0].cells[i], 'D1C4E9')
        set_cell_text(table4.rows[0].cells[i], header_text, bold=True, size=12)

    row_data = [
        ['Parasympathetic', 'Acetylcholine', '↑ Increased', '↑ Increased'],
        ['Sympathetic', 'Norepinephrine', '↓ Decreased', '↓ Decreased']
    ]

    for row_idx, row_content in enumerate(row_data, start=1):
        cells = table4.rows[row_idx].cells
        set_cell_background(cells[0], 'EDE7F6')
        set_cell_text(cells[0], row_content[0], bold=True)
        for col_idx in range(1, 4):
            set_cell_background(cells[col_idx], 'FFFFFF')
            set_cell_text(cells[col_idx], row_content[col_idx])

    # Save document
    doc.save(output_path)
    print(f"✅ Sample document created: {output_path}")

# =============================================================================
# MAIN
# =============================================================================

if __name__ == '__main__':
    output_path = '/Users/kimnguyen/Documents/Github/Skills/study-guides/templates-and-examples/Sample_2_LOs_Formatting_Demo.docx'
    create_sample_document(output_path)
