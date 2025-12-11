#!/usr/bin/env python3
"""
WORD LEARNING OBJECTIVE STUDY GUIDE EXAMPLE
Complete 3-section Word document with separated tables structure

IMPORTANT - INLINE GENERATION APPROACH:
=============================================
This file is a REFERENCE showing structure and styling.

When /LO-word runs, Claude:
  1. Reads source file completely (Read tool)
  2. Extracts ALL content (LOs, topics, comparisons, pearls)
  3. Generates a NEW Python script with extracted data HARDCODED
  4. Saves and executes that script → Word document created

This is NOT a template to modify. It shows what generated scripts should look like.

Compare to /anki: Same pattern - generates Python with hardcoded flashcard data.

Structure:
- Section 1: Learning Objectives (LO statements + summaries + clinical pearls + table links)
- Section 2: Reference Tables (all tables grouped by LO with back-links)
- Section 3: Key Comparisons (side-by-side tables)

See Word_LO_11-5_REVISED.txt for complete template requirements.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# =============================================================================
# HELPER FUNCTIONS
# =============================================================================

def set_cell_background(cell, color):
    """Set cell background color using hex"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def set_cell_text(cell, text, bold=False, size=11):
    """Set cell text - always black for readability"""
    cell.text = text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(size)
            run.font.name = 'Calibri'
            run.font.bold = bold
            # No color parameter - defaults to black

def add_bookmark(paragraph, bookmark_name):
    """Add a bookmark anchor to a paragraph for hyperlinking"""
    # Unique ID for bookmark
    bookmark_id = str(abs(hash(bookmark_name)) % 1000000)

    # Bookmark start
    bookmark_start = OxmlElement('w:bookmarkStart')
    bookmark_start.set(qn('w:id'), bookmark_id)
    bookmark_start.set(qn('w:name'), bookmark_name)

    # Bookmark end
    bookmark_end = OxmlElement('w:bookmarkEnd')
    bookmark_end.set(qn('w:id'), bookmark_id)

    # Insert at beginning and end of paragraph
    paragraph._element.insert(0, bookmark_start)
    paragraph._element.append(bookmark_end)

def add_hyperlink_to_bookmark(paragraph, display_text, bookmark_name, color='0563C1'):
    """Add clickable hyperlink to a bookmark within paragraph"""
    # Create hyperlink element
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), bookmark_name)

    # Create run for the text
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Hyperlink styling (blue, underlined)
    c = OxmlElement('w:color')
    c.set(qn('w:val'), color)
    rPr.append(c)

    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    # Font
    font = OxmlElement('w:rFonts')
    font.set(qn('w:ascii'), 'Calibri')
    rPr.append(font)

    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '24')  # 12pt
    rPr.append(sz)

    new_run.append(rPr)

    # Add text
    text_elem = OxmlElement('w:t')
    text_elem.text = display_text
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._element.append(hyperlink)

def add_plain_bullets(doc, items, bold_title=None):
    """Add plain bulleted list without colored box (for Clinical Pearls)"""
    if bold_title:
        para = doc.add_paragraph()
        run = para.add_run(bold_title)
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(12)

    for item in items:
        para = doc.add_paragraph(item, style='List Bullet')
        # Word default: bullet at 0.25", text at 0.5"
        para.paragraph_format.left_indent = Inches(0.5)
        para.paragraph_format.first_line_indent = Inches(-0.25)  # Hanging indent
        # Universal formatting: 1.5 spacing
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = 1.5
        for run in para.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(12)

def add_summary_simple(doc, summary_text):
    """Add simple paragraph-style summary (1-2 concepts) - NO 'Summary:' label"""
    summary_para = doc.add_paragraph(summary_text)
    # Universal formatting: 1.5 spacing
    summary_para.paragraph_format.space_before = Pt(0)
    summary_para.paragraph_format.space_after = Pt(0)
    summary_para.paragraph_format.line_spacing = 1.5
    summary_para.runs[0].font.name = 'Calibri'
    summary_para.runs[0].font.size = Pt(12)
    return summary_para

def add_summary_structured(doc, intro_text, items):
    """
    Add structured numbered summary with sub-bullets (3+ components) - NO 'Summary:' label
    Uses manual numbering to ensure numbering restarts at 1 for each LO
    Uses Word multilevel list defaults with hanging indents

    Args:
        doc: Document object
        intro_text: Introduction text (no label)
        items: List of dicts with structure:
               [
                   {'main': 'Main point 1', 'subs': ['Sub-point A', 'Sub-point B']},
                   {'main': 'Main point 2', 'subs': ['Sub-point C']},
               ]

    Note: Manual numbering is used instead of 'List Number' style because python-docx
          continues numbering across the document and has no native API to restart.
    """
    # Intro paragraph (no label)
    summary_intro = doc.add_paragraph(intro_text)
    summary_intro.paragraph_format.space_before = Pt(0)
    summary_intro.paragraph_format.space_after = Pt(0)
    summary_intro.paragraph_format.line_spacing = 1.5
    summary_intro.runs[0].font.name = 'Calibri'
    summary_intro.runs[0].font.size = Pt(12)

    # Numbered items with manual numbering
    for idx, item in enumerate(items, start=1):
        # Main numbered item with manual number
        p = doc.add_paragraph()

        # Add bold number
        num_run = p.add_run(f'{idx}. ')
        num_run.bold = True
        num_run.font.size = Pt(12)
        num_run.font.name = 'Calibri'

        # Add main text
        text_run = p.add_run(item['main'])
        text_run.font.size = Pt(12)
        text_run.font.name = 'Calibri'

        # Word default: number at 0", text at 0.25"
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)  # Hanging indent
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.5

        # Sub-bullets (if any)
        if 'subs' in item and item['subs']:
            for sub in item['subs']:
                sub_p = doc.add_paragraph(sub, style='List Bullet')
                # Word default: bullet at 0.25", text at 0.5"
                sub_p.paragraph_format.left_indent = Inches(0.5)
                sub_p.paragraph_format.first_line_indent = Inches(-0.25)  # Hanging indent
                # Universal formatting: 1.5 spacing
                sub_p.paragraph_format.space_before = Pt(0)
                sub_p.paragraph_format.space_after = Pt(0)
                sub_p.paragraph_format.line_spacing = 1.5
                sub_p.runs[0].font.name = 'Calibri'
                sub_p.runs[0].font.size = Pt(12)

# =============================================================================
# MAIN DOCUMENT CREATION
# =============================================================================

def create_word_study_guide(output_path):
    """Create complete Word LO study guide with separated tables structure"""

    doc = Document()

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # ==========================================================================
    # TITLE PAGE WITH TABLE OF CONTENTS
    # ==========================================================================
    title = doc.add_heading('Cardiovascular Pharmacology', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(118, 75, 162)
    title.runs[0].font.size = Pt(20)

    doc.add_paragraph()  # Spacing

    # Table of Contents heading
    toc_heading = doc.add_heading('Table of Contents', 2)
    toc_heading.runs[0].font.color.rgb = RGBColor(118, 75, 162)

    # Learning objectives for TOC (hardcoded example data)
    learning_objectives = [
        'Describe the mechanisms of action of beta blockers and their clinical uses',
        'Describe the regulation of digestive motility'
    ]

    # TOC entries (clickable)
    for idx, lo_text in enumerate(learning_objectives, start=1):
        toc_para = doc.add_paragraph()
        toc_para.paragraph_format.left_indent = Inches(0.25)

        # Number (not hyperlinked)
        num_run = toc_para.add_run(f'{idx}. ')
        num_run.font.name = 'Calibri'
        num_run.font.size = Pt(12)

        # LO text (hyperlinked to bookmark)
        bookmark_name = f'LO_{idx}'
        add_hyperlink_to_bookmark(toc_para, lo_text, bookmark_name, color='0563C1')

    doc.add_page_break()

    # ==========================================================================
    # LEARNING OBJECTIVES (no section heading - starts directly with LO 1)
    # ==========================================================================

    # --- Learning Objective 1 (with bookmark for TOC) ---
    obj1 = doc.add_heading('1. Describe the mechanisms of action of beta blockers and their clinical uses', 2)
    obj1.runs[0].font.color.rgb = RGBColor(118, 75, 162)
    obj1.runs[0].font.size = Pt(14)
    add_bookmark(obj1, 'LO_1')  # Add bookmark for TOC hyperlink

    # Summary - Simple format (1-2 concepts) - NO "Summary:" label
    add_summary_simple(doc,
        'Beta blockers competitively antagonize β-adrenergic receptors, reducing heart rate, '
        'contractility, and blood pressure. They are classified as selective (β₁) or non-selective '
        '(β₁ and β₂). Clinical uses include hypertension, angina, heart failure, and arrhythmias. '
        'Cardioselective agents (metoprolol, atenolol) are preferred in patients with asthma/COPD.')

    # Clinical Pearls (plain bullets, no colored box)
    # NOTE: Only 1 line space between sections within an LO (natural paragraph spacing)
    add_plain_bullets(doc, [
        'Beta blockers reduce mortality in heart failure (bisoprolol, carvedilol, metoprolol)',
        'Never stop beta blockers abruptly - risk of rebound tachycardia and hypertension',
        'Propranolol crosses BBB → useful for performance anxiety and tremor',
        'Cardioselective agents lose selectivity at high doses'
    ], bold_title='Clinical Pearls & High-Yield Points:')

    # Tables (clickable links to Reference Tables section)
    tables_heading = doc.add_paragraph()
    run = tables_heading.add_run('Tables:')
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(12)

    # Table links for LO 1
    for table_idx, table_title in enumerate([
        'Selective vs Non-Selective Beta Blockers'
    ], start=1):
        link_para = doc.add_paragraph()
        link_para.paragraph_format.left_indent = Inches(0.25)

        # Arrow symbol (not hyperlinked)
        arrow_run = link_para.add_run('→ ')
        arrow_run.font.name = 'Calibri'
        arrow_run.font.size = Pt(12)

        # Table title (hyperlinked to table in Reference Tables section)
        bookmark_name = f'LO_1_Table_{table_idx}'
        add_hyperlink_to_bookmark(link_para, table_title, bookmark_name, color='0563C1')

    # 1 line space between Learning Objectives
    doc.add_paragraph()

    # --- Learning Objective 2: STRUCTURED SUMMARY EXAMPLE (with bookmark) ---
    obj2 = doc.add_heading('2. Describe the regulation of digestive motility', 2)
    obj2.runs[0].font.color.rgb = RGBColor(118, 75, 162)
    obj2.runs[0].font.size = Pt(14)
    add_bookmark(obj2, 'LO_2')  # Add bookmark for TOC hyperlink

    # Summary - Structured format (3+ components with hierarchy) - NO "Summary:" label
    # Use when summary lists multiple components/layers/functions
    add_summary_structured(doc,
        'Digestive motility is regulated by three main mechanisms:',
        [
            {
                'main': 'Neural regulation',
                'subs': [
                    'Parasympathetic (vagus nerve) increases intestinal contraction',
                    'Sympathetic nervous system decreases motility during stress'
                ]
            },
            {
                'main': 'Hormonal regulation',
                'subs': [
                    'Secretin stimulates bicarbonate release from pancreas',
                    'Cholecystokinin (CCK) stimulates enzyme release'
                ]
            },
            {
                'main': 'Mechanical regulation',
                'subs': [
                    'Distension of intestinal wall triggers stretch reflexes'
                ]
            }
        ]
    )

    # Clinical Pearls (plain bullets, no colored box)
    # NOTE: Only 1 line space between sections within an LO (natural paragraph spacing)
    add_plain_bullets(doc, [
        'Vagal stimulation during eating initiates motility ("rest and digest")',
        'Sympathetic activation during stress shuts down digestion ("fight or flight")'
    ], bold_title='Clinical Pearls & High-Yield Points:')

    # Tables (clickable links to Reference Tables section)
    tables_heading = doc.add_paragraph()
    run = tables_heading.add_run('Tables:')
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(12)

    # Table links for LO 2
    for table_idx, table_title in enumerate([
        'Neural Regulation Details'
    ], start=1):
        link_para = doc.add_paragraph()
        link_para.paragraph_format.left_indent = Inches(0.25)

        # Arrow symbol (not hyperlinked)
        arrow_run = link_para.add_run('→ ')
        arrow_run.font.name = 'Calibri'
        arrow_run.font.size = Pt(12)

        # Table title (hyperlinked to table in Reference Tables section)
        bookmark_name = f'LO_2_Table_{table_idx}'
        add_hyperlink_to_bookmark(link_para, table_title, bookmark_name, color='0563C1')

    # 1 line space between Learning Objectives
    doc.add_paragraph()

    # ==========================================================================
    # REFERENCE TABLES SECTION (NEW)
    # ==========================================================================
    ref_tables_heading = doc.add_heading('Reference Tables', 1)
    ref_tables_heading.runs[0].font.color.rgb = RGBColor(118, 75, 162)

    doc.add_paragraph()

    # Table 1 for LO 1: Selective vs Non-Selective Beta Blockers
    table_heading_para = doc.add_paragraph()

    # Title run (bold, purple)
    title_run = table_heading_para.add_run('Selective vs Non-Selective Beta Blockers')
    title_run.bold = True
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(14)
    title_run.font.color.rgb = RGBColor(118, 75, 162)

    # Add bookmark to this paragraph
    add_bookmark(table_heading_para, 'LO_1_Table_1')

    # Add spacing, then back-link on same line
    table_heading_para.add_run('    ')  # Spacing

    # Back-link to LO
    add_hyperlink_to_bookmark(table_heading_para, 'LO 1', 'LO_1', color='0563C1')

    # Create the table
    table1 = doc.add_table(rows=5, cols=3)
    table1.style = 'Table Grid'

    # Headers (blue theme for pharmacology)
    headers = ['Feature', 'Selective (β₁)', 'Non-Selective (β₁+β₂)']
    for i, header_text in enumerate(headers):
        set_cell_background(table1.rows[0].cells[i], 'B3E5FC')
        set_cell_text(table1.rows[0].cells[i], header_text, bold=True, size=12)

    # Data rows (supports line breaks with \n and bullet lists)
    row_data = [
        ['Examples', 'Metoprolol, Atenolol, Bisoprolol', 'Propranolol, Nadolol, Timolol'],
        ['Receptor Target', 'β₁ (heart)', 'β₁ (heart) + β₂ (lungs, vessels)'],
        ['Respiratory Effects', 'Minimal bronchospasm risk', '⚠️ Bronchospasm risk'],
        ['Use in Asthma/COPD', '✅ Safer (caution still needed)', '🚫 Avoid - contraindicated'],
    ]

    for row_idx, row_content in enumerate(row_data, start=1):
        cells = table1.rows[row_idx].cells
        set_cell_background(cells[0], 'E1F5FE')  # Light blue
        set_cell_text(cells[0], row_content[0], bold=True, size=11)
        set_cell_background(cells[1], 'FFFFFF')
        set_cell_text(cells[1], row_content[1], size=11)
        set_cell_background(cells[2], 'FFFFFF')
        set_cell_text(cells[2], row_content[2], size=11)

    doc.add_paragraph()  # Spacing after table

    doc.add_paragraph()  # Extra spacing between LO table groups

    # Table 1 for LO 2: Neural Regulation Details
    table_heading_para = doc.add_paragraph()

    # Title run (bold, purple)
    title_run = table_heading_para.add_run('Neural Regulation Details')
    title_run.bold = True
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(14)
    title_run.font.color.rgb = RGBColor(118, 75, 162)

    # Add bookmark to this paragraph
    add_bookmark(table_heading_para, 'LO_2_Table_1')

    # Add spacing, then back-link on same line
    table_heading_para.add_run('    ')  # Spacing

    # Back-link to LO
    add_hyperlink_to_bookmark(table_heading_para, 'LO 2', 'LO_2', color='0563C1')

    # Create the table
    table2 = doc.add_table(rows=3, cols=3)
    table2.style = 'Table Grid'

    headers = ['Division', 'Neurotransmitters', 'Effects']
    for i, header_text in enumerate(headers):
        set_cell_background(table2.rows[0].cells[i], 'C8E6C9')  # Green
        set_cell_text(table2.rows[0].cells[i], header_text, bold=True, size=12)

    # Data rows with bulleted cells (line breaks supported with \n)
    row_data = [
        ['Parasympathetic', 'Acetylcholine', '• ↑ Motility\n• ↑ Secretion\n• ↑ Blood flow'],
        ['Sympathetic', 'Norepinephrine', '• ↓ Motility\n• ↓ Secretion\n• ↓ Blood flow']
    ]

    for row_idx, row_content in enumerate(row_data, start=1):
        cells = table2.rows[row_idx].cells
        set_cell_background(cells[0], 'E8F5E9')  # Light green
        set_cell_text(cells[0], row_content[0], bold=True, size=11)
        for col_idx in range(1, 3):
            set_cell_background(cells[col_idx], 'FFFFFF')
            set_cell_text(cells[col_idx], row_content[col_idx], size=11)

    doc.add_paragraph()  # Spacing after table

    doc.add_paragraph()  # Extra spacing between LO table groups

    # ==========================================================================
    # KEY COMPARISONS (section heading kept - needed for this section)
    # ==========================================================================
    heading2 = doc.add_heading('Key Comparisons', 1)
    heading2.runs[0].font.color.rgb = RGBColor(118, 75, 162)

    # Comparison: Cardioselective vs Non-Selective (no "Comparison 1:" prefix)
    doc.add_heading('Cardioselective vs Non-Selective Beta Blockers', 2)
    table_comp1 = doc.add_table(rows=6, cols=3)
    table_comp1.style = 'Table Grid'

    headers = ['Property', 'Cardioselective (β₁)', 'Non-Selective (β₁+β₂)']
    for i, header_text in enumerate(headers):
        set_cell_background(table_comp1.rows[0].cells[i], 'FFE0B2')  # Orange
        set_cell_text(table_comp1.rows[0].cells[i], header_text, bold=True, size=12)

    row_data = [
        ['Examples', 'Metoprolol, Atenolol, Bisoprolol', 'Propranolol, Nadolol, Carvedilol'],
        ['Target', 'β₁ receptors (heart)', 'β₁ + β₂ receptors'],
        ['Heart Effects', 'Decreased HR, contractility, BP', 'Decreased HR, contractility, BP'],
        ['Lung Effects', 'Minimal', 'Bronchospasm (β₂ blockade)'],
        ['Clinical Use', 'HTN, angina, HF, arrhythmia', 'HTN, anxiety, tremor, migraine'],
    ]

    for row_idx, row_content in enumerate(row_data, start=1):
        cells = table_comp1.rows[row_idx].cells
        set_cell_background(cells[0], 'FFF3E0')
        set_cell_text(cells[0], row_content[0], bold=True, size=11)
        set_cell_background(cells[1], 'FFFFFF')
        set_cell_text(cells[1], row_content[1], size=11)
        set_cell_background(cells[2], 'FFFFFF')
        set_cell_text(cells[2], row_content[2], size=11)

    # Save document
    doc.save(output_path)
    print(f"✅ Word LO study guide created: {output_path}")

# =============================================================================
# MAIN
# =============================================================================

if __name__ == '__main__':
    create_word_study_guide('Cardiovascular_Pharmacology_Study_Guide.docx')
