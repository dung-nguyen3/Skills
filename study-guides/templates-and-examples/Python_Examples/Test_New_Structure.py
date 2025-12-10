#!/usr/bin/env python3
"""
TEST SAMPLE - NEW WORD LO STRUCTURE
Tests all new features:
- No "Summary:" label
- Plain Clinical Pearls (no colored box)
- Related Tables links
- Reference Tables section with grouped tables and back-links
- No page breaks between LOs
- Table cells with line breaks and bullet points
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# =============================================================================
# HELPER FUNCTIONS
# =============================================================================

def set_cell_background(cell, color):
    """Set cell background color using hex"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def set_cell_text(cell, text, bold=False, size=11):
    """
    Set cell text - always black for readability
    Supports line breaks: text with \n renders on multiple lines
    """
    cell.text = text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(size)
            run.font.name = 'Calibri'
            run.bold = bold

def add_plain_bullets(doc, items, bold_title=None):
    """Add plain bulleted list without colored box (for Clinical Pearls)"""
    if bold_title:
        para = doc.add_paragraph()
        run = para.add_run(bold_title)
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(12)

    for item in items:
        para = doc.add_paragraph(item, style='List Bullet')
        # Explicitly set indentation to override style defaults
        para.paragraph_format.left_indent = Inches(0.25)
        para.paragraph_format.first_line_indent = Inches(0)
        for run in para.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(12)

def add_summary_simple(doc, summary_text):
    """Add simple paragraph-style summary (1-2 concepts) - NO 'Summary:' label"""
    summary_para = doc.add_paragraph(summary_text)
    summary_para.runs[0].font.name = 'Calibri'
    summary_para.runs[0].font.size = Pt(12)
    return summary_para

def add_summary_structured(doc, intro_text, items):
    """
    Add structured numbered summary with sub-bullets (3+ components) - NO 'Summary:' label
    Uses manual numbering to ensure numbering restarts at 1 for each LO

    Args:
        doc: Document object
        intro_text: Introduction text (no label)
        items: List of dicts with structure:
               [
                   {'main': 'Main point 1', 'subs': ['Sub-point A', 'Sub-point B']},
                   {'main': 'Main point 2', 'subs': ['Sub-point C']},
               ]
    """
    # Intro paragraph (no label)
    summary_intro = doc.add_paragraph(intro_text)
    summary_intro.runs[0].font.name = 'Calibri'
    summary_intro.runs[0].font.size = Pt(12)

    # Numbered items with manual numbering
    for idx, item in enumerate(items, start=1):
        # Main numbered item with manual number
        p = doc.add_paragraph()

        # Add bold number
        num_run = p.add_run(f'{idx}. ')
        num_run.bold = True
        num_run.font.size = Pt(12)
        num_run.font.name = 'Calibri'

        # Add main text
        text_run = p.add_run(item['main'])
        text_run.font.size = Pt(12)
        text_run.font.name = 'Calibri'

        # Indent the numbered item (consistent with Tables links)
        p.paragraph_format.left_indent = Inches(0.25)

        # Sub-bullets (if any)
        if 'subs' in item and item['subs']:
            for sub in item['subs']:
                sub_p = doc.add_paragraph(sub, style='List Bullet')
                # Indent sub-bullets further below main item
                sub_p.paragraph_format.left_indent = Inches(0.5)

def add_bookmark(paragraph, bookmark_name):
    """Add a bookmark anchor to a paragraph for hyperlinking"""
    bookmark_id = str(abs(hash(bookmark_name)) % 1000000)
    bookmark_start = OxmlElement('w:bookmarkStart')
    bookmark_start.set(qn('w:id'), bookmark_id)
    bookmark_start.set(qn('w:name'), bookmark_name)
    bookmark_end = OxmlElement('w:bookmarkEnd')
    bookmark_end.set(qn('w:id'), bookmark_id)
    paragraph._element.insert(0, bookmark_start)
    paragraph._element.append(bookmark_end)

def add_hyperlink_to_bookmark(paragraph, display_text, bookmark_name, color='0563C1'):
    """Add clickable hyperlink to a bookmark within paragraph"""
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), bookmark_name)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    c = OxmlElement('w:color')
    c.set(qn('w:val'), color)
    rPr.append(c)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    font = OxmlElement('w:rFonts')
    font.set(qn('w:ascii'), 'Calibri')
    rPr.append(font)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '24')
    rPr.append(sz)
    new_run.append(rPr)
    text_elem = OxmlElement('w:t')
    text_elem.text = display_text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._element.append(hyperlink)

# =============================================================================
# TEST DATA
# =============================================================================

learning_objectives = [
    "Identify the basic functions of the GI system.",
    "Describe the structure and layers of the GI tract wall."
]

lo1_summary = "The GI system performs four primary functions: motility (muscular contractions), secretion (enzymes and fluids), digestion (breakdown of nutrients), and absorption (nutrient uptake)."

lo1_clinical_pearls = [
    "Malabsorption can occur from defects in any of the four functions",
    "Most nutrients absorbed in small intestine (duodenum and jejunum)",
    "Fat-soluble vitamin deficiency suggests impaired fat absorption"
]

lo1_table1_title = "GI Functions Overview"
lo1_table1_data = [
    ['Function', 'Description', 'Key Points'],
    ['Motility', 'Muscular contractions moving contents through tract', 'Peristalsis and segmentation'],
    ['Secretion', 'Release of enzymes, acid, mucus, and bile', 'Digestive enzymes\nProtective mucus\nBile salts'],
    ['Digestion', 'Breakdown of macronutrients into absorbable units', ['Carbohydrates → monosaccharides', 'Proteins → amino acids', 'Fats → fatty acids']],
    ['Absorption', 'Uptake of nutrients, water, and electrolytes', 'Primarily in small intestine']
]

lo1_table2_title = "Digestive Breakdown Products"
lo1_table2_data = [
    ['Macronutrient', 'Final Product', 'Absorption Site'],
    ['Carbohydrates', 'Monosaccharides (glucose, fructose, galactose)', 'Small intestine'],
    ['Proteins', 'Amino acids and dipeptides', 'Small intestine'],
    ['Fats', 'Fatty acids and monoglycerides', 'Small intestine']
]

lo2_summary = "The GI tract wall consists of four layers from innermost to outermost: mucosa, submucosa, muscularis externa, and serosa/adventitia."

lo2_clinical_pearls = [
    "Inflammatory bowel diseases primarily affect the mucosa",
    "Achalasia involves dysfunction of the myenteric plexus",
    "Hirschsprung disease results from absence of ganglion cells"
]

lo2_table1_title = "GI Tract Wall Layers"
lo2_table1_data = [
    ['Layer', 'Components', 'Function'],
    ['Mucosa', 'Epithelium, lamina propria, muscularis mucosae', ['Absorption and secretion', 'Immune surveillance', 'Mucosal folding']],
    ['Submucosa', 'Connective tissue, blood vessels, submucosal plexus', 'Blood supply\nNervous control'],
    ['Muscularis Externa', 'Inner circular and outer longitudinal smooth muscle', 'Motility (peristalsis and segmentation)'],
    ['Serosa/Adventitia', 'Connective tissue covering', 'Protection and anchoring']
]

comparison_title = "Peristalsis vs Segmentation"
comparison_data = [
    ['Feature', 'Peristalsis', 'Segmentation'],
    ['Movement Pattern', 'Sequential waves pushing contents forward', 'Alternating contractions mixing contents'],
    ['Primary Function', 'Propulsion through GI tract', 'Mixing and nutrient exposure'],
    ['Location', 'Throughout GI tract', 'Primarily small intestine']
]

# =============================================================================
# DOCUMENT CREATION
# =============================================================================

def create_test_document(output_path):
    """Create test Word document with new structure"""

    doc = Document()

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # =========================================================================
    # PAGE 1: TITLE + TABLE OF CONTENTS
    # =========================================================================

    title = doc.add_heading('Test: New Word LO Structure', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(118, 75, 162)
    title.runs[0].font.size = Pt(20)

    doc.add_paragraph()

    # Table of Contents
    toc_heading = doc.add_heading('Table of Contents', 2)
    toc_heading.runs[0].font.color.rgb = RGBColor(118, 75, 162)

    for idx, lo_text in enumerate(learning_objectives, start=1):
        toc_para = doc.add_paragraph()
        toc_para.paragraph_format.left_indent = Inches(0.25)
        num_run = toc_para.add_run(f'{idx}. ')
        num_run.font.name = 'Calibri'
        num_run.font.size = Pt(12)
        add_hyperlink_to_bookmark(toc_para, lo_text, f'LO_{idx}', color='0563C1')

    doc.add_page_break()

    # =========================================================================
    # LEARNING OBJECTIVES SECTION
    # =========================================================================

    # LO 1
    obj1 = doc.add_heading(f'1. {learning_objectives[0]}', 2)
    obj1.runs[0].font.color.rgb = RGBColor(118, 75, 162)
    obj1.runs[0].font.size = Pt(14)
    add_bookmark(obj1, 'LO_1')

    # Summary - Structured format (lists four primary functions)
    add_summary_structured(doc,
        'The GI system performs four primary functions:',
        [
            {'main': 'Motility - muscular contractions moving contents through the tract'},
            {'main': 'Secretion - release of enzymes, acid, mucus, and bile'},
            {'main': 'Digestion - breakdown of macronutrients into absorbable units'},
            {'main': 'Absorption - uptake of nutrients, water, and electrolytes'}
        ]
    )

    # Clinical Pearls (plain bullets)
    # NOTE: Only 1 line space between sections (natural paragraph spacing)
    add_plain_bullets(doc, lo1_clinical_pearls,
                      bold_title='Clinical Pearls & High-Yield Points:')

    # Tables
    tables_heading = doc.add_paragraph()
    run = tables_heading.add_run('Tables:')
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(12)

    for table_idx, table_title in enumerate([lo1_table1_title, lo1_table2_title], start=1):
        link_para = doc.add_paragraph()
        link_para.paragraph_format.left_indent = Inches(0.25)
        arrow_run = link_para.add_run('→ ')
        arrow_run.font.name = 'Calibri'
        arrow_run.font.size = Pt(12)
        bookmark_name = f'LO_1_Table_{table_idx}'
        add_hyperlink_to_bookmark(link_para, table_title, bookmark_name, color='0563C1')

    # 1 line space between Learning Objectives
    doc.add_paragraph()

    # LO 2
    obj2 = doc.add_heading(f'2. {learning_objectives[1]}', 2)
    obj2.runs[0].font.color.rgb = RGBColor(118, 75, 162)
    obj2.runs[0].font.size = Pt(14)
    add_bookmark(obj2, 'LO_2')

    # Summary - Structured format (lists four layers)
    add_summary_structured(doc,
        'The GI tract wall consists of four layers from innermost to outermost:',
        [
            {'main': 'Mucosa - epithelium, lamina propria, muscularis mucosae'},
            {'main': 'Submucosa - connective tissue, blood vessels, submucosal plexus'},
            {'main': 'Muscularis Externa - inner circular and outer longitudinal smooth muscle'},
            {'main': 'Serosa/Adventitia - connective tissue covering'}
        ]
    )

    # Clinical Pearls (plain bullets)
    # NOTE: Only 1 line space between sections (natural paragraph spacing)
    add_plain_bullets(doc, lo2_clinical_pearls,
                      bold_title='Clinical Pearls & High-Yield Points:')

    # Tables
    tables_heading2 = doc.add_paragraph()
    run2 = tables_heading2.add_run('Tables:')
    run2.bold = True
    run2.font.name = 'Calibri'
    run2.font.size = Pt(12)

    link_para2 = doc.add_paragraph()
    link_para2.paragraph_format.left_indent = Inches(0.25)
    arrow_run2 = link_para2.add_run('→ ')
    arrow_run2.font.name = 'Calibri'
    arrow_run2.font.size = Pt(12)
    add_hyperlink_to_bookmark(link_para2, lo2_table1_title, 'LO_2_Table_1', color='0563C1')

    # 1 line space before Reference Tables section
    doc.add_paragraph()

    # =========================================================================
    # REFERENCE TABLES SECTION
    # =========================================================================

    ref_tables_heading = doc.add_heading('Reference Tables', 1)
    ref_tables_heading.runs[0].font.color.rgb = RGBColor(118, 75, 162)

    doc.add_paragraph()

    # Table 1 for LO 1
    table1_heading_para = doc.add_paragraph()
    title_run1 = table1_heading_para.add_run(lo1_table1_title)
    title_run1.bold = True
    title_run1.font.name = 'Calibri'
    title_run1.font.size = Pt(14)
    title_run1.font.color.rgb = RGBColor(118, 75, 162)
    add_bookmark(table1_heading_para, 'LO_1_Table_1')
    table1_heading_para.add_run('    ')
    add_hyperlink_to_bookmark(table1_heading_para, 'LO 1', 'LO_1', color='0563C1')

    table1 = doc.add_table(rows=len(lo1_table1_data), cols=3)
    table1.style = 'Table Grid'

    # Headers
    for i, header_text in enumerate(lo1_table1_data[0]):
        set_cell_background(table1.rows[0].cells[i], 'B3E5FC')
        set_cell_text(table1.rows[0].cells[i], header_text, bold=True, size=12)

    # Data rows (test line breaks and bullet points)
    for row_idx, row_data in enumerate(lo1_table1_data[1:], start=1):
        for col_idx, cell_data in enumerate(row_data):
            cell = table1.rows[row_idx].cells[col_idx]

            if isinstance(cell_data, list):
                # Bullet points
                cell.text = ''
                for item in cell_data:
                    para = cell.add_paragraph(item, style='List Bullet')
                    para.runs[0].font.name = 'Calibri'
                    para.runs[0].font.size = Pt(11)
            else:
                # Text (may contain \n for line breaks)
                set_cell_text(cell, cell_data, bold=False, size=11)

    doc.add_paragraph()

    # Table 2 for LO 1
    table2_heading_para = doc.add_paragraph()
    title_run2 = table2_heading_para.add_run(lo1_table2_title)
    title_run2.bold = True
    title_run2.font.name = 'Calibri'
    title_run2.font.size = Pt(14)
    title_run2.font.color.rgb = RGBColor(118, 75, 162)
    add_bookmark(table2_heading_para, 'LO_1_Table_2')
    table2_heading_para.add_run('    ')
    add_hyperlink_to_bookmark(table2_heading_para, 'LO 1', 'LO_1', color='0563C1')

    table2 = doc.add_table(rows=len(lo1_table2_data), cols=3)
    table2.style = 'Table Grid'

    for i, header_text in enumerate(lo1_table2_data[0]):
        set_cell_background(table2.rows[0].cells[i], 'B3E5FC')
        set_cell_text(table2.rows[0].cells[i], header_text, bold=True, size=12)

    for row_idx, row_data in enumerate(lo1_table2_data[1:], start=1):
        for col_idx, cell_data in enumerate(row_data):
            set_cell_text(table2.rows[row_idx].cells[col_idx], cell_data, bold=False, size=11)

    doc.add_paragraph()

    # Table 1 for LO 2
    table3_heading_para = doc.add_paragraph()
    title_run3 = table3_heading_para.add_run(lo2_table1_title)
    title_run3.bold = True
    title_run3.font.name = 'Calibri'
    title_run3.font.size = Pt(14)
    title_run3.font.color.rgb = RGBColor(118, 75, 162)
    add_bookmark(table3_heading_para, 'LO_2_Table_1')
    table3_heading_para.add_run('    ')
    add_hyperlink_to_bookmark(table3_heading_para, 'LO 2', 'LO_2', color='0563C1')

    table3 = doc.add_table(rows=len(lo2_table1_data), cols=3)
    table3.style = 'Table Grid'

    for i, header_text in enumerate(lo2_table1_data[0]):
        set_cell_background(table3.rows[0].cells[i], 'B3E5FC')
        set_cell_text(table3.rows[0].cells[i], header_text, bold=True, size=12)

    # Data rows (test line breaks and bullet points)
    for row_idx, row_data in enumerate(lo2_table1_data[1:], start=1):
        for col_idx, cell_data in enumerate(row_data):
            cell = table3.rows[row_idx].cells[col_idx]

            if isinstance(cell_data, list):
                # Bullet points
                cell.text = ''
                for item in cell_data:
                    para = cell.add_paragraph(item, style='List Bullet')
                    para.runs[0].font.name = 'Calibri'
                    para.runs[0].font.size = Pt(11)
            else:
                # Text (may contain \n for line breaks)
                set_cell_text(cell, cell_data, bold=False, size=11)

    doc.add_paragraph()

    # =========================================================================
    # KEY COMPARISONS SECTION
    # =========================================================================

    key_comp_heading = doc.add_heading('Key Comparisons', 1)
    key_comp_heading.runs[0].font.color.rgb = RGBColor(118, 75, 162)

    doc.add_heading(comparison_title, 2)

    comp_table = doc.add_table(rows=len(comparison_data), cols=3)
    comp_table.style = 'Table Grid'

    for i, header_text in enumerate(comparison_data[0]):
        set_cell_background(comp_table.rows[0].cells[i], 'D1C4E9')
        set_cell_text(comp_table.rows[0].cells[i], header_text, bold=True, size=12)

    for row_idx, row_data in enumerate(comparison_data[1:], start=1):
        for col_idx, cell_data in enumerate(row_data):
            set_cell_text(comp_table.rows[row_idx].cells[col_idx], cell_data, bold=False, size=11)

    # Save document
    doc.save(output_path)
    print(f"✅ Test document created: {output_path}")

# =============================================================================
# MAIN
# =============================================================================

if __name__ == '__main__':
    output_path = '/Users/kimnguyen/Documents/Github/Skills/study-guides/templates-and-examples/Test_New_Structure.docx'
    create_test_document(output_path)
