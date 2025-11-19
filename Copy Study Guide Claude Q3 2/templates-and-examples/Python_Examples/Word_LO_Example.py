#!/usr/bin/env python3
"""
WORD LEARNING OBJECTIVE STUDY GUIDE EXAMPLE
Complete 4-section Word document with color-coded tables and boxes

Structure:
- Section 1: Learning Objectives (detailed Q&A for each LO)
- Section 2: Key Comparisons (side-by-side tables)
- Section 3: Master Chart (comprehensive reference)
- Section 4: High-Yield Summary (color-coded boxes)

See Word_LO_11-5_REVISED.txt for complete template requirements.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# =============================================================================
# HELPER FUNCTIONS
# =============================================================================

def set_cell_background(cell, color):
    """Set cell background color using hex"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def set_cell_text(cell, text, bold=False, size=11):
    """Set cell text - always black for readability"""
    cell.text = text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(size)
            run.font.name = 'Calibri'
            run.bold = bold
            # No color parameter - defaults to black

def add_colored_box(doc, title, content_list, title_color, bg_color='F3E5F5'):
    """Add colored information box"""
    # Title
    para = doc.add_paragraph()
    run = para.add_run(title)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(*title_color)

    # Content box
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.3)
    para.paragraph_format.right_indent = Inches(0.3)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)

    # Background color
    pPr = para._element.get_or_add_pPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), bg_color)
    pPr.append(shading)

    for item in content_list:
        run = para.add_run(f'• {item}\n')
        run.font.size = Pt(11)
        run.font.name = 'Calibri'

# =============================================================================
# MAIN DOCUMENT CREATION
# =============================================================================

def create_word_study_guide(output_path):
    """Create complete Word LO study guide"""

    doc = Document()

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # ==========================================================================
    # TITLE PAGE
    # ==========================================================================
    title = doc.add_heading('Cardiovascular Pharmacology', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(118, 75, 162)
    title.runs[0].font.size = Pt(20)

    subtitle = doc.add_paragraph('Lecture 25')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(118, 75, 162)

    doc.add_page_break()

    # ==========================================================================
    # SECTION 1: LEARNING OBJECTIVES
    # ==========================================================================
    heading1 = doc.add_heading('Learning Objectives', 1)
    heading1.runs[0].font.color.rgb = RGBColor(118, 75, 162)

    # --- Learning Objective 1 ---
    obj1 = doc.add_heading('Learning Objective 1:', 2)
    obj1.runs[0].font.color.rgb = RGBColor(118, 75, 162)
    doc.add_paragraph('Describe the mechanisms of action of beta blockers and their clinical uses')

    # Summary
    summary_para = doc.add_paragraph()
    summary_run = summary_para.add_run('Summary: ')
    summary_run.bold = True
    summary_para.add_run('Beta blockers competitively antagonize β-adrenergic receptors, reducing heart rate, contractility, and blood pressure. They are classified as selective (β₁) or non-selective (β₁ and β₂). Clinical uses include hypertension, angina, heart failure, and arrhythmias. Cardioselective agents (metoprolol, atenolol) are preferred in patients with asthma/COPD.')

    doc.add_paragraph()

    # TABLE 1: Comparison Table
    doc.add_heading('TABLE 1: Selective vs Non-Selective Beta Blockers', 3)
    table1 = doc.add_table(rows=5, cols=3)
    table1.style = 'Table Grid'

    # Headers (blue theme for pharmacology)
    headers = ['Feature', 'Selective (β₁)', 'Non-Selective (β₁+β₂)']
    for i, header_text in enumerate(headers):
        set_cell_background(table1.rows[0].cells[i], 'B3E5FC')
        set_cell_text(table1.rows[0].cells[i], header_text, bold=True, size=12)

    # Data rows
    row_data = [
        ['Examples', 'Metoprolol, Atenolol, Bisoprolol', 'Propranolol, Nadolol, Timolol'],
        ['Receptor Target', 'β₁ (heart)', 'β₁ (heart) + β₂ (lungs, vessels)'],
        ['Respiratory Effects', 'Minimal bronchospasm risk', '⚠️ Bronchospasm risk'],
        ['Use in Asthma/COPD', '✅ Safer (caution still needed)', '🚫 Avoid - contraindicated'],
    ]

    for row_idx, row_content in enumerate(row_data, start=1):
        cells = table1.rows[row_idx].cells
        set_cell_background(cells[0], 'E1F5FE')  # Light blue
        set_cell_text(cells[0], row_content[0], bold=True)
        set_cell_background(cells[1], 'FFFFFF')
        set_cell_text(cells[1], row_content[1])
        set_cell_background(cells[2], 'FFFFFF')
        set_cell_text(cells[2], row_content[2])

    doc.add_paragraph()

    # Clinical Pearls Box
    add_colored_box(doc, 'Clinical Pearls & High-Yield Points:', [
        'Beta blockers reduce mortality in heart failure (bisoprolol, carvedilol, metoprolol)',
        'Never stop beta blockers abruptly - risk of rebound tachycardia and hypertension',
        'Propranolol crosses BBB → useful for performance anxiety and tremor',
        'Cardioselective agents lose selectivity at high doses'
    ], (0, 77, 64), 'E0F2F1')

    doc.add_paragraph()

    # Memory Tricks Box
    add_colored_box(doc, 'Memory Tricks & Mnemonics:', [
        '"Beta-1 = Be The One for the heart" (β₁ selective = cardioselective)',
        '"Propranolol = Prop up the brain" (lipophilic, crosses BBB)',
        '"Metoprolol = Metro (city) is selective" (cardioselective)',
        '"LOL drugs = Lowers Over Load" (all beta blockers end in -olol)'
    ], (230, 81, 0), 'FFF3E0')

    doc.add_paragraph()

    # Analogy Box
    add_colored_box(doc, 'Analogy:', [
        'Think of beta blockers like removing the gas pedal from a car. The sympathetic nervous system normally presses the pedal (β-receptors) to speed up the heart. Beta blockers take away that pedal, so the heart can\'t speed up even when adrenaline tries to press it. Selective blockers only remove the heart\'s pedal (β₁), while non-selective blockers remove pedals from both the heart (β₁) and lungs/vessels (β₂).'
    ], (118, 75, 162), 'F3E5F5')

    doc.add_page_break()

    # ==========================================================================
    # SECTION 2: KEY COMPARISONS
    # ==========================================================================
    heading2 = doc.add_heading('Key Comparisons', 1)
    heading2.runs[0].font.color.rgb = RGBColor(118, 75, 162)

    # Comparison 1: Cardioselective vs Non-Selective
    doc.add_heading('Cardioselective vs Non-Selective Beta Blockers', 2)
    table_comp1 = doc.add_table(rows=6, cols=3)
    table_comp1.style = 'Table Grid'

    headers = ['Property', 'Cardioselective (β₁)', 'Non-Selective (β₁+β₂)']
    for i, header_text in enumerate(headers):
        set_cell_background(table_comp1.rows[0].cells[i], 'FFE0B2')  # Orange
        set_cell_text(table_comp1.rows[0].cells[i], header_text, bold=True, size=12)

    row_data = [
        ['Examples', 'Metoprolol, Atenolol, Bisoprolol', 'Propranolol, Nadolol, Carvedilol'],
        ['Target', 'β₁ receptors (heart)', 'β₁ + β₂ receptors'],
        ['Heart Effects', 'Decreased HR, contractility, BP', 'Decreased HR, contractility, BP'],
        ['Lung Effects', 'Minimal', 'Bronchospasm (β₂ blockade)'],
        ['Clinical Use', 'HTN, angina, HF, arrhythmia', 'HTN, anxiety, tremor, migraine'],
    ]

    for row_idx, row_content in enumerate(row_data, start=1):
        cells = table_comp1.rows[row_idx].cells
        set_cell_background(cells[0], 'FFF3E0')
        set_cell_text(cells[0], row_content[0], bold=True)
        set_cell_background(cells[1], 'FFFFFF')
        set_cell_text(cells[1], row_content[1])
        set_cell_background(cells[2], 'FFFFFF')
        set_cell_text(cells[2], row_content[2])

    doc.add_page_break()

    # ==========================================================================
    # SECTION 3: MASTER CHART
    # ==========================================================================
    heading3 = doc.add_heading('Master Chart - Beta Blockers', 1)
    heading3.runs[0].font.color.rgb = RGBColor(118, 75, 162)

    doc.add_heading('Comprehensive Beta Blocker Reference', 2)

    # Master chart table
    table_master = doc.add_table(rows=6, cols=2)
    table_master.style = 'Table Grid'

    # Headers
    headers = ['Beta Blocker', 'Key Characteristics']
    for i, header_text in enumerate(headers):
        set_cell_background(table_master.rows[0].cells[i], 'EDE7F6')  # Light purple
        set_cell_text(table_master.rows[0].cells[i], header_text, bold=True, size=12)

    # Master chart data
    master_data = [
        ('Metoprolol', 'Cardioselective (β₁). Uses: HTN, angina, HF, MI. Metabolism: Hepatic (CYP2D6)'),
        ('Atenolol', 'Cardioselective (β₁). Uses: HTN, angina. Excretion: Renal (adjust in CKD)'),
        ('Bisoprolol', 'Cardioselective (β₁). Uses: HF (mortality benefit). Long half-life'),
        ('Propranolol', 'Non-selective. Lipophilic (crosses BBB). Uses: HTN, anxiety, tremor, migraine'),
        ('Carvedilol', 'Non-selective + α₁ blocker. Uses: HF (mortality benefit). Antioxidant properties'),
    ]

    for row_idx, row_content in enumerate(master_data, start=1):
        cells = table_master.rows[row_idx].cells
        set_cell_background(cells[0], 'F3E5F5')
        set_cell_text(cells[0], row_content[0], bold=True)
        set_cell_background(cells[1], 'FFFFFF')
        set_cell_text(cells[1], row_content[1])

    doc.add_page_break()

    # ==========================================================================
    # SECTION 4: HIGH-YIELD SUMMARY
    # ==========================================================================
    heading4 = doc.add_heading('High-Yield Summary', 1)
    heading4.runs[0].font.color.rgb = RGBColor(118, 75, 162)

    # Mechanism Must-Knows
    add_colored_box(doc, 'MECHANISM - Must Know:', [
        'All beta blockers competitively antagonize β-adrenergic receptors',
        'β₁ receptors: Heart (↑ HR, ↑ contractility)',
        'β₂ receptors: Lungs (bronchodilation), vessels (vasodilation)',
        'Cardioselective = β₁ only (lose selectivity at high doses)',
        'Non-selective = β₁ + β₂ blockade'
    ], (118, 75, 162), 'F3E5F5')

    doc.add_paragraph()

    # Critical Contraindications
    add_colored_box(doc, 'CONTRAINDICATIONS - Never Miss:', [
        '⚠️ Acute decompensated heart failure (can worsen initially)',
        '⚠️ Severe bradycardia (<50 bpm) or heart block',
        '⚠️ Severe asthma/COPD (non-selective agents absolutely contraindicated)',
        '⚠️ Peripheral vascular disease (can worsen claudication)'
    ], (183, 28, 28), 'FFEBEE')

    doc.add_paragraph()

    # Clinical Pearls Summary
    add_colored_box(doc, 'Clinical Pearls - SUMMARY:', [
        'Mortality benefit in HF: Bisoprolol, carvedilol, metoprolol succinate',
        'Post-MI: All patients should receive beta blocker (reduces mortality)',
        'Abrupt withdrawal → rebound tachycardia, hypertension, angina',
        'Propranolol for performance anxiety (crosses BBB)',
        'Lipophilic agents (propranolol, metoprolol) cross BBB',
        'Hydrophilic agents (atenolol, nadolol) renally excreted'
    ], (0, 77, 64), 'E0F2F1')

    doc.add_paragraph()

    # Quick Reference
    add_colored_box(doc, 'QUICK REFERENCE - Drug Selection:', [
        '🟢 HTN + no comorbidities: Any beta blocker',
        '🟢 HTN + asthma/COPD: Cardioselective (metoprolol, atenolol) with caution',
        '🟢 Heart failure: Bisoprolol, carvedilol, or metoprolol succinate',
        '🟢 Anxiety/tremor: Propranolol (crosses BBB)',
        '🟢 Kidney disease: Bisoprolol, metoprolol (hepatic metabolism)'
    ], (27, 94, 32), 'E8F5E9')

    # Save document
    doc.save(output_path)
    print(f"✅ Word LO study guide created: {output_path}")

# =============================================================================
# MAIN
# =============================================================================

if __name__ == '__main__':
    create_word_study_guide('Cardiovascular_Pharmacology_Study_Guide.docx')
