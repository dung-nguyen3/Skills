#!/usr/bin/env python3
"""
BACTERIAL GROWTH AND CULTURE - WORD LEARNING OBJECTIVE STUDY GUIDE
Generated from source file: 2 Bacterial growth and culture micro_text.txt

Structure:
- Section 1: Learning Objectives (detailed Q&A for each LO)
- Section 2: Key Comparisons (side-by-side tables)
- Section 3: Master Chart (comprehensive reference)
- Section 4: High-Yield Summary (color-coded boxes)
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# =============================================================================
# HELPER FUNCTIONS
# =============================================================================

def set_cell_background(cell, color):
    """Set cell background color using hex"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def set_cell_text(cell, text, bold=False, size=11):
    """Set cell text - always black for readability"""
    cell.text = text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(size)
            run.font.name = 'Calibri'
            run.bold = bold

def add_colored_box(doc, title, content_list, title_color, bg_color='F3E5F5'):
    """Add colored information box"""
    # Title
    para = doc.add_paragraph()
    run = para.add_run(title)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(*title_color)

    # Content box
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.3)
    para.paragraph_format.right_indent = Inches(0.3)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)

    # Background color
    pPr = para._element.get_or_add_pPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), bg_color)
    pPr.append(shading)

    for item in content_list:
        run = para.add_run(f'• {item}\n')
        run.font.size = Pt(11)
        run.font.name = 'Calibri'

# =============================================================================
# MAIN DOCUMENT CREATION
# =============================================================================

def create_bacterial_growth_study_guide(output_path):
    """Create complete Word LO study guide for Bacterial Growth and Culture"""

    doc = Document()

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # ==========================================================================
    # TITLE PAGE
    # ==========================================================================
    title = doc.add_heading('Bacterial Growth and Culture', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(118, 75, 162)
    title.runs[0].font.size = Pt(20)

    subtitle = doc.add_paragraph('Session 2 - Microbiology')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(118, 75, 162)

    doc.add_page_break()

    # ==========================================================================
    # SECTION 1: LEARNING OBJECTIVES
    # ==========================================================================
    heading1 = doc.add_heading('Learning Objectives', 1)
    heading1.runs[0].font.color.rgb = RGBColor(118, 75, 162)

    # --- Learning Objective 1 ---
    obj1 = doc.add_heading('Learning Objective 1:', 2)
    obj1.runs[0].font.color.rgb = RGBColor(118, 75, 162)
    doc.add_paragraph('Explain how an individual bacterial cell reproduces')

    # Summary
    summary_para = doc.add_paragraph()
    summary_run = summary_para.add_run('Summary: ')
    summary_run.bold = True
    summary_para.add_run('Bacterial cells reproduce through binary fission, a process of asexual reproduction. Starting with one cell, the bacterium replicates its circular DNA and increases biomass. The DNA copies separate, the plasma membrane forms between them, and a cross-wall (septum) begins to form. Finally, the cell divides completely into two genetically identical daughter cells.')

    doc.add_paragraph()

    # TABLE 1: Binary Fission Steps
    doc.add_heading('TABLE 1: Binary Fission Steps', 3)
    table1 = doc.add_table(rows=5, cols=2)
    table1.style = 'Table Grid'

    # Headers (green for normal/physiological process)
    headers = ['Step', 'Description']
    for i, header_text in enumerate(headers):
        set_cell_background(table1.rows[0].cells[i], 'C8E6C9')
        set_cell_text(table1.rows[0].cells[i], header_text, bold=True, size=12)

    # Data rows
    row_data = [
        ['1. Start', 'One happy little bacterium'],
        ['2. DNA Replication & Growth', 'Replicate DNA and increase biomass'],
        ['3. Segregation', 'Copies of DNA separate, plasma membranes form, cross-wall begins to form'],
        ['4. Division', 'Complete division into 2 new genetically identical cells'],
    ]

    for row_idx, row_content in enumerate(row_data, start=1):
        cells = table1.rows[row_idx].cells
        set_cell_background(cells[0], 'E8F5E9')
        set_cell_text(cells[0], row_content[0], bold=True)
        set_cell_background(cells[1], 'FFFFFF')
        set_cell_text(cells[1], row_content[1])

    doc.add_paragraph()

    # Clinical Pearls Box
    add_colored_box(doc, 'Clinical Pearls & High-Yield Points:', [
        'Binary fission produces genetically identical daughter cells (clones)',
        'E. coli doubling time: ~20 minutes under optimal conditions',
        'M. tuberculosis doubling time: ~24 hours (much slower)',
        'Rapid reproduction allows quick adaptation to environmental changes',
        'Doubling/generation time varies significantly between bacterial species'
    ], (0, 77, 64), 'E0F2F1')

    doc.add_paragraph()

    # Memory Tricks Box
    add_colored_box(doc, 'Memory Tricks & Mnemonics:', [
        '[Researched] "RESS" - Replication, Elongation, Septum formation, Separation',
        '[Researched] "Dogs Really Eat Slowly Split" - DNA replication, Really (elongation), Eat (Septum forms), Slowly Split (separation)',
        'Binary = 2, Fission = splitting (2 cells from 1)',
        'Think "Xerox copying" - exact genetic copies produced'
    ], (230, 81, 0), 'FFF3E0')

    doc.add_paragraph()

    # Analogy Box
    add_colored_box(doc, 'Analogy:', [
        'Binary fission is like photocopying a document. First, you prepare the original (DNA replication). Then you enlarge the machine to handle more paper (cell growth). Next, the copier splits the output into two trays (septum formation). Finally, you take two identical copies from separate trays (cell separation). Each copy is genetically identical to the original, just like bacterial daughter cells.'
    ], (118, 75, 162), 'F3E5F5')

    doc.add_page_break()

    # --- Learning Objective 2 ---
    obj2 = doc.add_heading('Learning Objective 2:', 2)
    obj2.runs[0].font.color.rgb = RGBColor(118, 75, 162)
    doc.add_paragraph('Explain how bacterial growth is measured')

    # Summary
    summary_para = doc.add_paragraph()
    summary_run = summary_para.add_run('Summary: ')
    summary_run.bold = True
    summary_para.add_run('Bacterial growth is measured using three primary methods: (1) Optical Density (OD) measures culture turbidity by light absorbance at 540-600nm; more cells = more absorbance. OD measures all cells (viable and non-viable). (2) Colony-Forming Units (CFU) counts viable cells by diluting culture and plating; each colony grows from one cell. (3) Biomass measures dry weight after washing and drying the sample; measures all cells (viable and non-viable).')

    doc.add_paragraph()

    # TABLE 1: Measurement Methods
    doc.add_heading('TABLE 1: Bacterial Growth Measurement Methods', 3)
    table2 = doc.add_table(rows=4, cols=4)
    table2.style = 'Table Grid'

    # Headers (blue for diagnostic/measurement)
    headers = ['Method', 'What It Measures', 'Viable vs Non-Viable', 'Technique']
    for i, header_text in enumerate(headers):
        set_cell_background(table2.rows[0].cells[i], 'B3E5FC')
        set_cell_text(table2.rows[0].cells[i], header_text, bold=True, size=12)

    # Data rows
    row_data = [
        ['Optical Density (OD)', 'Culture turbidity (light absorbance 540-600nm)', 'ALL cells (viable + non-viable)', 'Spectrophotometer measures absorbance. More cells = more turbidity = higher OD'],
        ['Colony-Forming Unit (CFU)', 'Number of viable cells', 'ONLY viable cells', 'Dilute culture and plate. Each colony from 1 cell. Count colonies'],
        ['Biomass', 'Dry weight of cells', 'ALL cells (viable + non-viable)', 'Sample is washed, dried, and weighed to get dry weight'],
    ]

    for row_idx, row_content in enumerate(row_data, start=1):
        cells = table2.rows[row_idx].cells
        set_cell_background(cells[0], 'E1F5FE')
        set_cell_text(cells[0], row_content[0], bold=True)
        for col_idx in range(1, 4):
            set_cell_background(cells[col_idx], 'FFFFFF')
            set_cell_text(cells[col_idx], row_content[col_idx])

    doc.add_paragraph()

    # Clinical Pearls Box
    add_colored_box(doc, 'Clinical Pearls & High-Yield Points:', [
        'CFU is the gold standard for measuring VIABLE bacteria (clinical labs use this)',
        'OD is faster but cannot distinguish living vs dead cells',
        'If testing antibiotic efficacy, use CFU (not OD) to measure viable cells',
        'Growth curves typically plot log of cell number vs time (logarithmic scale)',
        'Sample = culture; X axis = Time; Y axis = Growth (log of cells, OD, or biomass)'
    ], (0, 77, 64), 'E0F2F1')

    doc.add_paragraph()

    # Memory Tricks Box
    add_colored_box(doc, 'Memory Tricks & Mnemonics:', [
        '"OD = Obvious Density" (you can see turbidity with your eyes)',
        '"CFU = Cells Fully Up" (only living cells stand up and form colonies)',
        '"Biomass = Bye to Mass" (dry everything away except mass)',
        'Remember: OD and Biomass count dead cells too, CFU counts only viable cells'
    ], (230, 81, 0), 'FFF3E0')

    doc.add_paragraph()

    # Analogy Box
    add_colored_box(doc, 'Analogy:', [
        'Measuring bacterial growth is like counting fish in a pond. OD is like measuring how cloudy the water is (more fish = murkier water), but you can\'t tell if fish are alive. CFU is like catching fish one by one and counting only swimming ones (viable). Biomass is like draining the pond and weighing all the fish together (dead or alive). Each method gives different information.'
    ], (118, 75, 162), 'F3E5F5')

    doc.add_page_break()

    # --- Learning Objective 3 ---
    obj3 = doc.add_heading('Learning Objective 3:', 2)
    obj3.runs[0].font.color.rgb = RGBColor(118, 75, 162)
    doc.add_paragraph('Describe the stages of a bacterial growth curve and provide some examples of practical implications of this curve')

    # Summary
    summary_para = doc.add_paragraph()
    summary_run = summary_para.add_run('Summary: ')
    summary_run.bold = True
    summary_para.add_run('The bacterial growth curve has 4 phases: (1) Lag phase: no cell division, only biomass increase for biosynthesis; zero population growth. Length depends on nutrient form/availability and inoculum condition. (2) Log phase: exponential growth with rapid cell division; doubling time determined by division rate; pathogens produce virulence factors; primary metabolites made. (3) Stationary phase: energy and nutrients exhausted; death equals division; zero population growth; resources renewed by cell death; secondary metabolites (antibiotics, pigments) produced. (4) Death/decline phase: insufficient resources; exponential cell death (rate less than log phase growth); determined only by CFU/mL, not OD or biomass; negative population growth.')

    doc.add_paragraph()

    # TABLE 1: Growth Curve Phases
    doc.add_heading('TABLE 1: Bacterial Growth Curve Phases', 3)
    table3 = doc.add_table(rows=5, cols=5)
    table3.style = 'Table Grid'

    # Headers (purple for general topics)
    headers = ['Phase', 'Cell Division', 'Population Growth', 'Key Events', 'Practical Implications']
    for i, header_text in enumerate(headers):
        set_cell_background(table3.rows[0].cells[i], 'D1C4E9')
        set_cell_text(table3.rows[0].cells[i], header_text, bold=True, size=11)

    # Data rows
    row_data = [
        ['Lag', 'NO cell division', 'ZERO (flat)', 'Biosynthesis of macromolecules. Increase in biomass only', 'Adaptation period. Length varies with nutrient availability and inoculum condition'],
        ['Log (Exponential)', 'RAPID division', 'POSITIVE (exponential)', 'Doubling/generation time. Primary metabolites. Pathogens produce virulence factors', 'E. coli: ~20 min doubling. M. tuberculosis: ~24 hr. Best phase for industrial use'],
        ['Stationary', 'Death = Division', 'ZERO (flat)', 'Nutrients exhausted. Resources renewed by cell death. Secondary metabolites (antibiotics, pigments)', 'No change in CFU/mL, OD, or biomass. Antibiotics produced here'],
        ['Death/Decline', 'Death > Division', 'NEGATIVE (declining)', 'Insufficient resources. Exponential death (slower than log growth)', 'Only measured by CFU/mL. OD/biomass stay same (dead cells still present)'],
    ]

    for row_idx, row_content in enumerate(row_data, start=1):
        cells = table3.rows[row_idx].cells
        set_cell_background(cells[0], 'F3E5F5')
        set_cell_text(cells[0], row_content[0], bold=True)
        for col_idx in range(1, 5):
            set_cell_background(cells[col_idx], 'FFFFFF')
            set_cell_text(cells[col_idx], row_content[col_idx], size=10)

    doc.add_paragraph()

    # Clinical Pearls Box
    add_colored_box(doc, 'Clinical Pearls & High-Yield Points:', [
        'Antibiotics work best during LOG phase (actively dividing cells)',
        'Antibiotics produced during STATIONARY phase (secondary metabolites)',
        'Death phase only detected by CFU (viable cell count), not OD or biomass',
        'Virulence factors produced during LOG phase (when pathogens actively multiply)',
        'Industrial production uses log phase cells (most active metabolism)',
        'Late log phase = transition to stationary, secondary metabolites start'
    ], (0, 77, 64), 'E0F2F1')

    doc.add_paragraph()

    # Memory Tricks Box
    add_colored_box(doc, 'Memory Tricks & Mnemonics:', [
        '[Researched] "Lazy Lumberjacks Sleep Daily" - Lag, Log, Stationary, Death',
        '[Researched] "LLSD" - Lag, Log, Stationary, Death',
        'Lag = Lagging behind (no division yet, just preparing)',
        'Log = Logarithmic explosion (exponential growth)',
        'Stationary = Standing still (death = division, zero net growth)',
        'Death = Dying off (not enough food, exponential death)'
    ], (230, 81, 0), 'FFF3E0')

    doc.add_paragraph()

    # Analogy Box
    add_colored_box(doc, 'Analogy:', [
        'Think of bacterial growth like a restaurant opening day. Lag phase = staff preparing kitchen before customers arrive (no sales yet, just setup). Log phase = lunch rush with customers flooding in (exponential orders). Stationary phase = evening equilibrium where new customers equal leaving customers (steady state). Death phase = closing time when customers leave faster than new ones arrive (declining). Each phase has unique characteristics and timing.'
    ], (118, 75, 162), 'F3E5F5')

    doc.add_page_break()

    # --- Learning Objective 4 ---
    obj4 = doc.add_heading('Learning Objective 4:', 2)
    obj4.runs[0].font.color.rgb = RGBColor(118, 75, 162)
    doc.add_paragraph('Correlate fermentation and respiration with oxygen requirements')

    # Summary
    summary_para = doc.add_paragraph()
    summary_run = summary_para.add_run('Summary: ')
    summary_run.bold = True
    summary_para.add_run('Bacterial oxygen requirements correlate with metabolic pathways used. Obligate aerobes require oxygen (~20%) for aerobic respiration. Microaerophiles require reduced oxygen (1-10%) for respiration. Facultative aerobes/anaerobes use oxygen if available (aerobic respiration preferred) but can use fermentation or anaerobic respiration without oxygen. Aerotolerant anaerobes use fermentation only; do not utilize oxygen but tolerate its presence. Obligate anaerobes find oxygen toxic and die; they use fermentation or anaerobic respiration only; lack enzymes (superoxide dismutase, catalase, peroxidase) to counteract toxic reactive oxygen species (ROS).')

    doc.add_paragraph()

    # TABLE 1: Oxygen Requirements
    doc.add_heading('TABLE 1: Bacterial Oxygen Requirements and Metabolism', 3)
    table4 = doc.add_table(rows=6, cols=4)
    table4.style = 'Table Grid'

    # Headers (orange for classification)
    headers = ['Group', 'O2 Requirement', 'Enzymes for ROS?', 'Metabolic Pathways']
    for i, header_text in enumerate(headers):
        set_cell_background(table4.rows[0].cells[i], 'FFE0B2')
        set_cell_text(table4.rows[0].cells[i], header_text, bold=True, size=12)

    # Data rows
    row_data = [
        ['Obligate/Strict Aerobe', 'Required (~20%)', 'Yes', 'Aerobic respiration only'],
        ['Microaerophile', 'Required, but reduced levels (1-10%)', 'Yes', 'Aerobic respiration (at low O2)'],
        ['Facultative Aerobe/Anaerobe', 'Not required. Used if available', 'Yes', 'Aerobic respiration (with O2). Fermentation or anaerobic respiration (without O2)'],
        ['Aerotolerant Anaerobe', 'Not required. Not utilized', 'Most yes', 'Fermentation only (indifferent to O2 presence)'],
        ['Obligate/Strict Anaerobe', 'TOXIC', 'NO', 'Fermentation or anaerobic respiration only'],
    ]

    for row_idx, row_content in enumerate(row_data, start=1):
        cells = table4.rows[row_idx].cells
        set_cell_background(cells[0], 'FFF3E0')
        set_cell_text(cells[0], row_content[0], bold=True, size=11)
        for col_idx in range(1, 4):
            set_cell_background(cells[col_idx], 'FFFFFF')
            set_cell_text(cells[col_idx], row_content[col_idx], size=10)

    doc.add_paragraph()

    # Clinical Pearls Box
    add_colored_box(doc, 'Clinical Pearls & High-Yield Points:', [
        'Obligate anaerobes lack enzymes to neutralize ROS, so oxygen is lethal',
        'Facultative anaerobes prefer aerobic respiration (more ATP) but can switch',
        'Aminoglycosides do not work on anaerobes (require O2 to enter cells)',
        'Aerotolerant anaerobes have ROS enzymes but do not use oxygen for metabolism',
        'Examples: Aerobes (Mycobacterium tuberculosis), Anaerobes (Clostridium, Bacteroides)'
    ], (0, 77, 64), 'E0F2F1')

    doc.add_paragraph()

    # Memory Tricks Box
    add_colored_box(doc, 'Memory Tricks & Mnemonics:', [
        '[Researched] Obligate AEROBES: "Nagging Pests Must Breathe" - Nocardia, Pseudomonas, Mycobacterium tuberculosis, Bacillus',
        '[Researched] Obligate ANAEROBES: "Can\'t Breathe Fresh Air" - Clostridium, Bacteroides, Fusobacterium, Actinomyces',
        '[Researched] Aminoglycosides need O2 to enter cells, so they do NOT work on anaerobes',
        'Facultative = Flexible (can go either way)',
        'Microaerophile = Micro-oxygen (little bit of oxygen)',
        'Aerotolerant = Tolerates air but doesn\'t care (indifferent)'
    ], (230, 81, 0), 'FFF3E0')

    doc.add_paragraph()

    # Analogy Box
    add_colored_box(doc, 'Analogy:', [
        'Think of bacteria like people with different breathing preferences. Obligate aerobes are like mountain climbers who need full oxygen to function. Microaerophiles are like people who prefer thin mountain air (some oxygen, not too much). Facultative anaerobes are like adaptable people who prefer oxygen but can hold their breath and use fermentation when needed. Aerotolerant anaerobes are like people who can be around oxygen but never breathe it in. Obligate anaerobes are like people allergic to oxygen - it kills them because they lack protective enzymes.'
    ], (118, 75, 162), 'F3E5F5')

    doc.add_page_break()

    # --- Learning Objective 5 ---
    obj5 = doc.add_heading('Learning Objective 5:', 2)
    obj5.runs[0].font.color.rgb = RGBColor(118, 75, 162)
    doc.add_paragraph('Explain what a growth factor is (in relation to fastidious organisms)')

    # Summary
    summary_para = doc.add_paragraph()
    summary_run = summary_para.add_run('Summary: ')
    summary_run.bold = True
    summary_para.add_run('Growth factors are additional nutrients required by fastidious organisms. Fastidious organisms are "high maintenance, picky eaters, fussy" bacteria that require a complex diet and are difficult to grow in vitro. They require enriched media (complex media plus growth factors) and will not generally be found unless specifically looked for. Example: Haemophilus and Neisseria require hemin (factor X) and NAD (factor V) for growth. These organisms will not grow on blood agar plates, only on chocolate agar where red blood cells are lysed to release these growth factors.')

    doc.add_paragraph()

    # TABLE 1: Fastidious Organisms
    doc.add_heading('TABLE 1: Fastidious Organisms and Growth Factor Requirements', 3)
    table5 = doc.add_table(rows=4, cols=4)
    table5.style = 'Table Grid'

    # Headers (teal for special requirements)
    headers = ['Organism', 'Growth Factors Required', 'Media Type', 'Notes']
    for i, header_text in enumerate(headers):
        set_cell_background(table5.rows[0].cells[i], 'B2DFDB')
        set_cell_text(table5.rows[0].cells[i], header_text, bold=True, size=12)

    # Data rows
    row_data = [
        ['Haemophilus', 'Hemin (Factor X) + NAD (Factor V)', 'Chocolate agar', 'Will NOT grow on blood agar (RBCs not lysed)'],
        ['Neisseria', 'Vitamins, amino acids, coenzymes', 'Chocolate agar', 'Sensitive to fatty acids; requires enrichment'],
        ['Fastidious organisms (general)', 'Various growth factors', 'Enriched media (complex + growth factors)', 'Difficult to grow in vitro. Require special conditions'],
    ]

    for row_idx, row_content in enumerate(row_data, start=1):
        cells = table5.rows[row_idx].cells
        set_cell_background(cells[0], 'E0F2F1')
        set_cell_text(cells[0], row_content[0], bold=True)
        for col_idx in range(1, 4):
            set_cell_background(cells[col_idx], 'FFFFFF')
            set_cell_text(cells[col_idx], row_content[col_idx], size=10)

    doc.add_paragraph()

    # Clinical Pearls Box
    add_colored_box(doc, 'Clinical Pearls & High-Yield Points:', [
        'Chocolate agar = blood agar heated to 80°C to lyse RBCs and release factors',
        'Haemophilus requires BOTH Factor X (hemin) and Factor V (NAD)',
        'Haemophilus will grow around Staph aureus on blood agar (satellitism) because Staph lyses RBCs',
        'Fastidious = requires enriched media; will be missed on standard media',
        'Enriched media = complex media + growth factors'
    ], (0, 77, 64), 'E0F2F1')

    doc.add_paragraph()

    # Memory Tricks Box
    add_colored_box(doc, 'Memory Tricks & Mnemonics:', [
        '[Researched] Factor X = Hemin (think "X" marks the blood)',
        '[Researched] Factor V = NAD (think "V" for Vitamin-like coenzyme)',
        'Fastidious = Fussy eaters (need special diet)',
        'Chocolate agar = "Cooked blood" (heated to lyse RBCs)',
        'Haemophilus = "Heme-loving" (needs hemin from blood)',
        'Satellitism: Haemophilus satellites around Staph colonies on blood agar'
    ], (230, 81, 0), 'FFF3E0')

    doc.add_paragraph()

    # Analogy Box
    add_colored_box(doc, 'Analogy:', [
        'Fastidious bacteria are like elite athletes with strict dietary requirements. Regular food (complex media) is not enough - they need specialized supplements (growth factors). Haemophilus is like an athlete who needs both iron pills (Factor X/hemin) and B-vitamins (Factor V/NAD) to perform. Without both supplements, they cannot grow. Chocolate agar is like a pre-made meal with all supplements included. Blood agar is like having the supplements locked in capsules - bacteria cannot access them unless another organism (like Staph) breaks them open.'
    ], (118, 75, 162), 'F3E5F5')

    doc.add_page_break()

    # --- Learning Objective 6 ---
    obj6 = doc.add_heading('Learning Objective 6:', 2)
    obj6.runs[0].font.color.rgb = RGBColor(118, 75, 162)
    doc.add_paragraph('Define terms associated with growth requirements (such as obligate, facultative, categories for temperature and pH requirements, etc.)')

    # Summary
    summary_para = doc.add_paragraph()
    summary_run = summary_para.add_run('Summary: ')
    summary_run.bold = True
    summary_para.add_run('Obligate = without an alternative; limited to a narrow range of a given environmental parameter. Facultative = has alternatives; capable of surviving in different conditions. For each environmental parameter (temperature, pH, oxygen, etc.), bacteria have: MIN (minimum value supporting growth), OPT (optimum value for maximum growth), and MAX (maximum value supporting growth). Temperature categories: psychrophile (<20°C), mesophile (20-45°C, most pathogens ~30-37°C), thermophile (>45°C). pH categories: acidophile (low pH), neutralphile/neutrophile (neutral pH ~7, most pathogens), alkaliphile (high pH). Other terms: capnophile (requires high CO2 3-5%, examples: Neisseria, Haemophilus, Helicobacter), halophile (requires high salt), osmophile (requires high sugar/osmolarity).')

    doc.add_paragraph()

    # TABLE 1: Growth Requirement Terms
    doc.add_heading('TABLE 1: Growth Requirement Terminology', 3)
    table6 = doc.add_table(rows=3, cols=2)
    table6.style = 'Table Grid'

    # Headers (purple)
    headers = ['Term', 'Definition']
    for i, header_text in enumerate(headers):
        set_cell_background(table6.rows[0].cells[i], 'D1C4E9')
        set_cell_text(table6.rows[0].cells[i], header_text, bold=True, size=12)

    # Data rows
    row_data = [
        ['Obligate', 'WITHOUT an alternative. Limited to a narrow range of a given environmental parameter'],
        ['Facultative', 'HAS alternatives. Capable of surviving in different conditions'],
    ]

    for row_idx, row_content in enumerate(row_data, start=1):
        cells = table6.rows[row_idx].cells
        set_cell_background(cells[0], 'F3E5F5')
        set_cell_text(cells[0], row_content[0], bold=True)
        set_cell_background(cells[1], 'FFFFFF')
        set_cell_text(cells[1], row_content[1])

    doc.add_paragraph()

    # TABLE 2: Temperature Classifications
    doc.add_heading('TABLE 2: Temperature Classifications', 3)
    table7 = doc.add_table(rows=4, cols=3)
    table7.style = 'Table Grid'

    # Headers (green)
    headers = ['Classification', 'Temperature Range', 'Examples/Notes']
    for i, header_text in enumerate(headers):
        set_cell_background(table7.rows[0].cells[i], 'C8E6C9')
        set_cell_text(table7.rows[0].cells[i], header_text, bold=True, size=12)

    # Data rows
    row_data = [
        ['Psychrophile', '<20°C (cold-loving)', 'Cold environment bacteria'],
        ['Mesophile', '20-45°C (middle-loving)', 'MOST PATHOGENS. Range ~12-42°C. Optimum ~30-37°C'],
        ['Thermophile', '>45°C (heat-loving)', 'Hot spring bacteria (e.g., Bacillus spp.)'],
    ]

    for row_idx, row_content in enumerate(row_data, start=1):
        cells = table7.rows[row_idx].cells
        set_cell_background(cells[0], 'E8F5E9')
        set_cell_text(cells[0], row_content[0], bold=True)
        for col_idx in range(1, 3):
            set_cell_background(cells[col_idx], 'FFFFFF')
            set_cell_text(cells[col_idx], row_content[col_idx])

    doc.add_paragraph()

    # TABLE 3: pH Classifications
    doc.add_heading('TABLE 3: pH Classifications', 3)
    table8 = doc.add_table(rows=4, cols=3)
    table8.style = 'Table Grid'

    # Headers (blue)
    headers = ['Classification', 'pH Range', 'Examples/Notes']
    for i, header_text in enumerate(headers):
        set_cell_background(table8.rows[0].cells[i], 'B3E5FC')
        set_cell_text(table8.rows[0].cells[i], header_text, bold=True, size=12)

    # Data rows
    row_data = [
        ['Acidophile', 'Low pH (acidic)', 'Lactobacillus, Coxiella burnetii'],
        ['Neutralphile (Neutrophile)', 'Neutral pH (~7.0)', 'MOST PATHOGENS. Diagnostic media usually pH 7.0'],
        ['Alkaliphile', 'High pH (alkaline)', 'Bacillus spp., Vibrio cholerae'],
    ]

    for row_idx, row_content in enumerate(row_data, start=1):
        cells = table8.rows[row_idx].cells
        set_cell_background(cells[0], 'E1F5FE')
        set_cell_text(cells[0], row_content[0], bold=True)
        for col_idx in range(1, 3):
            set_cell_background(cells[col_idx], 'FFFFFF')
            set_cell_text(cells[col_idx], row_content[col_idx])

    doc.add_paragraph()

    # TABLE 4: Other Growth Requirements
    doc.add_heading('TABLE 4: Other Environmental Requirements', 3)
    table9 = doc.add_table(rows=4, cols=3)
    table9.style = 'Table Grid'

    # Headers (orange)
    headers = ['Term', 'Requirement', 'Examples/Notes']
    for i, header_text in enumerate(headers):
        set_cell_background(table9.rows[0].cells[i], 'FFE0B2')
        set_cell_text(table9.rows[0].cells[i], header_text, bold=True, size=12)

    # Data rows
    row_data = [
        ['Capnophile', 'High CO2 levels (3-5%)', 'Neisseria, Haemophilus, Helicobacter, Capnocytophaga. Requires candle jar, CO2 incubator'],
        ['Halophile', 'High salt concentration', 'Obligate (extreme 15-30%), moderate (3-15%), slight (0.5-3%). Facultative/halotolerant also exist'],
        ['Osmophile', 'High osmolarity/sugar', 'Require high sugar concentration'],
    ]

    for row_idx, row_content in enumerate(row_data, start=1):
        cells = table9.rows[row_idx].cells
        set_cell_background(cells[0], 'FFF3E0')
        set_cell_text(cells[0], row_content[0], bold=True)
        for col_idx in range(1, 3):
            set_cell_background(cells[col_idx], 'FFFFFF')
            set_cell_text(cells[col_idx], row_content[col_idx], size=10)

    doc.add_paragraph()

    # Clinical Pearls Box
    add_colored_box(doc, 'Clinical Pearls & High-Yield Points:', [
        'Most human pathogens are mesophiles (grow at body temperature 37°C)',
        'Most human pathogens are neutralphiles (grow at pH 7.0)',
        'Diagnostic media for human pathogens usually at pH 7.0 and 37°C',
        'Low temperature → loss of enzyme activity, decreased membrane fluidity',
        'High temperature → denaturing of proteins and enzymes',
        'MIN/OPT/MAX concept applies to ALL environmental parameters',
        'Capnophiles require special housing (candle jar or CO2 incubator)'
    ], (0, 77, 64), 'E0F2F1')

    doc.add_paragraph()

    # Memory Tricks Box
    add_colored_box(doc, 'Memory Tricks & Mnemonics:', [
        'Obligate = "Obey one way" (no alternatives)',
        'Facultative = "Faculty has options" (flexible, has alternatives)',
        'Psychro = "Psycho cold" (loves cold)',
        'Meso = "Middle" (middle temperature, where most pathogens live)',
        'Thermo = "Thermostat high" (heat-loving)',
        'Acido = "Acid" (low pH)',
        'Alkali = "Alkaline" (high pH)',
        'Capno = "Cap needs CO2" (high CO2 requirement)',
        'Halo = "Halo of salt" (salt-loving)'
    ], (230, 81, 0), 'FFF3E0')

    doc.add_paragraph()

    # Analogy Box
    add_colored_box(doc, 'Analogy:', [
        'Think of bacteria like Goldilocks looking for porridge. Each bacterium has its own "just right" conditions (OPT). Some bacteria are obligate - they only eat one type of porridge at one temperature (inflexible). Others are facultative - they can eat hot or cold porridge (flexible). Mesophiles are like humans who prefer room temperature (most pathogens). Psychrophiles love freezer-cold porridge. Thermophiles want boiling-hot porridge. The MIN-OPT-MAX range defines how picky each bacterium is about its environmental conditions.'
    ], (118, 75, 162), 'F3E5F5')

    doc.add_page_break()

    # --- Learning Objective 7 ---
    obj7 = doc.add_heading('Learning Objective 7:', 2)
    obj7.runs[0].font.color.rgb = RGBColor(118, 75, 162)
    doc.add_paragraph('Explain why oxygen is toxic to obligate anaerobes')

    # Summary
    summary_para = doc.add_paragraph()
    summary_run = summary_para.add_run('Summary: ')
    summary_run.bold = True
    summary_para.add_run('Oxygen is toxic to obligate anaerobes because reactions with oxygen often create reactive oxygen species (ROS) that can be toxic to the cell. Without enzymes to counteract these molecules (superoxide dismutase, catalase, peroxidase), ROS can be lethal. Obligate anaerobes lack these protective enzymes, so ROS accumulate and kill the cell. Other bacteria (obligate aerobes, microaerophiles, facultative anaerobes, and most aerotolerant anaerobes) possess these enzymes and can neutralize ROS, allowing them to survive in the presence of oxygen.')

    doc.add_paragraph()

    # TABLE 1: Reactive Oxygen Species and Enzymes
    doc.add_heading('TABLE 1: ROS and Protective Enzymes', 3)
    table10 = doc.add_table(rows=4, cols=3)
    table10.style = 'Table Grid'

    # Headers (red for toxic/pathology)
    headers = ['Enzyme', 'Function', 'Present In']
    for i, header_text in enumerate(headers):
        set_cell_background(table10.rows[0].cells[i], 'FFCDD2')
        set_cell_text(table10.rows[0].cells[i], header_text, bold=True, size=12)

    # Data rows
    row_data = [
        ['Superoxide dismutase', 'Breaks down superoxide radicals (O2-)', 'Obligate aerobes, microaerophiles, facultative, aerotolerant'],
        ['Catalase', 'Breaks down hydrogen peroxide (H2O2)', 'Obligate aerobes, facultative'],
        ['Peroxidase', 'Breaks down peroxides', 'Obligate aerobes, microaerophiles, facultative'],
    ]

    for row_idx, row_content in enumerate(row_data, start=1):
        cells = table10.rows[row_idx].cells
        set_cell_background(cells[0], 'FFEBEE')
        set_cell_text(cells[0], row_content[0], bold=True)
        for col_idx in range(1, 3):
            set_cell_background(cells[col_idx], 'FFFFFF')
            set_cell_text(cells[col_idx], row_content[col_idx])

    doc.add_paragraph()

    # Clinical Pearls Box
    add_colored_box(doc, 'Clinical Pearls & High-Yield Points:', [
        'Obligate anaerobes DIE in oxygen because they lack protective enzymes',
        'ROS = reactive oxygen species (superoxide, hydrogen peroxide, hydroxyl radicals)',
        'Catalase test: Add H2O2 to colonies. Bubbles = catalase present',
        'Obligate anaerobes are catalase NEGATIVE (no bubbles)',
        'Clinical implication: Anaerobic infections from deep tissue, abscesses (low O2)',
        'Examples of obligate anaerobes: Clostridium, Bacteroides, Actinomyces'
    ], (0, 77, 64), 'E0F2F1')

    doc.add_paragraph()

    # Memory Tricks Box
    add_colored_box(doc, 'Memory Tricks & Mnemonics:', [
        '[Researched] Obligate anaerobes "Can\'t Breathe Fresh Air" - Clostridium, Bacteroides, Fusobacterium, Actinomyces',
        'ROS = "Really Overly Scary" molecules (toxic if not neutralized)',
        'Obligate anaerobes = "No enzymes, no oxygen tolerance"',
        '"SCP" enzymes protect from ROS: Superoxide dismutase, Catalase, Peroxidase',
        'Catalase test: H2O2 → bubbles = enzyme present (like pouring peroxide on a wound)',
        'Think: Anaerobes are "unarmed" against oxygen (no protective enzymes)'
    ], (230, 81, 0), 'FFF3E0')

    doc.add_paragraph()

    # Analogy Box
    add_colored_box(doc, 'Analogy:', [
        'Think of reactive oxygen species (ROS) like toxic radiation. Most bacteria have protective "hazmat suits" (enzymes like superoxide dismutase, catalase, peroxidase) that neutralize the radiation. Obligate aerobes have full hazmat suits and thrive in radiation. Facultative bacteria have suits and can turn them on when needed. Obligate anaerobes have NO hazmat suits - they are completely unprotected. When exposed to oxygen, ROS accumulate like radiation poisoning, and without protective enzymes, these bacteria die. They must live in oxygen-free environments to survive.'
    ], (118, 75, 162), 'F3E5F5')

    doc.add_page_break()

    # --- Learning Objective 8 ---
    obj8 = doc.add_heading('Learning Objective 8:', 2)
    obj8.runs[0].font.color.rgb = RGBColor(118, 75, 162)
    doc.add_paragraph('Explain the differences between basic categories of media')

    # Summary
    summary_para = doc.add_paragraph()
    summary_run = summary_para.add_run('Summary: ')
    summary_run.bold = True
    summary_para.add_run('Media categories: (1) Defined media: known quantities of each component; defined carbon and nitrogen sources; eliminates variability; narrower growth range; more expensive. (2) Complex (undefined) media: some components not chemically defined (e.g., yeast extract); no exact formula; broader growth range; less expensive. (3) Enriched media: complex media plus growth factors; used for fastidious organisms. (4) Selective media: inhibits unwanted organisms (e.g., antibiotics, bile salts); only certain organisms able to grow. (5) Differential media: differentiates between organisms; visible change (color, precipitate); does NOT select for specific organisms or inhibit growth. Example: MacConkey agar is both selective (bile salts and crystal violet inhibit Gram+) and differential (lactose fermentation detected by pH indicator neutral red).')

    doc.add_paragraph()

    # TABLE 1: Defined vs Complex Media
    doc.add_heading('TABLE 1: Defined vs Complex Media', 3)
    table11 = doc.add_table(rows=6, cols=3)
    table11.style = 'Table Grid'

    # Headers (teal)
    headers = ['Feature', 'Defined Media', 'Complex (Undefined) Media']
    for i, header_text in enumerate(headers):
        set_cell_background(table11.rows[0].cells[i], 'B2DFDB')
        set_cell_text(table11.rows[0].cells[i], header_text, bold=True, size=12)

    # Data rows
    row_data = [
        ['Composition', 'Known quantities of each component', 'Some components not chemically defined (e.g., yeast extract)'],
        ['Formula', 'Exact formula known', 'No exact formula'],
        ['Carbon/Nitrogen', 'Defined sources', 'Undefined sources'],
        ['Growth Range', 'Usually narrower', 'Usually broader'],
        ['Cost', 'Usually more expensive', 'Usually less expensive'],
    ]

    for row_idx, row_content in enumerate(row_data, start=1):
        cells = table11.rows[row_idx].cells
        set_cell_background(cells[0], 'E0F2F1')
        set_cell_text(cells[0], row_content[0], bold=True)
        for col_idx in range(1, 3):
            set_cell_background(cells[col_idx], 'FFFFFF')
            set_cell_text(cells[col_idx], row_content[col_idx])

    doc.add_paragraph()

    # TABLE 2: Enriched, Selective, Differential Media
    doc.add_heading('TABLE 2: Enriched, Selective, and Differential Media', 3)
    table12 = doc.add_table(rows=4, cols=4)
    table12.style = 'Table Grid'

    # Headers (purple)
    headers = ['Media Type', 'Purpose', 'Components/Mechanism', 'Example']
    for i, header_text in enumerate(headers):
        set_cell_background(table12.rows[0].cells[i], 'D1C4E9')
        set_cell_text(table12.rows[0].cells[i], header_text, bold=True, size=11)

    # Data rows
    row_data = [
        ['Enriched', 'Grow fastidious organisms', 'Complex media + growth factors (hemin, NAD, etc.)', 'Chocolate agar (for Haemophilus, Neisseria)'],
        ['Selective', 'Select FOR specific organisms; inhibit unwanted organisms', 'Contains inhibitors (antibiotics, bile salts, crystal violet)', 'MacConkey agar (bile salts/crystal violet inhibit Gram+)'],
        ['Differential', 'Differentiate/distinguish between organisms', 'Visible change (color, precipitate) based on biochemical properties', 'MacConkey agar (lactose fermentation → pink colonies; non-fermenters → colorless)'],
    ]

    for row_idx, row_content in enumerate(row_data, start=1):
        cells = table12.rows[row_idx].cells
        set_cell_background(cells[0], 'F3E5F5')
        set_cell_text(cells[0], row_content[0], bold=True)
        for col_idx in range(1, 4):
            set_cell_background(cells[col_idx], 'FFFFFF')
            set_cell_text(cells[col_idx], row_content[col_idx], size=10)

    doc.add_paragraph()

    # Clinical Pearls Box
    add_colored_box(doc, 'Clinical Pearls & High-Yield Points:', [
        'MacConkey agar is BOTH selective AND differential (important concept)',
        'Selective media inhibits unwanted organisms; differential media distinguishes organisms',
        'Enriched media does NOT inhibit; it provides extra nutrients for fastidious organisms',
        'Blood agar is differential (hemolysis patterns) but not selective',
        'Chocolate agar is enriched (lysed RBCs provide growth factors)',
        'Interpretation is medium-specific (must know what the medium detects)',
        'Gram+ bacteria inhibited by bile salts and crystal violet in MacConkey'
    ], (0, 77, 64), 'E0F2F1')

    doc.add_paragraph()

    # Memory Tricks Box
    add_colored_box(doc, 'Memory Tricks & Mnemonics:', [
        'Defined = "Definition known" (exact formula)',
        'Complex = "Complicated, no exact formula"',
        'Enriched = "Enriched with extras" (growth factors added)',
        'Selective = "Select the winners, block the losers" (inhibits unwanted)',
        'Differential = "Different colors/appearances" (visual distinction)',
        'MacConkey: "Mac selects Gram- and differentiates lactose fermenters"',
        'Chocolate agar: "Chocolate is enriched" (not selective, just nutrient-rich)'
    ], (230, 81, 0), 'FFF3E0')

    doc.add_paragraph()

    # Analogy Box
    add_colored_box(doc, 'Analogy:', [
        'Think of bacterial media like restaurant menus. Defined media is like a recipe with exact measurements (100g flour, 2 eggs). Complex media is like "grandma\'s recipe" with vague instructions ("some sugar, a pinch of salt"). Enriched media is like a kids menu with extra vitamins added (for picky eaters/fastidious bacteria). Selective media is like a VIP club that only lets certain people in (bouncers = inhibitors). Differential media is like a sorting hat that identifies people by their reactions (lactose fermenters turn pink, others stay colorless). MacConkey agar is a VIP club that also sorts guests by what they drink!'
    ], (118, 75, 162), 'F3E5F5')

    doc.add_page_break()

    # --- Learning Objective 9 ---
    obj9 = doc.add_heading('Learning Objective 9:', 2)
    obj9.runs[0].font.color.rgb = RGBColor(118, 75, 162)
    doc.add_paragraph('Understand the different methods by which bacteria can obtain and use novel genetic information')

    # Summary
    summary_para = doc.add_paragraph()
    summary_run = summary_para.add_run('Summary: ')
    summary_run.bold = True
    summary_para.add_run('Bacteria obtain new DNA through horizontal gene exchange, which can impart positive or negative effects on the cell. Three methods: (1) Transformation: competent bacteria take up fragments of "naked DNA" (free from proteins) from the environment; some bacteria are naturally competent; can occur in soil, water, or other reservoirs; used in cloning. (2) Conjugation: one-way transfer of DNA from donor cell (F+) to recipient cell (F-) through conjugation pilus (direct contact); genes carried on large plasmids; proceeds as rolling circle replication (single-stranded DNA); occurs in almost all bacteria. (3) Transduction: genetic transfer mediated by viruses (bacteriophages); bacterial DNA packaged into phage particles; delivered to new cells through phage infection. Acquired DNA can be integrated into host genome through recombination. Plasmids are linear or circular DNA/RNA that replicate independently; may integrate into host genome; carry genetic information (virulence factors, antibiotic resistance, metabolic enzymes); R-factors (resistance plasmids) carry antibiotic resistance genes.')

    doc.add_paragraph()

    # TABLE 1: Horizontal Gene Exchange Methods
    doc.add_heading('TABLE 1: Methods of Horizontal Gene Exchange', 3)
    table13 = doc.add_table(rows=4, cols=5)
    table13.style = 'Table Grid'

    # Headers (orange)
    headers = ['Method', 'Mechanism', 'DNA Source', 'Requires Contact?', 'Key Features']
    for i, header_text in enumerate(headers):
        set_cell_background(table13.rows[0].cells[i], 'FFE0B2')
        set_cell_text(table13.rows[0].cells[i], header_text, bold=True, size=11)

    # Data rows
    row_data = [
        ['Transformation', 'Uptake of naked DNA from environment', 'Free DNA (lysed cells)', 'NO', 'Some bacteria naturally competent. Used in lab cloning'],
        ['Conjugation', 'Direct transfer via pilus', 'Donor cell (F+) plasmid', 'YES (direct contact)', 'Rolling circle replication. Single-stranded DNA. Most bacteria'],
        ['Transduction', 'Virus-mediated transfer', 'Bacterial DNA in phage', 'NO (via phage)', 'Bacteriophage packages bacterial DNA and delivers to new cell'],
    ]

    for row_idx, row_content in enumerate(row_data, start=1):
        cells = table13.rows[row_idx].cells
        set_cell_background(cells[0], 'FFF3E0')
        set_cell_text(cells[0], row_content[0], bold=True)
        for col_idx in range(1, 5):
            set_cell_background(cells[col_idx], 'FFFFFF')
            set_cell_text(cells[col_idx], row_content[col_idx], size=10)

    doc.add_paragraph()

    # TABLE 2: Plasmids
    doc.add_heading('TABLE 2: Plasmids and R-Factors', 3)
    table14 = doc.add_table(rows=6, cols=2)
    table14.style = 'Table Grid'

    # Headers (pink)
    headers = ['Feature', 'Description']
    for i, header_text in enumerate(headers):
        set_cell_background(table14.rows[0].cells[i], 'F8BBD0')
        set_cell_text(table14.rows[0].cells[i], header_text, bold=True, size=12)

    # Data rows
    row_data = [
        ['Structure', 'Linear or circular section of DNA or RNA'],
        ['Replication', 'Replicates as an independent unit (episomal replication)'],
        ['Integration', 'May integrate into host genome if proper sequences present'],
        ['Genetic Information', 'May be essential OR provide selective advantage (virulence factors, antibiotic resistance, metabolic enzymes)'],
        ['R-Factors (Resistance Plasmids)', 'RTFs carry 1 or more genes for antibiotic resistance. Example: Trimethoprim resistance'],
    ]

    for row_idx, row_content in enumerate(row_data, start=1):
        cells = table14.rows[row_idx].cells
        set_cell_background(cells[0], 'FCE4EC')
        set_cell_text(cells[0], row_content[0], bold=True)
        set_cell_background(cells[1], 'FFFFFF')
        set_cell_text(cells[1], row_content[1])

    doc.add_paragraph()

    # Clinical Pearls Box
    add_colored_box(doc, 'Clinical Pearls & High-Yield Points:', [
        'Horizontal gene transfer is how antibiotic resistance spreads between bacteria',
        'R-factors (resistance plasmids) can carry multiple antibiotic resistance genes',
        'Conjugation requires direct cell-to-cell contact via pilus',
        'Transformation does NOT require contact - bacteria pick up free DNA from environment',
        'Transduction requires bacteriophage as a vector',
        'Plasmids replicate independently of the bacterial chromosome',
        'Conjugation uses rolling circle replication (single-stranded DNA transferred)',
        'Clinical concern: Resistance genes transferred from susceptible bacteria → new resistant bacteria'
    ], (0, 77, 64), 'E0F2F1')

    doc.add_paragraph()

    # Memory Tricks Box
    add_colored_box(doc, 'Memory Tricks & Mnemonics:', [
        '[Researched] "Con"jugation = "Con"tact (direct cell-to-cell contact required)',
        '[Researched] Trans"D"uction = V"D" (Virus/bacteriophage-mediated, think STDs)',
        '[Researched] Trans"F"ormation = "F"ree DNA (naked DNA from environment, no contact)',
        '"Trans" = across; "Con" = with',
        'Transformation: bacteria "transform" by eating free DNA from dead neighbors',
        'Conjugation: bacteria "conjugate" (mate) via pilus bridge',
        'Transduction: virus "introduces" DNA like a delivery service',
        'R-factor: "R" = Resistance (antibiotic resistance plasmid)'
    ], (230, 81, 0), 'FFF3E0')

    doc.add_paragraph()

    # Analogy Box
    add_colored_box(doc, 'Analogy:', [
        'Think of genetic exchange like sharing homework answers. Transformation is like finding someone\'s discarded homework on the ground and copying it (free DNA from environment). Conjugation is like directly passing notes across desks using a pencil bridge (pilus = bridge for direct contact). Transduction is like a teacher accidentally delivering someone else\'s homework to you (virus = delivery service). All three methods spread information (genetic material), and in bacteria, this is how antibiotic resistance spreads through populations - one bacterium solves the problem (develops resistance), then shares the answer (resistance gene) with others.'
    ], (118, 75, 162), 'F3E5F5')

    doc.add_page_break()

    # ==========================================================================
    # SECTION 2: KEY COMPARISONS
    # ==========================================================================
    heading2 = doc.add_heading('Key Comparisons', 1)
    heading2.runs[0].font.color.rgb = RGBColor(118, 75, 162)

    # Comparison 1: Growth Measurement Methods
    doc.add_heading('Comparison 1: Bacterial Growth Measurement Methods', 2)
    table_comp1 = doc.add_table(rows=4, cols=4)
    table_comp1.style = 'Table Grid'

    headers = ['Method', 'Viable vs Non-Viable', 'Speed', 'Clinical Use']
    for i, header_text in enumerate(headers):
        set_cell_background(table_comp1.rows[0].cells[i], 'B3E5FC')
        set_cell_text(table_comp1.rows[0].cells[i], header_text, bold=True, size=12)

    row_data = [
        ['Optical Density (OD)', 'ALL cells (viable + non-viable)', 'FAST (minutes)', 'Research, not clinical'],
        ['Colony-Forming Unit (CFU)', 'ONLY viable cells', 'SLOW (overnight)', 'GOLD STANDARD for clinical labs'],
        ['Biomass', 'ALL cells (viable + non-viable)', 'SLOW (hours)', 'Research, not clinical'],
    ]

    for row_idx, row_content in enumerate(row_data, start=1):
        cells = table_comp1.rows[row_idx].cells
        set_cell_background(cells[0], 'E1F5FE')
        set_cell_text(cells[0], row_content[0], bold=True)
        for col_idx in range(1, 4):
            set_cell_background(cells[col_idx], 'FFFFFF')
            set_cell_text(cells[col_idx], row_content[col_idx])

    doc.add_paragraph()

    # Comparison 2: Oxygen Requirements
    doc.add_heading('Comparison 2: Bacterial Oxygen Requirements', 2)
    table_comp2 = doc.add_table(rows=6, cols=5)
    table_comp2.style = 'Table Grid'

    headers = ['Classification', 'O2 Needed?', 'O2 Toxic?', 'ROS Enzymes?', 'Examples']
    for i, header_text in enumerate(headers):
        set_cell_background(table_comp2.rows[0].cells[i], 'FFE0B2')
        set_cell_text(table_comp2.rows[0].cells[i], header_text, bold=True, size=11)

    row_data = [
        ['Obligate Aerobe', 'YES (required)', 'NO', 'YES', 'Mycobacterium tuberculosis, Pseudomonas, Nocardia'],
        ['Microaerophile', 'YES (low levels)', 'YES (high levels)', 'YES', 'Campylobacter, Helicobacter'],
        ['Facultative Aerobe/Anaerobe', 'NO (preferred)', 'NO', 'YES', 'E. coli, Staphylococcus, Streptococcus'],
        ['Aerotolerant Anaerobe', 'NO (not used)', 'NO', 'Most YES', 'Lactobacillus, Streptococcus'],
        ['Obligate Anaerobe', 'NO', 'YES (lethal)', 'NO', 'Clostridium, Bacteroides, Actinomyces'],
    ]

    for row_idx, row_content in enumerate(row_data, start=1):
        cells = table_comp2.rows[row_idx].cells
        set_cell_background(cells[0], 'FFF3E0')
        set_cell_text(cells[0], row_content[0], bold=True, size=10)
        for col_idx in range(1, 5):
            set_cell_background(cells[col_idx], 'FFFFFF')
            set_cell_text(cells[col_idx], row_content[col_idx], size=9)

    doc.add_paragraph()

    # Comparison 3: Culture Media Types
    doc.add_heading('Comparison 3: Selective vs Differential Media', 2)
    table_comp3 = doc.add_table(rows=3, cols=4)
    table_comp3.style = 'Table Grid'

    headers = ['Media Type', 'Purpose', 'Mechanism', 'Example']
    for i, header_text in enumerate(headers):
        set_cell_background(table_comp3.rows[0].cells[i], 'B2DFDB')
        set_cell_text(table_comp3.rows[0].cells[i], header_text, bold=True, size=12)

    row_data = [
        ['Selective', 'Inhibit unwanted organisms; select FOR specific organisms', 'Contains inhibitors (antibiotics, bile salts, dyes)', 'MacConkey (bile salts inhibit Gram+)'],
        ['Differential', 'Distinguish/differentiate between organisms', 'Visible change (color, precipitate) based on biochemical reaction', 'MacConkey (lactose fermenters = pink, non-fermenters = colorless)'],
    ]

    for row_idx, row_content in enumerate(row_data, start=1):
        cells = table_comp3.rows[row_idx].cells
        set_cell_background(cells[0], 'E0F2F1')
        set_cell_text(cells[0], row_content[0], bold=True)
        for col_idx in range(1, 4):
            set_cell_background(cells[col_idx], 'FFFFFF')
            set_cell_text(cells[col_idx], row_content[col_idx])

    doc.add_paragraph()

    # Comparison 4: Horizontal Gene Transfer Methods
    doc.add_heading('Comparison 4: Transformation vs Conjugation vs Transduction', 2)
    table_comp4 = doc.add_table(rows=4, cols=5)
    table_comp4.style = 'Table Grid'

    headers = ['Method', 'Requires Contact?', 'DNA Source', 'Mechanism', 'Clinical Significance']
    for i, header_text in enumerate(headers):
        set_cell_background(table_comp4.rows[0].cells[i], 'D1C4E9')
        set_cell_text(table_comp4.rows[0].cells[i], header_text, bold=True, size=11)

    row_data = [
        ['Transformation', 'NO', 'Free naked DNA (environment)', 'Competent bacteria uptake DNA from lysed cells', 'Lab cloning; natural in soil/water'],
        ['Conjugation', 'YES (direct)', 'Donor cell (F+) plasmid', 'Pilus bridge transfers DNA via rolling circle', 'Spreads antibiotic resistance between bacteria'],
        ['Transduction', 'NO (via virus)', 'Bacterial DNA in phage', 'Bacteriophage packages & delivers DNA', 'Spreads virulence factors and resistance'],
    ]

    for row_idx, row_content in enumerate(row_data, start=1):
        cells = table_comp4.rows[row_idx].cells
        set_cell_background(cells[0], 'F3E5F5')
        set_cell_text(cells[0], row_content[0], bold=True)
        for col_idx in range(1, 5):
            set_cell_background(cells[col_idx], 'FFFFFF')
            set_cell_text(cells[col_idx], row_content[col_idx], size=10)

    doc.add_page_break()

    # ==========================================================================
    # SECTION 3: MASTER CHART
    # ==========================================================================
    heading3 = doc.add_heading('Master Chart - Bacterial Growth and Culture', 1)
    heading3.runs[0].font.color.rgb = RGBColor(118, 75, 162)

    doc.add_heading('Comprehensive Reference Chart', 2)

    # Master chart table
    table_master = doc.add_table(rows=25, cols=2)
    table_master.style = 'Table Grid'

    # Headers
    headers = ['Concept/Topic', 'Key Details']
    for i, header_text in enumerate(headers):
        set_cell_background(table_master.rows[0].cells[i], 'EDE7F6')
        set_cell_text(table_master.rows[0].cells[i], header_text, bold=True, size=12)

    # Master chart data
    master_data = [
        ('Binary Fission', 'Bacterial reproduction: DNA replication → cell growth → segregation → division into 2 identical cells'),
        ('Doubling/Generation Time', 'Time for one cell division. E. coli: ~20 min. M. tuberculosis: ~24 hr'),
        ('Optical Density (OD)', 'Measures culture turbidity (light absorbance 540-600nm). All cells (viable + non-viable)'),
        ('Colony-Forming Unit (CFU)', 'Dilute and plate culture. Each colony from 1 cell. Measures ONLY viable cells. GOLD STANDARD'),
        ('Biomass', 'Dry weight after washing and drying. Measures all cells (viable + non-viable)'),
        ('Lag Phase', 'No cell division. Biosynthesis, biomass increase. Zero population growth. Length varies with nutrients and inoculum'),
        ('Log/Exponential Phase', 'Rapid exponential division. Primary metabolites. Virulence factors produced. Doubling time'),
        ('Stationary Phase', 'Death = division. Zero population growth. Nutrients exhausted. Secondary metabolites (antibiotics, pigments)'),
        ('Death/Decline Phase', 'Death > division. Negative population growth. Only detected by CFU (not OD/biomass)'),
        ('Obligate Aerobe', 'O2 REQUIRED (~20%). Has ROS enzymes. Examples: Mycobacterium tuberculosis, Pseudomonas, Nocardia'),
        ('Microaerophile', 'O2 required at LOW levels (1-10%). Has ROS enzymes. Examples: Campylobacter, Helicobacter'),
        ('Facultative Aerobe/Anaerobe', 'O2 not required but preferred. Has ROS enzymes. Examples: E. coli, Staphylococcus'),
        ('Aerotolerant Anaerobe', 'O2 not utilized but tolerated. Has ROS enzymes. Examples: Lactobacillus'),
        ('Obligate Anaerobe', 'O2 is TOXIC. NO ROS enzymes. Examples: Clostridium, Bacteroides, Actinomyces'),
        ('Reactive Oxygen Species (ROS)', 'Toxic molecules from O2 reactions. Neutralized by superoxide dismutase, catalase, peroxidase'),
        ('Fastidious Organisms', 'High maintenance, require growth factors. Examples: Haemophilus (X+V factors), Neisseria'),
        ('Growth Factors', 'Additional nutrients for fastidious organisms. Factor X = hemin, Factor V = NAD'),
        ('Defined Media', 'Known composition, defined C/N sources. Narrower growth range. More expensive'),
        ('Complex Media', 'Undefined components (yeast extract). Broader growth range. Less expensive'),
        ('Enriched Media', 'Complex + growth factors. For fastidious organisms. Example: Chocolate agar'),
        ('Selective Media', 'Inhibits unwanted organisms. Example: MacConkey (bile salts/crystal violet inhibit Gram+)'),
        ('Differential Media', 'Differentiates organisms by visible change. Example: MacConkey (lactose fermentation → pink)'),
        ('Transformation', 'Uptake of naked DNA from environment. No contact needed. Used in cloning'),
        ('Conjugation', 'Direct transfer via pilus. F+ → F-. Rolling circle replication. Spreads resistance'),
        ('Transduction', 'Virus-mediated DNA transfer. Bacteriophage packages bacterial DNA. Spreads virulence'),
    ]

    for row_idx, row_content in enumerate(master_data, start=1):
        cells = table_master.rows[row_idx].cells
        set_cell_background(cells[0], 'F3E5F5')
        set_cell_text(cells[0], row_content[0], bold=True, size=11)
        set_cell_background(cells[1], 'FFFFFF')
        set_cell_text(cells[1], row_content[1], size=10)

    doc.add_page_break()

    # ==========================================================================
    # SECTION 4: HIGH-YIELD SUMMARY
    # ==========================================================================
    heading4 = doc.add_heading('High-Yield Summary', 1)
    heading4.runs[0].font.color.rgb = RGBColor(118, 75, 162)

    # Reproduction Must-Knows
    add_colored_box(doc, 'REPRODUCTION - Must Know:', [
        'Binary fission produces 2 genetically identical daughter cells',
        'Steps: DNA replication → elongation → septum formation → separation',
        'Doubling time: E. coli ~20 min, M. tuberculosis ~24 hr',
        'Exponential growth means 1 → 2 → 4 → 8 → 16 (geometric progression)'
    ], (118, 75, 162), 'F3E5F5')

    doc.add_paragraph()

    # Growth Curve Must-Knows
    add_colored_box(doc, 'GROWTH CURVE PHASES - Must Know:', [
        'Lag: No division, only biosynthesis (adaptation period)',
        'Log: Exponential division, virulence factors, primary metabolites',
        'Stationary: Death = division, antibiotics/secondary metabolites produced',
        'Death: Death > division, only detected by CFU (not OD/biomass)',
        'Antibiotics work best during LOG phase (actively dividing cells)'
    ], (118, 75, 162), 'F3E5F5')

    doc.add_paragraph()

    # Oxygen Requirements - Critical
    add_colored_box(doc, 'OXYGEN REQUIREMENTS - Never Miss:', [
        '⚠️ Obligate anaerobes DIE in oxygen (lack ROS enzymes)',
        '⚠️ Facultative anaerobes PREFER oxygen but can survive without it',
        '⚠️ Aminoglycosides do NOT work on anaerobes (need O2 to enter cells)',
        '⚠️ ROS enzymes: Superoxide dismutase, Catalase, Peroxidase',
        '⚠️ Obligate aerobes: Mycobacterium tuberculosis, Pseudomonas, Nocardia',
        '⚠️ Obligate anaerobes: Clostridium, Bacteroides, Actinomyces'
    ], (183, 28, 28), 'FFEBEE')

    doc.add_paragraph()

    # Media Types - Clinical Pearls
    add_colored_box(doc, 'CULTURE MEDIA - Clinical Pearls:', [
        'MacConkey agar is BOTH selective (inhibits Gram+) AND differential (lactose fermentation)',
        'Chocolate agar is enriched (lysed RBCs release X and V factors)',
        'Selective media INHIBITS unwanted organisms (bile salts, antibiotics)',
        'Differential media DISTINGUISHES organisms (color change, precipitate)',
        'Fastidious organisms require enriched media (Haemophilus, Neisseria)',
        'Blood agar: Differential (hemolysis) but NOT selective'
    ], (0, 77, 64), 'E0F2F1')

    doc.add_paragraph()

    # Genetic Exchange - Clinical Significance
    add_colored_box(doc, 'GENETIC EXCHANGE - Clinical Significance:', [
        'Horizontal gene transfer spreads ANTIBIOTIC RESISTANCE between bacteria',
        'Transformation: Free DNA from environment (no contact)',
        'Conjugation: Direct contact via pilus (F+ → F-)',
        'Transduction: Virus-mediated DNA delivery',
        'R-factors (resistance plasmids) carry multiple antibiotic resistance genes',
        'Plasmids replicate independently of bacterial chromosome'
    ], (0, 77, 64), 'E0F2F1')

    doc.add_paragraph()

    # Measurement Methods
    add_colored_box(doc, 'MEASUREMENT METHODS - Quick Reference:', [
        '🟢 CFU = GOLD STANDARD (only viable cells, clinical labs use this)',
        '🟢 OD = FAST but measures all cells (viable + non-viable)',
        '🟢 Biomass = Dry weight (all cells, not clinical)',
        '🟢 Testing antibiotic efficacy? Use CFU (not OD)',
        '🟢 Growth curves plot log scale (logarithmic)'
    ], (27, 94, 32), 'E8F5E9')

    doc.add_paragraph()

    # Environmental Factors
    add_colored_box(doc, 'ENVIRONMENTAL FACTORS - Quick Reference:', [
        'Most human pathogens: Mesophiles (30-37°C) and Neutralphiles (pH 7.0)',
        'Obligate = one way only (inflexible)',
        'Facultative = flexible (has alternatives)',
        'Capnophiles require high CO2 (Neisseria, Haemophilus, Helicobacter)',
        'MIN-OPT-MAX concept applies to ALL environmental parameters'
    ], (27, 94, 32), 'E8F5E9')

    # Save document
    doc.save(output_path)
    print(f"Word LO study guide created: {output_path}")

# =============================================================================
# MAIN
# =============================================================================

if __name__ == '__main__':
    output_dir = '/Users/kimnguyen/Library/Mobile Documents/com~apple~CloudDocs/Documents/Midwestern/Quarter 3/Microbiology /Claude Study Tools/'
    import os
    os.makedirs(output_dir, exist_ok=True)

    output_file = os.path.join(output_dir, 'Bacterial_Growth_Culture_LO_Guide.docx')
    create_bacterial_growth_study_guide(output_file)
    print(f"\nFile saved to: {output_file}")
