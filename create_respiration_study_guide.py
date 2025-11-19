#!/usr/bin/env python3
"""
Create Word LO Study Guide: Lecture 33 - Introduction to Respiration
Source: intro to respiration physio.txt
Template: Word LO 11-5.txt
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_background(cell, color):
    """Set cell background color using hex color string"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def set_cell_text(cell, text, bold=False, color=None, size=11):
    """Set cell text with formatting"""
    cell.text = text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(size)
            run.font.name = 'Calibri'
            run.bold = bold
            if color:
                run.font.color.rgb = RGBColor(*color)

def add_colored_box(doc, title, content_list, title_color, bg_color='F3E5F5'):
    """Add a colored information box"""
    # Add title
    para = doc.add_paragraph()
    run = para.add_run(title)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(*title_color)

    # Add content box
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.3)
    para.paragraph_format.right_indent = Inches(0.3)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)

    # Set background color for the paragraph
    pPr = para._element.get_or_add_pPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), bg_color)
    pPr.append(shading)

    for item in content_list:
        run = para.add_run(f'• {item}\n')
        run.font.size = Pt(11)
        run.font.name = 'Calibri'

# Create document
doc = Document()

# Set margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

# Title
title = doc.add_heading('Introduction to Respiration', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = RGBColor(118, 75, 162)
title.runs[0].font.size = Pt(20)

subtitle = doc.add_paragraph('Lecture 33')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.runs[0]
subtitle_run.font.size = Pt(14)
subtitle_run.font.color.rgb = RGBColor(118, 75, 162)

doc.add_page_break()

# ==========================
# TAB 1: LEARNING OBJECTIVES
# ==========================

heading1 = doc.add_heading('Learning Objectives', 1)
heading1.runs[0].font.color.rgb = RGBColor(118, 75, 162)

# ========== LO1: Steps of External Respiration ==========
obj1 = doc.add_heading('Learning Objective 1:', 2)
obj1.runs[0].font.color.rgb = RGBColor(118, 75, 162)
doc.add_paragraph('Explain the steps of external respiration')

# Summary
summary_para = doc.add_paragraph()
summary_run = summary_para.add_run('Summary: ')
summary_run.bold = True
summary_para.add_run('External respiration involves four major steps that efficiently deliver oxygen from air to tissues and remove carbon dioxide from tissues to air. The four steps are: (1) Ventilation between atmosphere and alveoli, (2) Exchange of O₂ and CO₂ between air in alveoli and blood via diffusion, (3) Transport of O₂ and CO₂ between lungs and tissue (O₂ primarily bound to hemoglobin, CO₂ primarily as bicarbonate), and (4) Exchange of O₂ and CO₂ between blood and tissue via diffusion.')

doc.add_paragraph()

# TABLE 1: Steps of External Respiration
doc.add_heading('TABLE 1: Steps of External Respiration', 3)
table1 = doc.add_table(rows=5, cols=3)
table1.style = 'Table Grid'

# Headers
headers = ['Step', 'Process', 'Key Details']
for i, header_text in enumerate(headers):
    set_cell_background(table1.rows[0].cells[i], 'D1C4E9')  # Purple
    set_cell_text(table1.rows[0].cells[i], header_text, bold=True, color=(255, 255, 255), size=12)

# Data rows
row_data = [
    ['Step 1', 'Ventilation', 'Air movement between atmosphere and alveoli. Ventilation rate can be adjusted to meet body needs for O₂ uptake and CO₂ removal'],
    ['Step 2', 'Gas Exchange (Alveoli-Blood)', 'Diffusion of O₂ and CO₂ between alveolar air and pulmonary blood'],
    ['Step 3', 'Transport', 'O₂ primarily bound to hemoglobin; CO₂ primarily transported as bicarbonate in blood'],
    ['Step 4', 'Gas Exchange (Blood-Tissue)', 'Diffusion of O₂ and CO₂ between systemic blood and tissue cells']
]

for row_idx, row_content in enumerate(row_data, start=1):
    cells = table1.rows[row_idx].cells
    set_cell_background(cells[0], 'F3E5F5')  # Light purple
    set_cell_text(cells[0], row_content[0], bold=True)
    set_cell_background(cells[1], 'FFFFFF')
    set_cell_text(cells[1], row_content[1], bold=True)
    set_cell_background(cells[2], 'FFFFFF')
    set_cell_text(cells[2], row_content[2])

doc.add_paragraph()

# Clinical Pearls
add_colored_box(doc, 'Clinical Pearls & High-Yield Points:', [
    'External respiration = all steps getting O₂ from atmosphere to mitochondria (and CO₂ out). Internal respiration = cellular use of O₂ in mitochondria',
    'Step 2 and 4 both use diffusion - passive process driven by partial pressure gradients',
    'If ventilation fails (Step 1) → hypoxemia and hypercapnia develop quickly',
    'Hemoglobin is critical - without it, blood could only carry ~3 mL O₂/L (vs. normal ~200 mL O₂/L)'
], (0, 77, 64), 'E0F2F1')

doc.add_paragraph()

# Memory Tricks
add_colored_box(doc, 'Memory Tricks & Mnemonics:', [
    '"VETIN" for the 4 steps: V = Ventilation, E = External gas exchange (alveoli-blood), T = Transport in blood, IN = INternal gas exchange (blood-tissue)',
    '"Diffusion happens Twice" - Step 2 (lungs) and Step 4 (tissues)'
], (230, 81, 0), 'FFF3E0')

doc.add_paragraph()

# Analogy
add_colored_box(doc, 'Analogy:', [
    'Think of external respiration like a delivery service: (1) The truck picks up packages from the warehouse (ventilation brings air to alveoli), (2) Packages are loaded onto the truck (O₂ diffuses into blood), (3) The truck drives to destinations (blood transports gases), (4) Packages are delivered to customers (O₂ diffuses to tissues). The empty boxes (CO₂) go back the same route in reverse.'
], (118, 75, 162), 'F3E5F5')

doc.add_page_break()

# ========== LO2: Respiratory Quotient ==========
obj2 = doc.add_heading('Learning Objective 2:', 2)
obj2.runs[0].font.color.rgb = RGBColor(118, 75, 162)
doc.add_paragraph('Describe the respiratory quotient (RQ)')

# Summary
summary_para = doc.add_paragraph()
summary_run = summary_para.add_run('Summary: ')
summary_run.bold = True
summary_para.add_run('The respiratory quotient (RQ) is the ratio of CO₂ produced to O₂ consumed during metabolism (VCO₂/VO₂). RQ varies by macronutrient: RQ = 1.0 for carbohydrates, RQ = 0.7 for fats, and RQ = 0.8 for proteins or mixed diet. For example, if consuming 250 mL/min O₂ and producing 200 mL/min CO₂, RQ = 200/250 = 0.8.')

doc.add_paragraph()

# TABLE 1: Respiratory Quotient by Macronutrient
doc.add_heading('TABLE 1: Respiratory Quotient by Macronutrient', 3)
table2 = doc.add_table(rows=4, cols=3)
table2.style = 'Table Grid'

# Headers
headers = ['Substrate', 'RQ Value', 'Meaning']
for i, header_text in enumerate(headers):
    set_cell_background(table2.rows[0].cells[i], 'C8E6C9')  # Green
    set_cell_text(table2.rows[0].cells[i], header_text, bold=True, color=(255, 255, 255), size=12)

# Data rows
row_data = [
    ['Carbohydrates', '1.0', 'Equal CO₂ produced and O₂ consumed'],
    ['Fats (Lipids)', '0.7', 'Less CO₂ produced per O₂ consumed'],
    ['Proteins (or mixed diet)', '0.8', 'Intermediate value, typical for normal mixed diet']
]

for row_idx, row_content in enumerate(row_data, start=1):
    cells = table2.rows[row_idx].cells
    set_cell_background(cells[0], 'E8F5E9')  # Light green
    set_cell_text(cells[0], row_content[0], bold=True)
    set_cell_background(cells[1], 'FFFFFF')
    set_cell_text(cells[1], row_content[1], bold=True)
    set_cell_background(cells[2], 'FFFFFF')
    set_cell_text(cells[2], row_content[2])

doc.add_paragraph()

# Clinical Pearls
add_colored_box(doc, 'Clinical Pearls & High-Yield Points:', [
    'RQ = 0.8 is the "normal" value used clinically (mixed diet of proteins, fats, and carbohydrates)',
    'RQ tells you what fuel the body is burning: exercising athlete on high-carb diet → RQ closer to 1.0; fasting/ketogenic state → RQ closer to 0.7',
    'Typical resting metabolism: VO₂ = 250 mL/min, VCO₂ = 200 mL/min → RQ = 0.8'
], (0, 77, 64), 'E0F2F1')

doc.add_paragraph()

# Memory Tricks
add_colored_box(doc, 'Memory Tricks & Mnemonics:', [
    '"Perfect 1.0 for Carbs" - Carbohydrates have RQ = 1.0 (the only perfect whole number)',
    '"Fat is 0.7-lean" - Fats have the lowest RQ at 0.7',
    '"Pro-8-in" - Proteins are 0.8',
    'Order from low to high: "Fat (0.7) → Protein (0.8) → Carbs (1.0)"'
], (230, 81, 0), 'FFF3E0')

doc.add_paragraph()

# Analogy
add_colored_box(doc, 'Analogy:', [
    'Think of RQ like fuel efficiency in cars. Fats are like diesel fuel - they\'re energy-dense and need more oxygen to burn completely, producing less exhaust (CO₂) per oxygen used (RQ 0.7). Carbohydrates are like gasoline - they burn "cleaner" with equal oxygen in and CO₂ out (RQ 1.0). Your body is a hybrid that can switch fuels depending on what\'s available.'
], (118, 75, 162), 'F3E5F5')

doc.add_page_break()

# ========== LO3: Organization of Respiratory System ==========
obj3 = doc.add_heading('Learning Objective 3:', 2)
obj3.runs[0].font.color.rgb = RGBColor(118, 75, 162)
doc.add_paragraph('Compare and contrast the components and role of the conducting zone and ventilatory zones of the lung')

# Summary
summary_para = doc.add_paragraph()
summary_run = summary_para.add_run('Summary: ')
summary_run.bold = True
summary_para.add_run('The respiratory system has two functional zones. The conducting zone (generations 0-16) includes trachea, bronchi, bronchioles, and terminal bronchioles - it moves air but does NOT participate in gas exchange, representing anatomical dead space (~150 mL). The respiratory zone (generations 17-23) includes respiratory bronchioles, alveolar ducts, and alveolar sacs - it\'s where gas exchange occurs via diffusion. The conducting zone warms, humidifies, filters air, decreases airflow velocity, and regulates airway resistance. The respiratory zone has ~300 million alveoli per lung with total surface area of 50-100 m² (tennis court size).')

doc.add_paragraph()

# TABLE 1: Conducting Zone vs Respiratory Zone
doc.add_heading('TABLE 1: Conducting Zone vs Respiratory Zone', 3)
table3 = doc.add_table(rows=7, cols=3)
table3.style = 'Table Grid'

# Headers
headers = ['Feature', 'Conducting Zone', 'Respiratory Zone']
for i, header_text in enumerate(headers):
    set_cell_background(table3.rows[0].cells[i], 'B3E5FC')  # Blue
    set_cell_text(table3.rows[0].cells[i], header_text, bold=True, color=(255, 255, 255), size=12)

# Data rows
row_data = [
    ['Structures', 'Trachea, bronchi, bronchioles, terminal bronchioles (generations 0-16)', 'Respiratory bronchioles, alveolar ducts, alveolar sacs (generations 17-23)'],
    ['Primary Function', 'Air transport and conditioning', 'Gas exchange via diffusion'],
    ['Gas Exchange', 'NO gas exchange', 'YES - O₂ and CO₂ diffusion'],
    ['Volume', '~150 mL (anatomical dead space)', 'Alveolar ventilation volume'],
    ['Key Functions', 'Humidification, warming, filtration, ↓ airflow velocity, regulate airway resistance', 'Maximize surface area for diffusion'],
    ['Special Features', 'Smooth muscle, cilia, mucous-secreting goblet cells', '~300 million alveoli, 50-100 m² total surface area, Type I & II alveolar cells']
]

for row_idx, row_content in enumerate(row_data, start=1):
    cells = table3.rows[row_idx].cells
    set_cell_background(cells[0], 'E1F5FE')  # Light blue
    set_cell_text(cells[0], row_content[0], bold=True)
    set_cell_background(cells[1], 'FFFFFF')
    set_cell_text(cells[1], row_content[1])
    set_cell_background(cells[2], 'FFFFFF')
    set_cell_text(cells[2], row_content[2])

doc.add_paragraph()

# TABLE 2: Alveolar Cell Types
doc.add_heading('TABLE 2: Alveolar Cell Types', 3)
table3b = doc.add_table(rows=3, cols=3)
table3b.style = 'Table Grid'

# Headers
headers = ['Cell Type', 'Structure', 'Function']
for i, header_text in enumerate(headers):
    set_cell_background(table3b.rows[0].cells[i], 'C8E6C9')  # Green
    set_cell_text(table3b.rows[0].cells[i], header_text, bold=True, color=(255, 255, 255), size=12)

# Data rows
row_data = [
    ['Type I Alveolar Cell', 'Thin, flat epithelial cells', 'Form alveolar epithelium for gas exchange'],
    ['Type II Alveolar Cell', 'Cuboidal cells', 'Synthesize pulmonary surfactant']
]

for row_idx, row_content in enumerate(row_data, start=1):
    cells = table3b.rows[row_idx].cells
    set_cell_background(cells[0], 'E8F5E9')  # Light green
    set_cell_text(cells[0], row_content[0], bold=True)
    set_cell_background(cells[1], 'FFFFFF')
    set_cell_text(cells[1], row_content[1])
    set_cell_background(cells[2], 'FFFFFF')
    set_cell_text(cells[2], row_content[2])

doc.add_paragraph()

# Clinical Pearls
add_colored_box(doc, 'Clinical Pearls & High-Yield Points:', [
    'Dead space (~150 mL) = conducting zone. This is why tidal volume must be >150 mL to bring fresh air to alveoli',
    'Alveolar-capillary barrier is only 0.5 μm thick - one of thinnest barriers in body, optimized for rapid diffusion',
    'RBCs spend only ~0.75 seconds in pulmonary capillaries, traversing 2-3 alveoli - gas exchange must be FAST',
    'Type II cells (surfactant producers) are critical - without surfactant, alveoli would collapse (think: neonatal respiratory distress syndrome)',
    'Huge surface area (50-100 m²) = tennis court in your chest! This is why lungs are so efficient at gas exchange'
], (0, 77, 64), 'E0F2F1')

doc.add_paragraph()

# Memory Tricks
add_colored_box(doc, 'Memory Tricks & Mnemonics:', [
    '"Conducting = No Gas Exchange, Respiratory = YES Gas Exchange"',
    '"Dead space = 1/3 of tidal volume" (150 mL out of 500 mL)',
    '"Type I = thin for exchange, Type II = produces surfactant (2 = surfactant)"',
    '"The alveoli are the size of a Tennis court" (50-100 m²)'
], (230, 81, 0), 'FFF3E0')

doc.add_paragraph()

# Analogy
add_colored_box(doc, 'Analogy:', [
    'Think of the respiratory system like a tree: The conducting zone is the trunk and branches - they deliver air but don\'t do photosynthesis. The respiratory zone is the leaves - where all the action happens (gas exchange = photosynthesis). Just like a tree has millions of leaves for maximum surface area, your lungs have 300 million alveoli creating a tennis court of surface area for gas exchange.'
], (118, 75, 162), 'F3E5F5')

doc.add_page_break()

# ========== LO4: ANS Control of Airway Resistance ==========
obj4 = doc.add_heading('Learning Objective 4:', 2)
obj4.runs[0].font.color.rgb = RGBColor(118, 75, 162)
doc.add_paragraph('Discuss the role of the autonomic nervous system in controlling airway resistance')

# Summary
summary_para = doc.add_paragraph()
summary_run = summary_para.add_run('Summary: ')
summary_run.bold = True
summary_para.add_run('Smooth muscle in bronchi and bronchioles is controlled by the autonomic nervous system. Sympathetic nervous system (SNS) activation of β₂-adrenergic receptors causes bronchodilation and decreases airway resistance. Parasympathetic nervous system (PNS) activation of M₃ muscarinic receptors causes bronchoconstriction and increases airway resistance. Asthma is caused by inappropriate bronchial smooth muscle constriction and is commonly treated with albuterol (β₂-adrenergic agonist) to cause bronchodilation.')

doc.add_paragraph()

# TABLE 1: ANS Control of Airways
doc.add_heading('TABLE 1: Autonomic Control of Airway Resistance', 3)
table4 = doc.add_table(rows=3, cols=4)
table4.style = 'Table Grid'

# Headers
headers = ['ANS Branch', 'Receptor', 'Effect on Airways', 'Effect on Resistance']
for i, header_text in enumerate(headers):
    set_cell_background(table4.rows[0].cells[i], 'FFE0B2')  # Orange
    set_cell_text(table4.rows[0].cells[i], header_text, bold=True, color=(255, 255, 255), size=12)

# Data rows
row_data = [
    ['Sympathetic (SNS)', 'β₂-adrenergic', 'Bronchodilation', 'Resistance ↓↓'],
    ['Parasympathetic (PNS)', 'M₃ muscarinic', 'Bronchoconstriction', 'Resistance ↑↑']
]

for row_idx, row_content in enumerate(row_data, start=1):
    cells = table4.rows[row_idx].cells
    set_cell_background(cells[0], 'FFF3E0')  # Light orange
    set_cell_text(cells[0], row_content[0], bold=True)
    set_cell_background(cells[1], 'FFFFFF')
    set_cell_text(cells[1], row_content[1])
    set_cell_background(cells[2], 'FFFFFF')
    set_cell_text(cells[2], row_content[2])
    set_cell_background(cells[3], 'FFFFFF')
    set_cell_text(cells[3], row_content[3])

doc.add_paragraph()

# TABLE 2: Asthma
doc.add_heading('TABLE 2: Asthma - Pathophysiology and Treatment', 3)
table4b = doc.add_table(rows=2, cols=3)
table4b.style = 'Table Grid'

# Headers
headers = ['Disease', 'Pathophysiology', 'Treatment']
for i, header_text in enumerate(headers):
    set_cell_background(table4b.rows[0].cells[i], 'FFCDD2')  # Red
    set_cell_text(table4b.rows[0].cells[i], header_text, bold=True, color=(255, 255, 255), size=12)

# Data rows
row_data = [
    ['Asthma', 'Bronchial smooth muscle constriction due to chronic inflammatory condition. Features: mucosal edema, increased mucus, contracted hypertrophied muscle', '🟢 Albuterol (Ventolin) - β₂-adrenergic receptor agonist → bronchodilation']
]

for row_idx, row_content in enumerate(row_data, start=1):
    cells = table4b.rows[row_idx].cells
    set_cell_background(cells[0], 'FFEBEE')  # Light red
    set_cell_text(cells[0], row_content[0], bold=True)
    set_cell_background(cells[1], 'FFFFFF')
    set_cell_text(cells[1], row_content[1])
    set_cell_background(cells[2], 'FFFFFF')
    set_cell_text(cells[2], row_content[2])

doc.add_paragraph()

# Clinical Pearls
add_colored_box(doc, 'Clinical Pearls & High-Yield Points:', [
    'Remember: "Sympathetic = fight or flight = need more O₂ = dilate airways"',
    'β₂-agonists (albuterol, salbutamol) are rescue inhalers - they work within minutes by relaxing bronchial smooth muscle',
    'Asthma is a chronic inflammatory disease with acute bronchospasm - treat inflammation with steroids, bronchospasm with β₂-agonists',
    'Anticholinergics (block M₃ receptors) also treat asthma by preventing parasympathetic bronchoconstriction'
], (0, 77, 64), 'E0F2F1')

doc.add_paragraph()

# Memory Tricks
add_colored_box(doc, 'Memory Tricks & Mnemonics:', [
    '"Beta-2 makes airways Breathe Better" - β₂-agonists cause bronchodilation',
    '"M3 = Must constrict" - M₃ muscarinic activation causes bronchoconstriction',
    '"Sympathetic = Dilate, Parasympathetic = Constrict" (opposite of pupils!)',
    '"Albuterol opens Airways" - classic asthma rescue inhaler'
], (230, 81, 0), 'FFF3E0')

doc.add_paragraph()

# Analogy
add_colored_box(doc, 'Analogy:', [
    'Think of your airways like adjustable garden hoses. The sympathetic nervous system (fight-or-flight) opens the nozzle wide (bronchodilation) so you can breathe in lots of air quickly during exercise or stress. The parasympathetic system (rest-and-digest) narrows the nozzle slightly (bronchoconstriction) during rest. In asthma, the nozzle gets stuck in the narrowed position and albuterol acts like opening the nozzle back up.'
], (118, 75, 162), 'F3E5F5')

doc.add_page_break()

# ========== LO5: Mucociliary Clearance ==========
obj5 = doc.add_heading('Learning Objective 5:', 2)
obj5.runs[0].font.color.rgb = RGBColor(118, 75, 162)
doc.add_paragraph('Describe the role of mucociliary clearance in removing foreign particles')

# Summary
summary_para = doc.add_paragraph()
summary_run = summary_para.add_run('Summary: ')
summary_run.bold = True
summary_para.add_run('Mucociliary clearance (the "mucociliary escalator") removes inhaled particles and pathogens from airways. Effective clearance requires both ciliary activity and respiratory tract fluids. Ciliated cells produce periciliary fluid (sol layer) with low viscosity optimal for ciliary beating. Goblet cells produce mucus (gel layer, 5-10 μm thick) with high viscosity that traps inhaled materials. Cilia beat at 10-20 times per second, propelling mucus upward to pharynx where it\'s swallowed or coughed up. Cigarette smoking impairs mucociliary clearance by increasing airway wall thickness, increasing goblet cells, and decreasing ciliary cell activity and number.')

doc.add_paragraph()

# TABLE 1: Components of Mucociliary Clearance
doc.add_heading('TABLE 1: Components of Mucociliary Clearance System', 3)
table5 = doc.add_table(rows=4, cols=3)
table5.style = 'Table Grid'

# Headers
headers = ['Component', 'Source', 'Function']
for i, header_text in enumerate(headers):
    set_cell_background(table5.rows[0].cells[i], 'B2DFDB')  # Teal
    set_cell_text(table5.rows[0].cells[i], header_text, bold=True, color=(255, 255, 255), size=12)

# Data rows
row_data = [
    ['Cilia', 'Ciliated epithelial cells (~200 cilia/cell)', 'Beat 10-20 times/second to propel mucus upward'],
    ['Periciliary Fluid (Sol Layer)', 'Ciliated cells', 'Low viscosity fluid optimal for ciliary beating'],
    ['Mucus (Gel Layer)', 'Goblet cells', '5-10 μm thick, high viscosity, traps inhaled particles and pathogens']
]

for row_idx, row_content in enumerate(row_data, start=1):
    cells = table5.rows[row_idx].cells
    set_cell_background(cells[0], 'E0F2F1')  # Light teal
    set_cell_text(cells[0], row_content[0], bold=True)
    set_cell_background(cells[1], 'FFFFFF')
    set_cell_text(cells[1], row_content[1])
    set_cell_background(cells[2], 'FFFFFF')
    set_cell_text(cells[2], row_content[2])

doc.add_paragraph()

# TABLE 2: Effects of Smoking
doc.add_heading('TABLE 2: Cigarette Smoking Effects on Mucociliary Clearance', 3)
table5b = doc.add_table(rows=4, cols=2)
table5b.style = 'Table Grid'

# Headers
headers = ['Smoking Effect', 'Result']
for i, header_text in enumerate(headers):
    set_cell_background(table5b.rows[0].cells[i], 'FFCDD2')  # Red
    set_cell_text(table5b.rows[0].cells[i], header_text, bold=True, color=(255, 255, 255), size=12)

# Data rows
row_data = [
    ['Airway wall thickness ↑↑', 'Narrowed airways'],
    ['Number of goblet cells ↑↑', 'Excessive mucus production'],
    ['Ciliary cell activity ↓↓ and number ↓↓', 'Impaired mucus clearance']
]

for row_idx, row_content in enumerate(row_data, start=1):
    cells = table5b.rows[row_idx].cells
    set_cell_background(cells[0], 'FFEBEE')  # Light red
    set_cell_text(cells[0], row_content[0], bold=True)
    set_cell_background(cells[1], 'FFFFFF')
    set_cell_text(cells[1], row_content[1])

doc.add_paragraph()

# Clinical Pearls
add_colored_box(doc, 'Clinical Pearls & High-Yield Points:', [
    'Mucociliary clearance is your airway\'s "self-cleaning" mechanism - first line of defense against inhaled pathogens',
    'Cilia beat continuously upward toward pharynx - particles trapped in mucus get swallowed (destroyed by stomach acid) or coughed out',
    'Chronic smoking → "smoker\'s cough" due to impaired clearance and excess mucus that can\'t be cleared effectively',
    'Cystic fibrosis patients have thick, dehydrated mucus that can\'t be cleared properly → chronic infections',
    'Each ciliated cell has ~200 cilia beating in coordinated waves like a field of wheat'
], (0, 77, 64), 'E0F2F1')

doc.add_paragraph()

# Memory Tricks
add_colored_box(doc, 'Memory Tricks & Mnemonics:', [
    '"Mucociliary Escalator" - Just like a mall escalator, it only goes UP (to pharynx)',
    '"Sol = Solution (thin/watery), Gel = Gelatin (thick/sticky)" - helps remember viscosity',
    '"Goblet cells make Glue (mucus) to Grab particles"',
    '"Smoking Stops the Sweepers" - impairs ciliary clearance'
], (230, 81, 0), 'FFF3E0')

doc.add_paragraph()

# Analogy
add_colored_box(doc, 'Analogy:', [
    'Think of mucociliary clearance like a conveyor belt at a grocery store checkout. The cilia are the moving belt (constantly running upward), the periciliary fluid is the lubricant keeping the belt moving smoothly, and the mucus is sticky flypaper on the belt that catches anything that lands on it. Dust, bacteria, and debris get stuck to the flypaper and are carried up and out. When you smoke, it\'s like throwing sand in the motor - the belt slows down, gets gunked up, and can\'t clear debris effectively anymore.'
], (118, 75, 162), 'F3E5F5')

doc.add_page_break()

# ========== LO6: Fick's Law ==========
obj6 = doc.add_heading('Learning Objective 6:', 2)
obj6.runs[0].font.color.rgb = RGBColor(118, 75, 162)
doc.add_paragraph('Determine the factors that will influence diffusion of a gas (Fick\'s Law)')

# Summary
summary_para = doc.add_paragraph()
summary_run = summary_para.add_run('Summary: ')
summary_run.bold = True
summary_para.add_run('Fick\'s Law describes factors affecting gas diffusion rate: V̇ₓ = (D × A × ΔP) / ΔX. Diffusion is directly proportional to: (1) Diffusion coefficient (D) of the gas, (2) Surface area (A) available for exchange, and (3) Partial pressure gradient (ΔP). Diffusion is inversely proportional to membrane thickness (ΔX). The alveolar-capillary barrier is optimized for diffusion with thin walls (0.5 μm) and huge surface area (50-100 m²). Clinical factors affecting diffusion: increased fibrotic tissue → increased thickness → decreased diffusion; lung resection → decreased surface area → decreased diffusion; high altitude → decreased pressure gradient → decreased diffusion; exercise → increased surface area → increased diffusion.')

doc.add_paragraph()

# TABLE 1: Fick's Law Components
doc.add_heading('TABLE 1: Fick\'s Law - Factors Affecting Gas Diffusion', 3)
table6 = doc.add_table(rows=5, cols=4)
table6.style = 'Table Grid'

# Headers
headers = ['Factor', 'Symbol', 'Relationship', 'Effect on Diffusion']
for i, header_text in enumerate(headers):
    set_cell_background(table6.rows[0].cells[i], 'D1C4E9')  # Purple
    set_cell_text(table6.rows[0].cells[i], header_text, bold=True, color=(255, 255, 255), size=12)

# Data rows
row_data = [
    ['Diffusion Coefficient', 'D', 'Directly proportional', 'Higher D → More diffusion'],
    ['Surface Area', 'A', 'Directly proportional', 'Larger A → More diffusion'],
    ['Pressure Gradient', 'ΔP', 'Directly proportional', 'Greater ΔP → More diffusion'],
    ['Membrane Thickness', 'ΔX', 'Inversely proportional', 'Greater ΔX → Less diffusion']
]

for row_idx, row_content in enumerate(row_data, start=1):
    cells = table6.rows[row_idx].cells
    set_cell_background(cells[0], 'F3E5F5')  # Light purple
    set_cell_text(cells[0], row_content[0], bold=True)
    set_cell_background(cells[1], 'FFFFFF')
    set_cell_text(cells[1], row_content[1], bold=True)
    set_cell_background(cells[2], 'FFFFFF')
    set_cell_text(cells[2], row_content[2])
    set_cell_background(cells[3], 'FFFFFF')
    set_cell_text(cells[3], row_content[3])

doc.add_paragraph()

# TABLE 2: Clinical Applications
doc.add_heading('TABLE 2: Clinical Conditions Affecting Diffusion', 3)
table6b = doc.add_table(rows=5, cols=3)
table6b.style = 'Table Grid'

# Headers
headers = ['Condition', 'Fick\'s Law Factor Affected', 'Effect on Diffusion']
for i, header_text in enumerate(headers):
    set_cell_background(table6b.rows[0].cells[i], 'FFCDD2')  # Red
    set_cell_text(table6b.rows[0].cells[i], header_text, bold=True, color=(255, 255, 255), size=12)

# Data rows
row_data = [
    ['Restrictive lung disease (fibrosis)', 'Thickness ↑↑', 'Diffusion ↓↓'],
    ['Lung resection (lobectomy)', 'Surface area ↓↓', 'Diffusion ↓↓'],
    ['High altitude (Cuzco, Peru)', 'Pressure gradient ↓↓', 'Diffusion ↓↓'],
    ['Exercise', 'Surface area ↑↑', 'Diffusion ↑↑']
]

for row_idx, row_content in enumerate(row_data, start=1):
    cells = table6b.rows[row_idx].cells
    set_cell_background(cells[0], 'FFEBEE')  # Light red
    set_cell_text(cells[0], row_content[0], bold=True)
    set_cell_background(cells[1], 'FFFFFF')
    set_cell_text(cells[1], row_content[1])
    set_cell_background(cells[2], 'FFFFFF')
    set_cell_text(cells[2], row_content[2])

doc.add_paragraph()

# Clinical Pearls
add_colored_box(doc, 'Clinical Pearls & High-Yield Points:', [
    'Lungs are perfectly designed for diffusion: thin barrier (0.5 μm) + huge surface area (tennis court) = rapid, efficient gas exchange',
    'Pulmonary fibrosis/interstitial lung disease → thickened alveolar-capillary membrane → impaired diffusion → hypoxemia',
    'At high altitude, PO₂ drops (even though % O₂ still 21%) → decreased driving pressure → less diffusion → altitude sickness',
    'During exercise, recruitment of additional capillaries increases functional surface area → enhanced diffusion capacity',
    'RBCs have only 0.75 seconds in pulmonary capillaries - thin barrier and large surface area allow complete equilibration in this short time'
], (0, 77, 64), 'E0F2F1')

doc.add_paragraph()

# Memory Tricks
add_colored_box(doc, 'Memory Tricks & Mnemonics:', [
    '"Big Surface, Small Distance, Big Pressure difference = Fast Diffusion" (A ↑, ΔX ↓, ΔP ↑ = V̇ ↑)',
    '"Thick walls are Bad" - Inversely proportional (fibrosis decreases diffusion)',
    '"High altitude = Low pressure = Slow diffusion"',
    'Think "ASAP" for factors: Area, Surface properties (D), And Pressure gradient (direct); thickness is the exception (inverse)'
], (230, 81, 0), 'FFF3E0')

doc.add_paragraph()

# Analogy
add_colored_box(doc, 'Analogy:', [
    'Fick\'s Law is like filtering coffee: (1) Larger filter (surface area) = faster brewing, (2) Thinner filter paper (membrane thickness) = faster flow, (3) More pressure difference (gravity pulling water down) = faster extraction. Your alveoli use the same principles: giant coffee filter (100 m²), ultra-thin paper (0.5 μm), and strong pressure gradient to rapidly extract O₂ from air into blood.'
], (118, 75, 162), 'F3E5F5')

doc.add_page_break()

# ========== LO7: Partial Pressure ==========
obj7 = doc.add_heading('Learning Objective 7:', 2)
obj7.runs[0].font.color.rgb = RGBColor(118, 75, 162)
doc.add_paragraph('Define partial pressure of a gas')

# Summary
summary_para = doc.add_paragraph()
summary_run = summary_para.add_run('Summary: ')
summary_run.bold = True
summary_para.add_run('Partial pressure (Pgas) is the pressure a single gas exerts in a mixture of gases. Atmospheric pressure (barometric pressure) is the sum of all partial pressures of gases in the atmosphere. At sea level, atmospheric pressure ≈ 760 mmHg. Air is 21% O₂ and 78% nitrogen. To calculate partial pressure: Pgas = fractional concentration × Patm. At sea level: PO₂ = 0.21 × 760 mmHg ≈ 160 mmHg. At high altitude (Mt. Everest, ~250 mmHg atmospheric pressure), O₂ percentage stays 21%, but PO₂ = 0.21 × 250 ≈ 53 mmHg (much lower).')

doc.add_paragraph()

# TABLE 1: Partial Pressure Calculations
doc.add_heading('TABLE 1: Partial Pressure at Different Altitudes', 3)
table7 = doc.add_table(rows=3, cols=4)
table7.style = 'Table Grid'

# Headers
headers = ['Location', 'Atmospheric Pressure', 'O₂ Percentage', 'PO₂']
for i, header_text in enumerate(headers):
    set_cell_background(table7.rows[0].cells[i], 'B3E5FC')  # Blue
    set_cell_text(table7.rows[0].cells[i], header_text, bold=True, color=(255, 255, 255), size=12)

# Data rows
row_data = [
    ['Sea Level', '~760 mmHg', '21%', '~160 mmHg (0.21 × 760)'],
    ['Mt. Everest Summit', '~250 mmHg', '21%', '~53 mmHg (0.21 × 250)']
]

for row_idx, row_content in enumerate(row_data, start=1):
    cells = table7.rows[row_idx].cells
    set_cell_background(cells[0], 'E1F5FE')  # Light blue
    set_cell_text(cells[0], row_content[0], bold=True)
    set_cell_background(cells[1], 'FFFFFF')
    set_cell_text(cells[1], row_content[1])
    set_cell_background(cells[2], 'FFFFFF')
    set_cell_text(cells[2], row_content[2])
    set_cell_background(cells[3], 'FFFFFF')
    set_cell_text(cells[3], row_content[3])

doc.add_paragraph()

# TABLE 2: Atmospheric Gas Composition
doc.add_heading('TABLE 2: Atmospheric Gas Composition at Sea Level', 3)
table7b = doc.add_table(rows=4, cols=3)
table7b.style = 'Table Grid'

# Headers
headers = ['Gas', 'Percentage', 'Partial Pressure (at 760 mmHg)']
for i, header_text in enumerate(headers):
    set_cell_background(table7b.rows[0].cells[i], 'C8E6C9')  # Green
    set_cell_text(table7b.rows[0].cells[i], header_text, bold=True, color=(255, 255, 255), size=12)

# Data rows
row_data = [
    ['Nitrogen (N₂)', '78%', '~593 mmHg'],
    ['Oxygen (O₂)', '21%', '~160 mmHg'],
    ['Other (Argon, CO₂, etc.)', '~1%', '~7 mmHg']
]

for row_idx, row_content in enumerate(row_data, start=1):
    cells = table7b.rows[row_idx].cells
    set_cell_background(cells[0], 'E8F5E9')  # Light green
    set_cell_text(cells[0], row_content[0], bold=True)
    set_cell_background(cells[1], 'FFFFFF')
    set_cell_text(cells[1], row_content[1])
    set_cell_background(cells[2], 'FFFFFF')
    set_cell_text(cells[2], row_content[2])

doc.add_paragraph()

# Clinical Pearls
add_colored_box(doc, 'Clinical Pearls & High-Yield Points:', [
    'KEY CONCEPT: At high altitude, O₂ percentage stays the same (21%), but PO₂ drops dramatically due to lower atmospheric pressure',
    'This is why climbers on Everest need supplemental O₂ - not because there\'s less O₂ percentage, but because lower atmospheric pressure means lower PO₂',
    'Partial pressure drives diffusion - O₂ moves from high PO₂ (alveoli, ~100 mmHg) to low PO₂ (venous blood, ~40 mmHg)',
    'Nitrogen is the most abundant gas in atmosphere (78%) but doesn\'t participate in respiration - it\'s inert',
    'At sea level: atmospheric pressure = 760 mmHg, PO₂ = 160 mmHg, alveolar PO₂ = ~100 mmHg (due to mixing with residual air and water vapor)'
], (0, 77, 64), 'E0F2F1')

doc.add_paragraph()

# Memory Tricks
add_colored_box(doc, 'Memory Tricks & Mnemonics:', [
    '"21% at sea level = ~160 mmHg PO₂" (0.21 × 760 ≈ 160)',
    '"Nitrogen = 78%, Oxygen = 21%, Other = 1%" (78-21-1 rule)',
    '"Same percentage, Different pressure" - O₂ % constant at altitude, but PO₂ drops',
    'Formula: "Pgas = Fraction × Patm" (partial pressure = fraction of gas × atmospheric pressure)'
], (230, 81, 0), 'FFF3E0')

doc.add_paragraph()

# Analogy
add_colored_box(doc, 'Analogy:', [
    'Think of atmospheric pressure like water depth in a pool. At sea level, you\'re at the bottom of an "ocean of air" - lots of pressure (760 mmHg). On Mt. Everest, you\'re near the surface - less air above you, less pressure (250 mmHg). The percentage of O₂ in air is like the percentage of chlorine in the pool - it stays constant throughout, but the total amount you\'re exposed to depends on the depth/pressure. At Everest\'s "shallow" depth, even though there\'s still 21% O₂, there\'s less total air pressure pushing it into your lungs.'
], (118, 75, 162), 'F3E5F5')

doc.add_page_break()

# ========== LO8: Tidal Volume and Dead Space ==========
obj8 = doc.add_heading('Learning Objective 8:', 2)
obj8.runs[0].font.color.rgb = RGBColor(118, 75, 162)
doc.add_paragraph('Describe tidal volume and dead space')

# Summary
summary_para = doc.add_paragraph()
summary_run = summary_para.add_run('Summary: ')
summary_run.bold = True
summary_para.add_run('Tidal volume is the amount of air that moves in and out per breath during normal breathing (~500 mL). Humans use tidal (bidirectional) ventilation through the same path. To adjust for increased metabolic demand during exercise, we can adjust breathing rate or tidal volume. Dead space is the volume of air in conducting airways (~150 mL) that doesn\'t participate in gas exchange. Dead space is approximately 1/3 of tidal volume (150 mL out of 500 mL). The airway dead space must be inhaled before any fresh inspired air can reach the lungs for gas exchange, so tidal volume must be large enough to account for dead space to supply sufficient oxygen.')

doc.add_paragraph()

# TABLE 1: Tidal Volume and Dead Space
doc.add_heading('TABLE 1: Tidal Volume and Dead Space Values', 3)
table8 = doc.add_table(rows=3, cols=3)
table8.style = 'Table Grid'

# Headers
headers = ['Parameter', 'Normal Value', 'Description']
for i, header_text in enumerate(headers):
    set_cell_background(table8.rows[0].cells[i], 'FFE0B2')  # Orange
    set_cell_text(table8.rows[0].cells[i], header_text, bold=True, color=(255, 255, 255), size=12)

# Data rows
row_data = [
    ['Tidal Volume (TV)', '~500 mL', 'Air moved in/out per breath during normal breathing'],
    ['Anatomical Dead Space', '~150 mL', 'Air in conducting airways that doesn\'t participate in gas exchange (~1/3 of TV)']
]

for row_idx, row_content in enumerate(row_data, start=1):
    cells = table8.rows[row_idx].cells
    set_cell_background(cells[0], 'FFF3E0')  # Light orange
    set_cell_text(cells[0], row_content[0], bold=True)
    set_cell_background(cells[1], 'FFFFFF')
    set_cell_text(cells[1], row_content[1], bold=True)
    set_cell_background(cells[2], 'FFFFFF')
    set_cell_text(cells[2], row_content[2])

doc.add_paragraph()

# TABLE 2: Lung Volumes
doc.add_heading('TABLE 2: Lung Volumes Beyond Tidal Volume', 3)
table8b = doc.add_table(rows=4, cols=2)
table8b.style = 'Table Grid'

# Headers
headers = ['Volume', 'Description']
for i, header_text in enumerate(headers):
    set_cell_background(table8b.rows[0].cells[i], 'B3E5FC')  # Blue
    set_cell_text(table8b.rows[0].cells[i], header_text, bold=True, color=(255, 255, 255), size=12)

# Data rows
row_data = [
    ['Inspiratory Reserve Volume', 'Extra air that can be inhaled beyond normal tidal breath'],
    ['Expiratory Reserve Volume', 'Extra air that can be exhaled beyond normal tidal breath'],
    ['Residual Volume', 'Air remaining in lungs after maximal exhalation (cannot be exhaled)']
]

for row_idx, row_content in enumerate(row_data, start=1):
    cells = table8b.rows[row_idx].cells
    set_cell_background(cells[0], 'E1F5FE')  # Light blue
    set_cell_text(cells[0], row_content[0], bold=True)
    set_cell_background(cells[1], 'FFFFFF')
    set_cell_text(cells[1], row_content[1])

doc.add_paragraph()

# Clinical Pearls
add_colored_box(doc, 'Clinical Pearls & High-Yield Points:', [
    'The "Rule of 1/3": Dead space = 1/3 of tidal volume (150 mL / 500 mL = 30%)',
    'Only 350 mL of the 500 mL tidal breath actually reaches alveoli for gas exchange (500 - 150 = 350 mL)',
    'Shallow breathing is inefficient: if TV drops to 150 mL, ALL air goes to dead space → zero alveolar ventilation!',
    'During exercise: increase both rate AND depth (tidal volume) to maximize alveolar ventilation',
    'Dead space is why patients with tracheostomy (bypasses upper airway) need less tidal volume - reduced dead space'
], (0, 77, 64), 'E0F2F1')

doc.add_paragraph()

# Memory Tricks
add_colored_box(doc, 'Memory Tricks & Mnemonics:', [
    '"500 and 150" or "5 and 1.5" - Normal tidal volume and dead space',
    '"Dead space = 1/3 tidal volume" (easiest way to remember the relationship)',
    '"Tidal = Tide coming in and out" - bidirectional breathing',
    '"150 mL is wasted space, 350 mL does the work" (500 - 150 = 350 for gas exchange)'
], (230, 81, 0), 'FFF3E0')

doc.add_paragraph()

# Analogy
add_colored_box(doc, 'Analogy:', [
    'Think of breathing like filling a water bottle with a long straw. The straw is your dead space (conducting airways, ~150 mL) and the bottle is your alveoli. Every time you pour water (inhale), you must first fill the straw before any water reaches the bottle. If you only pour 150 mL, it all stays in the straw - nothing reaches the bottle! You need to pour at least 500 mL (tidal volume) so that 350 mL actually reaches the bottle after filling the straw. The straw contents get pushed out on the next pour (exhale), but it\'s "wasted" volume that never reached the destination.'
], (118, 75, 162), 'F3E5F5')

doc.add_page_break()

# ==========================
# TAB 2: KEY COMPARISONS
# ==========================

heading2 = doc.add_heading('Key Comparisons', 1)
heading2.runs[0].font.color.rgb = RGBColor(118, 75, 162)

# Comparison 1: Conducting Zone vs Respiratory Zone (repeated for easy reference)
doc.add_heading('Conducting Zone vs Respiratory Zone', 2)
table_comp1 = doc.add_table(rows=7, cols=3)
table_comp1.style = 'Table Grid'

headers = ['Feature', 'Conducting Zone', 'Respiratory Zone']
for i, header_text in enumerate(headers):
    set_cell_background(table_comp1.rows[0].cells[i], 'B3E5FC')
    set_cell_text(table_comp1.rows[0].cells[i], header_text, bold=True, color=(255, 255, 255), size=12)

row_data = [
    ['Structures', 'Trachea, bronchi, bronchioles, terminal bronchioles (gen 0-16)', 'Respiratory bronchioles, alveolar ducts, alveolar sacs (gen 17-23)'],
    ['Primary Function', 'Air transport and conditioning', 'Gas exchange via diffusion'],
    ['Gas Exchange', 'NO gas exchange', 'YES - O₂ and CO₂ diffusion'],
    ['Volume', '~150 mL (anatomical dead space)', 'Alveolar ventilation volume'],
    ['Key Functions', 'Humidification, warming, filtration, ↓ airflow velocity, regulate airway resistance', 'Maximize surface area for diffusion'],
    ['Special Features', 'Smooth muscle, cilia, mucous-secreting goblet cells', '~300 million alveoli, 50-100 m² total surface area, Type I & II alveolar cells']
]

for row_idx, row_content in enumerate(row_data, start=1):
    cells = table_comp1.rows[row_idx].cells
    set_cell_background(cells[0], 'E1F5FE')
    set_cell_text(cells[0], row_content[0], bold=True)
    set_cell_background(cells[1], 'FFFFFF')
    set_cell_text(cells[1], row_content[1])
    set_cell_background(cells[2], 'FFFFFF')
    set_cell_text(cells[2], row_content[2])

doc.add_paragraph()

# Comparison 2: Type I vs Type II Alveolar Cells
doc.add_heading('Type I vs Type II Alveolar Cells', 2)
table_comp2 = doc.add_table(rows=4, cols=3)
table_comp2.style = 'Table Grid'

headers = ['Feature', 'Type I Alveolar Cell', 'Type II Alveolar Cell']
for i, header_text in enumerate(headers):
    set_cell_background(table_comp2.rows[0].cells[i], 'C8E6C9')
    set_cell_text(table_comp2.rows[0].cells[i], header_text, bold=True, color=(255, 255, 255), size=12)

row_data = [
    ['Structure', 'Thin, flat epithelial cells', 'Cuboidal cells'],
    ['Function', 'Form alveolar epithelium for gas exchange', 'Synthesize pulmonary surfactant'],
    ['Coverage', 'Cover ~95% of alveolar surface', 'Less numerous but metabolically active']
]

for row_idx, row_content in enumerate(row_data, start=1):
    cells = table_comp2.rows[row_idx].cells
    set_cell_background(cells[0], 'E8F5E9')
    set_cell_text(cells[0], row_content[0], bold=True)
    set_cell_background(cells[1], 'FFFFFF')
    set_cell_text(cells[1], row_content[1])
    set_cell_background(cells[2], 'FFFFFF')
    set_cell_text(cells[2], row_content[2])

doc.add_paragraph()

# Comparison 3: Sympathetic vs Parasympathetic Control of Airways
doc.add_heading('Sympathetic vs Parasympathetic Control of Airways', 2)
table_comp3 = doc.add_table(rows=5, cols=3)
table_comp3.style = 'Table Grid'

headers = ['Feature', 'Sympathetic (SNS)', 'Parasympathetic (PNS)']
for i, header_text in enumerate(headers):
    set_cell_background(table_comp3.rows[0].cells[i], 'FFE0B2')
    set_cell_text(table_comp3.rows[0].cells[i], header_text, bold=True, color=(255, 255, 255), size=12)

row_data = [
    ['Receptor', 'β₂-adrenergic', 'M₃ muscarinic'],
    ['Effect on Airways', 'Bronchodilation', 'Bronchoconstriction'],
    ['Effect on Resistance', 'Resistance ↓↓', 'Resistance ↑↑'],
    ['Clinical Example', 'Albuterol (β₂-agonist) for asthma rescue', 'Anticholinergics block M₃ to prevent constriction']
]

for row_idx, row_content in enumerate(row_data, start=1):
    cells = table_comp3.rows[row_idx].cells
    set_cell_background(cells[0], 'FFF3E0')
    set_cell_text(cells[0], row_content[0], bold=True)
    set_cell_background(cells[1], 'FFFFFF')
    set_cell_text(cells[1], row_content[1])
    set_cell_background(cells[2], 'FFFFFF')
    set_cell_text(cells[2], row_content[2])

doc.add_paragraph()

# Comparison 4: External vs Internal Respiration
doc.add_heading('External vs Internal Respiration', 2)
table_comp4 = doc.add_table(rows=4, cols=3)
table_comp4.style = 'Table Grid'

headers = ['Feature', 'External Respiration', 'Internal Respiration']
for i, header_text in enumerate(headers):
    set_cell_background(table_comp4.rows[0].cells[i], 'D1C4E9')
    set_cell_text(table_comp4.rows[0].cells[i], header_text, bold=True, color=(255, 255, 255), size=12)

row_data = [
    ['Definition', 'Getting O₂ from atmosphere to mitochondria (and CO₂ out)', 'Cellular use of O₂ in mitochondria'],
    ['Location', 'Lungs, blood, tissues', 'Mitochondria within cells'],
    ['Process', '4 steps: Ventilation → Gas exchange → Transport → Gas exchange', 'Glucose + O₂ → CO₂ + H₂O + ATP']
]

for row_idx, row_content in enumerate(row_data, start=1):
    cells = table_comp4.rows[row_idx].cells
    set_cell_background(cells[0], 'F3E5F5')
    set_cell_text(cells[0], row_content[0], bold=True)
    set_cell_background(cells[1], 'FFFFFF')
    set_cell_text(cells[1], row_content[1])
    set_cell_background(cells[2], 'FFFFFF')
    set_cell_text(cells[2], row_content[2])

doc.add_page_break()

# ==========================
# TAB 3: MASTER CHART
# ==========================

heading3 = doc.add_heading('Master Chart - Introduction to Respiration', 1)
heading3.runs[0].font.color.rgb = RGBColor(118, 75, 162)

doc.add_heading('Comprehensive Respiratory Physiology Concepts', 2)

# Create master chart table
table_master = doc.add_table(rows=19, cols=2)
table_master.style = 'Table Grid'

# Headers
headers = ['Concept', 'Key Details']
for i, header_text in enumerate(headers):
    set_cell_background(table_master.rows[0].cells[i], 'EDE7F6')  # Light purple
    set_cell_text(table_master.rows[0].cells[i], header_text, bold=True, color=(118, 75, 162), size=12)

# Master chart data
master_data = [
    ['External Respiration Steps', '4 steps: (1) Ventilation, (2) Gas exchange (alveoli-blood), (3) Transport, (4) Gas exchange (blood-tissue)'],
    ['Internal Respiration', 'Glucose + O₂ → CO₂ + H₂O + ATP in mitochondria'],
    ['Respiratory Quotient (RQ)', 'VCO₂/VO₂ ratio. Carbs = 1.0, Fats = 0.7, Proteins/mixed = 0.8'],
    ['Conducting Zone', 'Gen 0-16: Trachea to terminal bronchioles. NO gas exchange. Functions: warm, humidify, filter air'],
    ['Respiratory Zone', 'Gen 17-23: Respiratory bronchioles to alveolar sacs. Gas exchange occurs here. 50-100 m² surface area'],
    ['Anatomical Dead Space', '~150 mL conducting zone air that doesn\'t participate in gas exchange (~1/3 tidal volume)'],
    ['Type I Alveolar Cells', 'Thin, flat epithelial cells forming alveolar epithelium for gas exchange'],
    ['Type II Alveolar Cells', 'Cuboidal cells that synthesize pulmonary surfactant'],
    ['Sympathetic Control (β₂)', 'β₂-adrenergic activation → bronchodilation → ↓ resistance. Example: Albuterol'],
    ['Parasympathetic Control (M₃)', 'M₃ muscarinic activation → bronchoconstriction → ↑ resistance'],
    ['Asthma', 'Bronchial smooth muscle constriction + chronic inflammation. Treat with β₂-agonists (albuterol)'],
    ['Mucociliary Clearance', 'Cilia (beat 10-20×/sec) + mucus trap and remove particles. Impaired by smoking'],
    ['Periciliary Fluid (Sol Layer)', 'Low viscosity fluid from ciliated cells, optimal for ciliary beating'],
    ['Mucus (Gel Layer)', '5-10 μm thick from goblet cells. High viscosity, traps particles'],
    ['Fick\'s Law', 'V̇ = (D × A × ΔP) / ΔX. Diffusion ∝ area & pressure gradient; ∝ 1/thickness'],
    ['Partial Pressure', 'Pgas = fraction × Patm. At sea level (760 mmHg): PO₂ = 0.21 × 760 = 160 mmHg'],
    ['Tidal Volume', '~500 mL air moved in/out per normal breath'],
    ['Alveolar-Capillary Barrier', '0.5 μm thick (3 layers: alveolar epithelium, interstitial fluid, capillary endothelium). RBCs spend ~0.75 sec']
]

for row_idx, row_content in enumerate(master_data, start=1):
    cells = table_master.rows[row_idx].cells
    set_cell_background(cells[0], 'F3E5F5')  # Light purple
    set_cell_text(cells[0], row_content[0], bold=True, size=11)
    set_cell_background(cells[1], 'FFFFFF')
    set_cell_text(cells[1], row_content[1], size=11)

doc.add_page_break()

# ==========================
# TAB 4: HIGH-YIELD SUMMARY
# ==========================

heading4 = doc.add_heading('High-Yield Summary', 1)
heading4.runs[0].font.color.rgb = RGBColor(118, 75, 162)

# Anatomy Must-Knows
add_colored_box(doc, 'ANATOMY - Must Know:', [
    'Conducting zone (gen 0-16) = NO gas exchange, anatomical dead space (~150 mL)',
    'Respiratory zone (gen 17-23) = gas exchange, 300 million alveoli, 50-100 m² surface area',
    'Alveolar-capillary barrier = 0.5 μm thick (thinnest in body for rapid diffusion)',
    'Type I alveolar cells = gas exchange surface; Type II = surfactant production',
    'Three layers of barrier: alveolar epithelium, interstitial fluid, capillary endothelium'
], (118, 75, 162), 'F3E5F5')

doc.add_paragraph()

# Critical Formulas & Values
add_colored_box(doc, 'FORMULAS & KEY VALUES - Memorize:', [
    'RQ = VCO₂ / VO₂ (Carbs = 1.0, Fats = 0.7, Proteins/mixed = 0.8)',
    'Fick\'s Law: V̇ = (D × A × ΔP) / ΔX',
    'Partial Pressure: Pgas = fraction × Patm',
    'Sea level: Patm = 760 mmHg, PO₂ = 160 mmHg (21% × 760)',
    'Tidal volume = 500 mL, Dead space = 150 mL (~1/3 TV)',
    'Normal metabolism: VO₂ = 250 mL/min, VCO₂ = 200 mL/min → RQ = 0.8'
], (0, 77, 64), 'E0F2F1')

doc.add_paragraph()

# Clinical Pearls Summary
add_colored_box(doc, 'Clinical Pearls & High-Yield Points - SUMMARY:', [
    'External respiration (4 steps getting O₂ in, CO₂ out) ≠ Internal respiration (mitochondrial ATP production)',
    'Dead space = wasted ventilation. Shallow breathing (<150 mL) = zero alveolar ventilation!',
    'Sympathetic = bronchodilation (β₂), Parasympathetic = bronchoconstriction (M₃)',
    'Asthma = bronchospasm. First-line rescue: Albuterol (β₂-agonist)',
    'Smoking → ↑ mucus, ↓ cilia → impaired clearance → chronic infections',
    'High altitude: O₂ % stays 21%, but ↓ Patm → ↓ PO₂ → hypoxemia',
    'Fick\'s Law: ↑ area & ↑ pressure gradient = ↑ diffusion; ↑ thickness = ↓ diffusion',
    'Pulmonary fibrosis → ↑ barrier thickness → impaired diffusion → hypoxemia'
], (0, 77, 64), 'E0F2F1')

doc.add_paragraph()

# Pathophysiology
add_colored_box(doc, 'PATHOPHYSIOLOGY - Disease Connections:', [
    'Asthma: Chronic inflammation + bronchospasm → ↑ airway resistance → dyspnea',
    'Restrictive lung disease (fibrosis): ↑ alveolar-capillary thickness → ↓ diffusion',
    'Emphysema: Loss of alveolar surface area → ↓ diffusion capacity',
    'Neonatal RDS: Insufficient surfactant (Type II cell deficiency) → alveolar collapse',
    'Cystic fibrosis: Thick mucus → impaired mucociliary clearance → chronic infections',
    'High altitude sickness: ↓ PO₂ → ↓ O₂ delivery → hypoxemia'
], (183, 28, 28), 'FFEBEE')

doc.add_paragraph()

# Quick Facts
add_colored_box(doc, 'QUICK FACTS - Respiratory Physiology:', [
    'Air composition: 78% N₂, 21% O₂, 1% other',
    'RBCs in pulmonary capillaries for only 0.75 seconds (traverse 2-3 alveoli)',
    'Each ciliated cell has ~200 cilia beating 10-20 times/second',
    'Mucociliary escalator moves mucus upward to pharynx (swallowed or coughed)',
    'Exercise increases diffusion by recruiting more capillaries → ↑ surface area',
    'Mt. Everest summit: Patm ~250 mmHg, PO₂ ~53 mmHg (vs. 160 mmHg at sea level)'
], (230, 81, 0), 'FFF3E0')

doc.add_paragraph()

# Treatment Pearls
add_colored_box(doc, 'TREATMENT - Clinical Applications:', [
    '🟢 Asthma rescue: Albuterol (β₂-agonist) → bronchodilation',
    'Asthma maintenance: Inhaled corticosteroids (treat inflammation)',
    'Anticholinergics (block M₃) also used for bronchodilation',
    'Supplemental O₂ at high altitude or in lung disease to ↑ PO₂ gradient',
    'Mucolytics and chest physiotherapy help clear thick mucus (CF, COPD)'
], (27, 94, 32), 'E8F5E9')

# Save document
output_path = '/home/user/Skills/Example claude study guides/Lecture_33_Introduction_to_Respiration_Study_Guide.docx'
doc.save(output_path)
print(f'✅ Study guide created: {output_path}')
