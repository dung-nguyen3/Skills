#!/usr/bin/env python3
"""
Create a Word template (.dotx) with custom table styles and quick tables
for study guide creation.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT

def set_cell_background(cell, hex_color):
    """Set cell background color using hex color string"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), hex_color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def set_cell_text(cell, text, bold=False, color=None, size=12, font_name='Calibri'):
    """Set cell text with formatting"""
    cell.text = text
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in paragraph.runs:
            run.font.size = Pt(size)
            run.font.name = font_name
            run.bold = bold
            if color:
                run.font.color.rgb = RGBColor(*color)

def create_comparison_table(doc, header_color_hex, header_color_rgb, header_text_color=(255, 255, 255), data_bg='FFFFFF'):
    """Create a 3-column comparison table template"""
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Inches(2.0)
        row.cells[1].width = Inches(2.5)
        row.cells[2].width = Inches(2.5)

    # Headers
    headers = ['Feature', 'Item 1', 'Item 2']
    for i, header_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_background(cell, header_color_hex)
        set_cell_text(cell, header_text, bold=True, color=header_text_color)

    # Row labels (first column)
    row_labels = ['Category 1', 'Category 2', 'Category 3', 'Category 4']
    light_bg = header_color_hex[0] + 'E' + header_color_hex[2:]  # Lighter version

    for row_idx, label in enumerate(row_labels, start=1):
        # First column (feature/category)
        set_cell_background(table.rows[row_idx].cells[0], light_bg)
        set_cell_text(table.rows[row_idx].cells[0], label, bold=True)

        # Data columns
        for col_idx in [1, 2]:
            set_cell_background(table.rows[row_idx].cells[col_idx], data_bg)
            set_cell_text(table.rows[row_idx].cells[col_idx], '')

    return table

def create_category_details_table(doc, header_color_hex, header_color_rgb, header_text_color=(255, 255, 255)):
    """Create a 2-column Category/Details table template"""
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Inches(2.0)
        row.cells[1].width = Inches(5.0)

    # Headers
    headers = ['Category', 'Details']
    for i, header_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_background(cell, header_color_hex)
        set_cell_text(cell, header_text, bold=True, color=header_text_color)

    # Row labels
    categories = ['Epidemiology', 'Risk Factors', 'Clinical Presentation', 'Diagnostic Criteria', 'Complications']
    light_bg = header_color_hex[0] + 'E' + header_color_hex[2:]  # Lighter version

    for row_idx, category in enumerate(categories, start=1):
        # Category column
        set_cell_background(table.rows[row_idx].cells[0], light_bg)
        set_cell_text(table.rows[row_idx].cells[0], category, bold=True)

        # Details column
        set_cell_background(table.rows[row_idx].cells[1], 'FFFFFF')
        set_cell_text(table.rows[row_idx].cells[1], '')

    return table

def create_master_chart_table(doc):
    """Create a master chart table template"""
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Inches(2.5)
        row.cells[1].width = Inches(4.5)

    # Headers - purple theme for master charts
    headers = ['Condition/Topic', 'Key Features']
    for i, header_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_background(cell, 'D1C4E9')  # Purple
        set_cell_text(cell, header_text, bold=True, color=(74, 20, 140))

    # Data rows
    for row_idx in range(1, 6):
        # Condition column
        set_cell_background(table.rows[row_idx].cells[0], 'EDE7F6')  # Light purple
        set_cell_text(table.rows[row_idx].cells[0], '', bold=True)

        # Details column
        set_cell_background(table.rows[row_idx].cells[1], 'FFFFFF')
        set_cell_text(table.rows[row_idx].cells[1], '')

    return table

def main():
    """Create the Word template with all table styles"""
    doc = Document()

    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Title
    title = doc.add_heading('Study Guide Template', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(118, 75, 162)

    # Instructions
    para = doc.add_paragraph()
    para.add_run('This template contains pre-formatted table styles for creating study guides.\n\n').font.size = Pt(11)
    para.add_run('How to use:\n').bold = True

    instructions = doc.add_paragraph(style='List Bullet')
    instructions.add_run('Insert → Table → Quick Tables → Select a pre-built table').font.size = Pt(11)

    instructions = doc.add_paragraph(style='List Bullet')
    instructions.add_run('Or create any table → Table Design → Apply a Study Guide style').font.size = Pt(11)

    doc.add_paragraph()

    # Section 1: Comparison Tables
    heading = doc.add_heading('Quick Tables - Comparison Format (3 columns)', 1)
    heading.runs[0].font.color.rgb = RGBColor(118, 75, 162)

    para = doc.add_paragraph('Use for comparing 2 items (e.g., Hodgkin vs Non-Hodgkin, Condition A vs Condition B)')
    para.runs[0].font.size = Pt(11)
    doc.add_paragraph()

    # Purple header comparison table
    doc.add_heading('Purple Header - General/Main Topics', 2)
    create_comparison_table(doc, 'D1C4E9', (209, 196, 233), (74, 20, 140))
    doc.add_paragraph()

    # Red header comparison table
    doc.add_heading('Red Header - Pathology/Abnormal Findings', 2)
    create_comparison_table(doc, 'FFCDD2', (255, 205, 210), (183, 28, 28))
    doc.add_paragraph()

    # Blue header comparison table
    doc.add_heading('Blue Header - Diagnostic/Exam Techniques', 2)
    create_comparison_table(doc, 'B3E5FC', (179, 229, 252), (1, 87, 155))
    doc.add_paragraph()

    # Green header comparison table
    doc.add_heading('Green Header - Normal Findings/Anatomy', 2)
    create_comparison_table(doc, 'C8E6C9', (200, 230, 201), (27, 94, 32))
    doc.add_paragraph()

    doc.add_page_break()

    # Section 2: Category/Details Tables
    heading = doc.add_heading('Quick Tables - Category/Details Format (2 columns)', 1)
    heading.runs[0].font.color.rgb = RGBColor(118, 75, 162)

    para = doc.add_paragraph('Use for defining/describing a single condition or concept')
    para.runs[0].font.size = Pt(11)
    doc.add_paragraph()

    # Purple header category table
    doc.add_heading('Purple Header - General/Main Topics', 2)
    create_category_details_table(doc, 'D1C4E9', (209, 196, 233), (74, 20, 140))
    doc.add_paragraph()

    # Red header category table
    doc.add_heading('Red Header - Pathology/Abnormal Findings', 2)
    create_category_details_table(doc, 'FFCDD2', (255, 205, 210), (183, 28, 28))
    doc.add_paragraph()

    # Blue header category table
    doc.add_heading('Blue Header - Diagnostic/Exam Techniques', 2)
    create_category_details_table(doc, 'B3E5FC', (179, 229, 252), (1, 87, 155))
    doc.add_paragraph()

    # Green header category table
    doc.add_heading('Green Header - Normal Findings/Anatomy', 2)
    create_category_details_table(doc, 'C8E6C9', (200, 230, 201), (27, 94, 32))
    doc.add_paragraph()

    doc.add_page_break()

    # Section 3: Master Chart
    heading = doc.add_heading('Quick Tables - Master Chart Format', 1)
    heading.runs[0].font.color.rgb = RGBColor(118, 75, 162)

    para = doc.add_paragraph('Use for comprehensive lists of all conditions/topics')
    para.runs[0].font.size = Pt(11)
    doc.add_paragraph()

    doc.add_heading('Master Chart Table', 2)
    create_master_chart_table(doc)
    doc.add_paragraph()

    # Save as regular .docx (python-docx doesn't support .dotx properly)
    output_path = '/Users/kimnguyen/Documents/html-lo-workspace/Claude Templates/Study Guide Table Templates.docx'
    doc.save(output_path)
    print(f'✅ Template created: {output_path}')
    print('\nHow to use:')
    print('1. Open this file to see all pre-formatted table examples')
    print('2. Copy the table style you need from this file')
    print('3. Paste into your study guide document')
    print('4. Edit the content as needed - all formatting is already applied!')

if __name__ == '__main__':
    main()
