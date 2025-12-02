#!/usr/bin/env python3
"""
Word Learning Objectives Study Guide - Bacterial Morphology & Virulence Factors
Source: 1 Intro-Bacterial Morphology microbiology_text.txt
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# =============================================================================
# HELPER FUNCTIONS
# =============================================================================

def set_cell_background(cell, color):
    """Set cell background color using hex"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def set_cell_text(cell, text, bold=False, size=11):
    """Set cell text - always black for readability"""
    cell.text = text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(size)
            run.font.name = 'Calibri'
            run.bold = bold

def add_colored_box(doc, title, content_list, title_color, bg_color='F3E5F5'):
    """Add colored information box"""
    # Title
    para = doc.add_paragraph()
    run = para.add_run(title)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(*title_color)

    # Content box
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.3)
    para.paragraph_format.right_indent = Inches(0.3)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)

    # Background color
    pPr = para._element.get_or_add_pPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), bg_color)
    pPr.append(shading)

    for item in content_list:
        run = para.add_run(f'• {item}\n')
        run.font.size = Pt(11)
        run.font.name = 'Calibri'

# =============================================================================
# MAIN DOCUMENT CREATION
# =============================================================================

def create_bacterial_morphology_guide(output_path):
    """Create Word LO study guide for bacterial morphology"""

    doc = Document()

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # ==========================================================================
    # TITLE PAGE
    # ==========================================================================
    title = doc.add_heading('Bacterial Morphology & Virulence Factors', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(118, 75, 162)
    title.runs[0].font.size = Pt(20)

    subtitle = doc.add_paragraph('Session 1 - Introduction to Microbiology')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(118, 75, 162)

    doc.add_page_break()

    # ==========================================================================
    # SECTION 1: LEARNING OBJECTIVES
    # ==========================================================================
    heading1 = doc.add_heading('Learning Objectives', 1)
    heading1.runs[0].font.color.rgb = RGBColor(118, 75, 162)

    # --- Learning Objective 1 ---
    obj1 = doc.add_heading('Learning Objective 1:', 2)
    obj1.runs[0].font.color.rgb = RGBColor(118, 75, 162)
    doc.add_paragraph('Describe and visually identify the various shapes and arrangements of bacterial cells.')

    # Summary
    summary_para = doc.add_paragraph()
    summary_run = summary_para.add_run('Summary: ')
    summary_run.bold = True
    summary_para.add_run('Bacteria are microscopic prokaryotic organisms ranging from 0.2-2.0 µm in diameter. They exhibit distinct shapes: cocci (spheres), bacilli (rods), vibrios (comma-shaped), spirilla (rigid spirals), spirochetes (flexible spirals), and coccobacilli (short rods/ovals). Some bacteria are pleomorphic with no defined shape. Bacterial arrangements depend on the plane of cell division: diplo- (pairs), strepto- (chains), staphylo- (grape-like clusters), tetrads (packets of 4), and sarcina (packets of 8). These morphological features are important for laboratory identification and diagnosis of infectious diseases.')

    doc.add_paragraph()

    # TABLE 1: Bacterial Shapes
    doc.add_heading('TABLE 1: Bacterial Shapes', 3)
    table1 = doc.add_table(rows=7, cols=2)
    table1.style = 'Table Grid'

    # Headers
    headers = ['Shape', 'Description']
    for i, header_text in enumerate(headers):
        set_cell_background(table1.rows[0].cells[i], 'C8E6C9')  # Green for anatomy/structure
        set_cell_text(table1.rows[0].cells[i], header_text, bold=True, size=12)

    # Data rows
    row_data = [
        ['Coccus (plural: cocci)', 'Spheres/round'],
        ['Bacillus (plural: bacilli)', 'Rod-shaped'],
        ['Vibrio (plural: vibrios)', 'Comma-shaped'],
        ['Spirillum (plural: spirilla)', 'Rigid spiral rod'],
        ['Spirochete (plural: spirochetes)', 'Flexible spiral rod'],
        ['Coccobacillus', 'Short rod/oval'],
    ]

    for row_idx, row_content in enumerate(row_data, start=1):
        cells = table1.rows[row_idx].cells
        set_cell_background(cells[0], 'E8F5E9')  # Light green
        set_cell_text(cells[0], row_content[0], bold=True)
        set_cell_background(cells[1], 'FFFFFF')
        set_cell_text(cells[1], row_content[1])

    doc.add_paragraph()

    # TABLE 2: Bacterial Arrangements
    doc.add_heading('TABLE 2: Bacterial Arrangements', 3)
    table2 = doc.add_table(rows=6, cols=3)
    table2.style = 'Table Grid'

    # Headers
    headers = ['Arrangement', 'Description', 'Plane of Division']
    for i, header_text in enumerate(headers):
        set_cell_background(table2.rows[0].cells[i], 'C8E6C9')
        set_cell_text(table2.rows[0].cells[i], header_text, bold=True, size=12)

    # Data rows
    row_data = [
        ['Diplo-', 'Pairs', 'One plane'],
        ['Strepto-', 'Chains', 'One plane'],
        ['Staphylo-', 'Grape-like clusters', 'Two or more planes'],
        ['Tetrads', 'Packets of 4 cells', 'Two planes'],
        ['Sarcina', 'Packets of 8 cells', 'Three planes'],
    ]

    for row_idx, row_content in enumerate(row_data, start=1):
        cells = table2.rows[row_idx].cells
        set_cell_background(cells[0], 'E8F5E9')
        set_cell_text(cells[0], row_content[0], bold=True)
        for col_idx in range(1, 3):
            set_cell_background(cells[col_idx], 'FFFFFF')
            set_cell_text(cells[col_idx], row_content[col_idx])

    doc.add_paragraph()

    # Clinical Pearls Box
    add_colored_box(doc, 'Clinical Pearls & High-Yield Points:', [
        'Morphology and arrangement are key steps in laboratory identification of bacteria from clinical specimens',
        'Gram-positive cocci in clusters (staphylo-) = think Staphylococcus aureus or coagulase-negative staph',
        'Gram-positive cocci in chains (strepto-) = think Streptococcus species (pyogenes, pneumoniae, agalactiae)',
        'Gram-negative diplococci = think Neisseria species (gonorrhoeae or meningitidis)',
        'Size range 0.2-2.0 µm makes bacteria visible under light microscopy (unlike viruses which require EM)'
    ], (0, 77, 64), 'E0F2F1')

    doc.add_paragraph()

    # Memory Tricks Box
    add_colored_box(doc, 'Memory Tricks & Mnemonics:', [
        '"Diplo- = Duo = 2" (pairs of bacteria)',
        '"Strepto- = String = chains" (like beads on a string)',
        '"Staphylo- = Staph like grapes" (grape-like clusters)',
        '"Cocci = Cue balls = round" (pool balls are spherical)',
        '"Bacilli = Baseball bats = rods" (both are rod-shaped)',
        'For spirochetes: "Lyme Steals Syphilis" - Leptospira, Treponema, Borrelia are the 3 main spirochetes',
        '[Researched from medical education sources]'
    ], (230, 81, 0), 'FFF3E0')

    doc.add_paragraph()

    # Analogy Box
    add_colored_box(doc, 'Analogy:', [
        'Think of bacterial shapes like different pasta types. Cocci are like small round pearl couscous. Bacilli are like penne or rigatoni tubes. Spirochetes are like fusilli spirals that twist and turn. Just as you can identify pasta by its shape, microbiologists identify bacteria by their characteristic forms under the microscope. And just as pasta can be arranged differently on a plate (single pieces, pairs, or clusters), bacteria arrange themselves in predictable patterns based on how they divide.',
        '[Researched analogy concept]'
    ], (118, 75, 162), 'F3E5F5')

    doc.add_page_break()

    # --- Learning Objective 2 ---
    obj2 = doc.add_heading('Learning Objective 2:', 2)
    obj2.runs[0].font.color.rgb = RGBColor(118, 75, 162)
    doc.add_paragraph('Understand the basic structure, function(s) & clinical relevance (where applicable) of the structures that compose gram-negative, gram-positive, and acid-fast cell walls: Peptidoglycan (including general synthesis and the effect of beta-lactams and lysozyme on peptidoglycan), Outer membrane and LPS, Mycolic acids.')

    # Summary
    summary_para = doc.add_paragraph()
    summary_run = summary_para.add_run('Summary: ')
    summary_run.bold = True
    summary_para.add_run('Bacterial cell walls provide rigidity, shape, and protection from osmotic pressure. There are three main types: (1) Gram-positive walls have thick peptidoglycan with teichoic acid; (2) Gram-negative walls have thin peptidoglycan plus an outer membrane containing lipopolysaccharide (LPS/endotoxin); (3) Acid-fast walls have thin peptidoglycan plus a mycolic acid outer layer. Peptidoglycan is a polymer of NAG-NAM with tetrapeptide sidechains and pentapeptide cross-links. Beta-lactam antibiotics (penicillin) and lysozyme weaken cell walls by interfering with peptidoglycan synthesis/structure. LPS in Gram-negatives triggers immune responses; its lipid A component causes toxic shock. Mycolic acids in acid-fast bacteria provide resistance to desiccation, antibiotics, and phagocytosis.')

    doc.add_paragraph()

    # TABLE 1: Three Cell Wall Types Comparison
    doc.add_heading('TABLE 1: Comparison of Three Cell Wall Types', 3)
    table3 = doc.add_table(rows=6, cols=4)
    table3.style = 'Table Grid'

    # Headers
    headers = ['Feature', 'Gram-Positive', 'Gram-Negative', 'Acid-Fast']
    for i, header_text in enumerate(headers):
        set_cell_background(table3.rows[0].cells[i], 'B3E5FC')  # Blue for diagnostic
        set_cell_text(table3.rows[0].cells[i], header_text, bold=True, size=12)

    # Data rows
    row_data = [
        ['Peptidoglycan Layer', 'Thick', 'Thin', 'Thin'],
        ['Teichoic Acid', 'Present', 'Absent', 'Absent'],
        ['Outer Membrane', 'Absent', 'Present (with LPS)', 'Absent'],
        ['Mycolic Acid Layer', 'Absent', 'Absent', 'Present (outer layer)'],
        ['Examples', 'Staphylococcus, Streptococcus', 'E. coli, Salmonella', 'Mycobacterium tuberculosis'],
    ]

    for row_idx, row_content in enumerate(row_data, start=1):
        cells = table3.rows[row_idx].cells
        set_cell_background(cells[0], 'E1F5FE')
        set_cell_text(cells[0], row_content[0], bold=True)
        for col_idx in range(1, 4):
            set_cell_background(cells[col_idx], 'FFFFFF')
            set_cell_text(cells[col_idx], row_content[col_idx])

    doc.add_paragraph()

    # TABLE 2: Peptidoglycan Structure
    doc.add_heading('TABLE 2: Peptidoglycan Components', 3)
    table4 = doc.add_table(rows=4, cols=2)
    table4.style = 'Table Grid'

    # Headers
    headers = ['Component', 'Description']
    for i, header_text in enumerate(headers):
        set_cell_background(table4.rows[0].cells[i], 'C8E6C9')
        set_cell_text(table4.rows[0].cells[i], header_text, bold=True, size=12)

    # Data rows
    row_data = [
        ['Backbone', 'Alternating N-acetylglucosamine (NAG) and N-acetylmuramic acid (NAM)'],
        ['Tetrapeptide Sidechains', 'Set of identical tetrapeptide sidechains attached to NAM'],
        ['Pentapeptide Cross-links', 'Set of identical pentapeptide cross-links connecting chains'],
    ]

    for row_idx, row_content in enumerate(row_data, start=1):
        cells = table4.rows[row_idx].cells
        set_cell_background(cells[0], 'E8F5E9')
        set_cell_text(cells[0], row_content[0], bold=True)
        set_cell_background(cells[1], 'FFFFFF')
        set_cell_text(cells[1], row_content[1])

    doc.add_paragraph()

    # TABLE 3: LPS Structure
    doc.add_heading('TABLE 3: Lipopolysaccharide (LPS) Components', 3)
    table5 = doc.add_table(rows=4, cols=2)
    table5.style = 'Table Grid'

    # Headers
    headers = ['LPS Component', 'Function/Clinical Significance']
    for i, header_text in enumerate(headers):
        set_cell_background(table5.rows[0].cells[i], 'FFCDD2')  # Red for pathology
        set_cell_text(table5.rows[0].cells[i], header_text, bold=True, size=12)

    # Data rows
    row_data = [
        ['LPS (overall)', 'Also called endotoxin - signals immune system of Gram-negative bacterial invasion'],
        ['Lipid A', 'Causes toxic shock when large amounts accumulate in bloodstream'],
        ['O antigen', 'Immunogenic - recognized by the immune system'],
    ]

    for row_idx, row_content in enumerate(row_data, start=1):
        cells = table5.rows[row_idx].cells
        set_cell_background(cells[0], 'FFEBEE')
        set_cell_text(cells[0], row_content[0], bold=True)
        set_cell_background(cells[1], 'FFFFFF')
        set_cell_text(cells[1], row_content[1])

    doc.add_paragraph()

    # Clinical Pearls Box
    add_colored_box(doc, 'Clinical Pearls & High-Yield Points:', [
        'Beta-lactam antibiotics (penicillin, vancomycin) kill bacteria by interfering with peptidoglycan synthesis',
        'Lysozyme in tears, saliva, and immune cells weakens cell walls by breaking down peptidoglycan',
        'Gram-negative septic shock is caused by LPS (endotoxin) release - specifically the lipid A component',
        'Outer membrane in Gram-negatives acts as a protective barrier against many antibiotics',
        'Mycolic acid layer in acid-fast bacteria makes them resistant to standard Gram staining and many disinfectants',
        'Acid-fast bacteria (Mycobacterium) require special Ziehl-Neelsen staining due to waxy mycolic acid layer'
    ], (0, 77, 64), 'E0F2F1')

    doc.add_paragraph()

    # Memory Tricks Box
    add_colored_box(doc, 'Memory Tricks & Mnemonics:', [
        '"Long PowerPoint (PPT)" - LPS and outer membrane in Gram-negative; thick Peptidoglycan and Teichoic acids in Gram-positive [Researched mnemonic]',
        '"P for Purple" - Gram-Positive stains Purple (both start with P) [Researched mnemonic]',
        '"Positive = Thick, Negative = Thin" - Gram-positive has thick peptidoglycan, Gram-negative has thin [Researched mnemonic]',
        '"NAG-NAM makes you NAG!" - Peptidoglycan backbone alternates NAG and NAM',
        '"LiPid A = Lethal Poison Always" - Lipid A component of LPS causes toxic shock',
        '"Mycolic = My coat keeps me alive" - Mycolic acid layer protects acid-fast bacteria'
    ], (230, 81, 0), 'FFF3E0')

    doc.add_paragraph()

    # Analogy Box
    add_colored_box(doc, 'Analogy:', [
        'Think of peptidoglycan as a chain-link fence surrounding the bacterial cell. The NAG-NAM backbone is like the horizontal wires, and the peptide cross-links are the vertical posts connecting everything together. Beta-lactam antibiotics are like wire cutters that snip the cross-links, weakening the fence until it collapses. Gram-positive bacteria have a thick, multilayer fence (thick peptidoglycan), while Gram-negative bacteria have a thin fence but add a second outer wall (outer membrane with LPS) like a brick wall outside the chain-link fence. Acid-fast bacteria coat their thin fence with a waxy waterproof layer (mycolic acids) like applying waterproof sealant.',
        '[Researched analogy concept]'
    ], (118, 75, 162), 'F3E5F5')

    doc.add_page_break()

    # --- Learning Objective 3 ---
    obj3 = doc.add_heading('Learning Objective 3:', 2)
    obj3.runs[0].font.color.rgb = RGBColor(118, 75, 162)
    doc.add_paragraph('How to perform and interpret Gram-stain and acid-fast stains?')

    # Summary
    summary_para = doc.add_paragraph()
    summary_run = summary_para.add_run('Summary: ')
    summary_run.bold = True
    summary_para.add_run('The Gram stain is a differential stain developed by Hans Christian Gram in 1884 using crystal violet (primary stain) and safranin (counterstain). Gram-positive bacteria retain crystal violet and appear purple due to their thick peptidoglycan layer. Gram-negative bacteria lose crystal violet during decolorization and take up safranin, appearing red/pink. Clinical relevance includes initial bacterial identification, determining which additional tests are needed, and initiating empiric antibiotic therapy before final identification. Acid-fast staining (Ziehl-Neelsen stain) is used for Mycobacterium species that do not stain well with Gram stain due to their waxy mycolic acid layer. Acid-fast bacteria retain the primary stain after acid-alcohol treatment and appear red.')

    doc.add_paragraph()

    # TABLE 1: Gram Stain Procedure and Interpretation
    doc.add_heading('TABLE 1: Gram Stain - Procedure and Results', 3)
    table6 = doc.add_table(rows=4, cols=3)
    table6.style = 'Table Grid'

    # Headers
    headers = ['Feature', 'Gram-Positive Result', 'Gram-Negative Result']
    for i, header_text in enumerate(headers):
        set_cell_background(table6.rows[0].cells[i], 'B3E5FC')
        set_cell_text(table6.rows[0].cells[i], header_text, bold=True, size=12)

    # Data rows
    row_data = [
        ['Primary Stain', 'Crystal violet - RETAINED', 'Crystal violet - LOST during decolorization'],
        ['Counterstain', 'Safranin (not visible)', 'Safranin - TAKEN UP'],
        ['Final Color', 'Purple', 'Red/Pink'],
    ]

    for row_idx, row_content in enumerate(row_data, start=1):
        cells = table6.rows[row_idx].cells
        set_cell_background(cells[0], 'E1F5FE')
        set_cell_text(cells[0], row_content[0], bold=True)
        for col_idx in range(1, 3):
            set_cell_background(cells[col_idx], 'FFFFFF')
            set_cell_text(cells[col_idx], row_content[col_idx])

    doc.add_paragraph()

    # TABLE 2: Clinical Relevance of Gram Stain
    doc.add_heading('TABLE 2: Clinical Relevance of Gram Stain', 3)
    table7 = doc.add_table(rows=4, cols=2)
    table7.style = 'Table Grid'

    # Headers
    headers = ['Clinical Use', 'Details']
    for i, header_text in enumerate(headers):
        set_cell_background(table7.rows[0].cells[i], 'B2DFDB')  # Teal for procedures
        set_cell_text(table7.rows[0].cells[i], header_text, bold=True, size=12)

    # Data rows
    row_data = [
        ['Initial Identification', 'First step in identifying bacteria from clinical specimens'],
        ['Guide Additional Testing', 'Indicates which additional tests need to be performed for definitive ID'],
        ['Empiric Antibiotic Therapy', 'Allows initiation of antibiotic therapy even before final identification is made'],
    ]

    for row_idx, row_content in enumerate(row_data, start=1):
        cells = table7.rows[row_idx].cells
        set_cell_background(cells[0], 'E0F2F1')
        set_cell_text(cells[0], row_content[0], bold=True)
        set_cell_background(cells[1], 'FFFFFF')
        set_cell_text(cells[1], row_content[1])

    doc.add_paragraph()

    # TABLE 3: Acid-Fast Stain
    doc.add_heading('TABLE 3: Acid-Fast Stain (Ziehl-Neelsen)', 3)
    table8 = doc.add_table(rows=5, cols=2)
    table8.style = 'Table Grid'

    # Headers
    headers = ['Feature', 'Details']
    for i, header_text in enumerate(headers):
        set_cell_background(table8.rows[0].cells[i], 'FFE0B2')  # Orange for special stains
        set_cell_text(table8.rows[0].cells[i], header_text, bold=True, size=12)

    # Data rows
    row_data = [
        ['Purpose', 'Used for Mycobacterium species (M. tuberculosis, M. leprae)'],
        ['Why Needed', 'Acid-fast bacteria do not stain well with Gram stain due to waxy mycolic acid layer'],
        ['Procedure', 'Uses carbolfuchsin (primary stain) with heat, acid-alcohol decolorization, methylene blue counterstain'],
        ['Result', 'Acid-fast bacteria appear RED; non-acid-fast bacteria appear BLUE'],
    ]

    for row_idx, row_content in enumerate(row_data, start=1):
        cells = table8.rows[row_idx].cells
        set_cell_background(cells[0], 'FFF3E0')
        set_cell_text(cells[0], row_content[0], bold=True)
        set_cell_background(cells[1], 'FFFFFF')
        set_cell_text(cells[1], row_content[1])

    doc.add_paragraph()

    # Clinical Pearls Box
    add_colored_box(doc, 'Clinical Pearls & High-Yield Points:', [
        'If Gram stain shows Gram-positive cocci in clusters, start empiric anti-staphylococcal therapy (e.g., vancomycin if MRSA suspected)',
        'If CSF Gram stain shows Gram-negative diplococci, think Neisseria meningitidis → start ceftriaxone immediately',
        'Acid-fast positive sputum = presumptive TB diagnosis → start respiratory isolation and anti-TB therapy pending culture',
        'Some bacteria are "Gram-variable" or stain poorly with Gram stain (Mycoplasma has no cell wall, Legionella stains poorly)',
        'Quality of Gram stain depends on specimen quality and technician skill - negative Gram stain does not rule out bacterial infection'
    ], (0, 77, 64), 'E0F2F1')

    doc.add_paragraph()

    # Memory Tricks Box
    add_colored_box(doc, 'Memory Tricks & Mnemonics:', [
        '"Gram-Positive = Purple = P & P" [Researched mnemonic]',
        '"Gram-Negative = piNk/red = N in both words" [Researched mnemonic]',
        '"Safranin = Safety Net for Negatives" - Safranin is the counterstain that colors Gram-negatives',
        '"Crystal violet Sticks in Thick peptidoglycan" - Why Gram-positives retain the purple stain',
        '"Acid-Fast = Always Find Mycobacteria" - Acid-fast stain is specifically for Mycobacterium',
        '"TB = Think about Being Red" - Tuberculosis bacteria appear red on acid-fast stain'
    ], (230, 81, 0), 'FFF3E0')

    doc.add_paragraph()

    # Analogy Box
    add_colored_box(doc, 'Analogy:', [
        'Think of Gram staining like dying fabric. Gram-positive bacteria are like thick wool sweaters that soak up purple dye and hold it tightly even when you try to wash it out (decolorization step). Gram-negative bacteria are like thin silk shirts - they absorb the purple dye initially, but it washes right out when you rinse them. Since the silk is now colorless, you add a second dye (safranin/pink) so you can see it. The thick wool sweater was already so dark purple that the pink dye doesn\'t show up on it.',
        '[Researched analogy concept]'
    ], (118, 75, 162), 'F3E5F5')

    doc.add_page_break()

    # --- Learning Objective 4 ---
    obj4 = doc.add_heading('Learning Objective 4:', 2)
    obj4.runs[0].font.color.rgb = RGBColor(118, 75, 162)
    doc.add_paragraph('Describe the structures and functions of bacterial external and internal cell components.')

    # Summary
    summary_para = doc.add_paragraph()
    summary_run = summary_para.add_run('Summary: ')
    summary_run.bold = True
    summary_para.add_run('External structures: (1) Glycocalyx (capsule or slime layer) promotes adherence for biofilm formation, protects against desiccation and toxins, and inhibits phagocytosis; (2) Flagella are protein filaments that provide motility with arrangements including monotrichous (single polar), lophotrichous (multiple polar), amphitrichous (both poles), peritrichous (distributed over entire cell), or atrichous (no flagella); (3) Pili/fimbriae are rigid hair-like structures - ordinary pili serve as adhesins for attachment to host cells, while sex pili mediate conjugation (genetic exchange) in Gram-negatives. Internal structures: (1) Nucleoid contains a single circular chromosome (supercoiled dsDNA); (2) Plasmids are smaller circular dsDNA that can carry antibiotic resistance genes; (3) 70S ribosomes (30S + 50S subunits) perform protein synthesis and are antibiotic targets; (4) Endospores are dormant cells produced by Bacillus, Clostridium, and Clostridioides that are extremely resistant to heat, desiccation, chemicals, and radiation.')

    doc.add_paragraph()

    # TABLE 1: External Structures
    doc.add_heading('TABLE 1: Bacterial External Structures', 3)
    table9 = doc.add_table(rows=4, cols=3)
    table9.style = 'Table Grid'

    # Headers
    headers = ['Structure', 'Function', 'Clinical Significance']
    for i, header_text in enumerate(headers):
        set_cell_background(table9.rows[0].cells[i], 'C8E6C9')
        set_cell_text(table9.rows[0].cells[i], header_text, bold=True, size=12)

    # Data rows
    row_data = [
        ['Glycocalyx (Capsule/Slime Layer)', 'Adherence (biofilm), protection from desiccation, barrier to toxins', 'Capsule inhibits phagocytosis (virulence factor); biofilm formation on medical devices'],
        ['Flagella', 'Motility', 'Allows bacteria to move toward nutrients, away from toxins; helps colonization'],
        ['Pili/Fimbriae', 'Ordinary pili: adherence to host cells; Sex pilus: conjugation (gene transfer)', 'Adhesion to host tissues enables colonization; conjugation spreads antibiotic resistance'],
    ]

    for row_idx, row_content in enumerate(row_data, start=1):
        cells = table9.rows[row_idx].cells
        set_cell_background(cells[0], 'E8F5E9')
        set_cell_text(cells[0], row_content[0], bold=True)
        for col_idx in range(1, 3):
            set_cell_background(cells[col_idx], 'FFFFFF')
            set_cell_text(cells[col_idx], row_content[col_idx])

    doc.add_paragraph()

    # TABLE 2: Flagellar Arrangements
    doc.add_heading('TABLE 2: Flagellar Arrangements', 3)
    table10 = doc.add_table(rows=6, cols=2)
    table10.style = 'Table Grid'

    # Headers
    headers = ['Arrangement', 'Description']
    for i, header_text in enumerate(headers):
        set_cell_background(table10.rows[0].cells[i], 'FFE0B2')
        set_cell_text(table10.rows[0].cells[i], header_text, bold=True, size=12)

    # Data rows
    row_data = [
        ['Atrichous', 'No flagellum'],
        ['Monotrichous', 'Single polar flagellum (one end)'],
        ['Lophotrichous', 'Multiple polar flagella (tuft at one end)'],
        ['Amphitrichous', 'Single flagellum or tuft at two opposite poles'],
        ['Peritrichous', 'Multiple flagella distributed over entire cell'],
    ]

    for row_idx, row_content in enumerate(row_data, start=1):
        cells = table10.rows[row_idx].cells
        set_cell_background(cells[0], 'FFF3E0')
        set_cell_text(cells[0], row_content[0], bold=True)
        set_cell_background(cells[1], 'FFFFFF')
        set_cell_text(cells[1], row_content[1])

    doc.add_paragraph()

    # TABLE 3: Internal Structures
    doc.add_heading('TABLE 3: Bacterial Internal Structures', 3)
    table11 = doc.add_table(rows=5, cols=3)
    table11.style = 'Table Grid'

    # Headers
    headers = ['Structure', 'Description', 'Clinical Significance']
    for i, header_text in enumerate(headers):
        set_cell_background(table11.rows[0].cells[i], 'D1C4E9')  # Purple for general topics
        set_cell_text(table11.rows[0].cells[i], header_text, bold=True, size=12)

    # Data rows
    row_data = [
        ['Nucleoid (Chromosome)', 'Single circular double-stranded DNA; supercoiled with binding proteins', 'Contains essential genes for bacterial survival and replication'],
        ['Plasmids', 'Circular supercoiled dsDNA; smaller than chromosome', 'Often carry antibiotic resistance genes; spread via conjugation'],
        ['Ribosomes (70S)', 'Composed of 30S + 50S subunits', 'Target of many antibiotics (do not affect human 80S ribosomes)'],
        ['Endospores', 'Dormant cells; extremely resistant to environmental stresses', 'Produced by Bacillus, Clostridium, Clostridioides; survive harsh conditions including boiling'],
    ]

    for row_idx, row_content in enumerate(row_data, start=1):
        cells = table11.rows[row_idx].cells
        set_cell_background(cells[0], 'F3E5F5')
        set_cell_text(cells[0], row_content[0], bold=True)
        for col_idx in range(1, 3):
            set_cell_background(cells[col_idx], 'FFFFFF')
            set_cell_text(cells[col_idx], row_content[col_idx])

    doc.add_paragraph()

    # Clinical Pearls Box
    add_colored_box(doc, 'Clinical Pearls & High-Yield Points:', [
        'Encapsulated bacteria (e.g., S. pneumoniae, H. influenzae, N. meningitidis) are more virulent because capsules prevent phagocytosis',
        'Biofilm formation on indwelling catheters and prosthetic devices makes infections extremely difficult to eradicate',
        'Plasmid-mediated antibiotic resistance can spread rapidly between bacteria via conjugation (horizontal gene transfer)',
        'Antibiotics targeting 70S ribosomes (e.g., tetracyclines, aminoglycosides, macrolides) selectively kill bacteria without harming human cells',
        'Endospore-forming bacteria (Clostridium difficile, C. tetani, C. botulinum, Bacillus anthracis) cause serious infections and require special sterilization procedures',
        'Endospores can survive boiling water (100°C) but are killed by autoclaving (121°C at 15 psi for 15-20 minutes)'
    ], (0, 77, 64), 'E0F2F1')

    doc.add_paragraph()

    # Memory Tricks Box
    add_colored_box(doc, 'Memory Tricks & Mnemonics:', [
        'Flagella arrangements: "Mono- = one, Amphi- = both sides, Lopho- = tuft, Peri- = around" [Etymology-based mnemonic]',
        '"BaD Ox closed geo tomato inside" - Endospore formers: Bacillus, Clostridium, Coxiella, Desulfotomaculum, Geobacillus [Researched mnemonic]',
        '"70S ribosomes = 7 letters in SMALLER (bacterial ribosomes are smaller than 80S eukaryotic)"',
        '"Plasmids = Portable Antibiotic Resistance" - Plasmids carry and transfer resistance genes',
        '"Capsule = Slippery Coat = Can\'t Catch" - Encapsulated bacteria evade phagocytosis',
        '"Pili = Sticky Fingers" - Pili help bacteria stick to surfaces and host cells'
    ], (230, 81, 0), 'FFF3E0')

    doc.add_paragraph()

    # Analogy Box
    add_colored_box(doc, 'Analogy:', [
        'Think of a bacterial cell like a tiny submarine. The flagella are propellers that move it through water. The pili are like grappling hooks that let it attach to surfaces or other submarines. The capsule is a slippery outer coating like oil on the hull that makes it hard for enemy ships (phagocytes) to grab. Inside the submarine, the chromosome is the main control manual (permanent instructions), while plasmids are like removable USB drives with extra software - including programs that make the submarine resistant to depth charges (antibiotics). The ribosomes are the manufacturing plants inside making all the proteins the submarine needs. And endospores? They\'re like emergency life pods that can survive in the vacuum of space and wake up decades later when conditions improve.',
        '[Researched analogy concept]'
    ], (118, 75, 162), 'F3E5F5')

    doc.add_page_break()

    # --- Learning Objective 5 ---
    obj5 = doc.add_heading('Learning Objective 5:', 2)
    obj5.runs[0].font.color.rgb = RGBColor(118, 75, 162)
    doc.add_paragraph('Define virulence factors, and understand the basic functions of various bacterial virulence factors.')

    # Summary
    summary_para = doc.add_paragraph()
    summary_run = summary_para.add_run('Summary: ')
    summary_run.bold = True
    summary_para.add_run('Virulence factors are properties (gene products) that enable microbes to establish themselves on/within a host and enhance their potential to cause disease. Main categories include: (1) Adhesins - proteins/glycoproteins that bind pathogens to host cell receptors (examples: fimbriae/pili for tissue-specific attachment, glycocalyx for biofilm formation); (2) Invasins - proteins/enzymes that damage host cells locally and allow spread through tissues (examples: collagenase breaks down collagen, neuraminidase degrades intracellular "glue", streptokinase/staphylokinase break down blood clots); (3) Toxins - high-activity proteins with specificity (exotoxins are secreted like diphtheria toxin, botulinum toxin, cholera enterotoxin; endotoxins are LPS in Gram-negative outer membranes with lipid A causing toxicity); (4) Antiphagocytic factors - capsules make bacterial surfaces slippery so phagocytes cannot grab them, proteins like staphylococcal protein A and streptococcal protein M prevent phagocytosis.')

    doc.add_paragraph()

    # TABLE 1: Categories of Virulence Factors
    doc.add_heading('TABLE 1: Four Main Categories of Virulence Factors', 3)
    table12 = doc.add_table(rows=5, cols=3)
    table12.style = 'Table Grid'

    # Headers
    headers = ['Virulence Factor', 'Function', 'Examples']
    for i, header_text in enumerate(headers):
        set_cell_background(table12.rows[0].cells[i], 'FFCDD2')  # Red for pathology
        set_cell_text(table12.rows[0].cells[i], header_text, bold=True, size=12)

    # Data rows
    row_data = [
        ['Adhesins', 'Bind pathogens to receptors on host cells and tissues', 'Fimbriae, pili (E. coli), glycocalyx (S. aureus capsule), biofilms'],
        ['Invasins', 'Proteins/enzymes that damage host cells locally and allow tissue spread', 'Collagenase, neuraminidase, streptokinase, staphylokinase'],
        ['Toxins', 'Proteins with high activity and specificity that cause damage', 'Exotoxins: diphtheria toxin, botulinum toxin, cholera toxin; Endotoxins: LPS (lipid A)'],
        ['Antiphagocytic Factors', 'Prevent phagocytosis by immune cells', 'Capsule (slippery surface), protein A (Staph), protein M (Strep)'],
    ]

    for row_idx, row_content in enumerate(row_data, start=1):
        cells = table12.rows[row_idx].cells
        set_cell_background(cells[0], 'FFEBEE')
        set_cell_text(cells[0], row_content[0], bold=True)
        for col_idx in range(1, 3):
            set_cell_background(cells[col_idx], 'FFFFFF')
            set_cell_text(cells[col_idx], row_content[col_idx])

    doc.add_paragraph()

    # TABLE 2: Adhesins Detail
    doc.add_heading('TABLE 2: Adhesins - Attachment Mechanisms', 3)
    table13 = doc.add_table(rows=4, cols=3)
    table13.style = 'Table Grid'

    # Headers
    headers = ['Type of Adhesin', 'Mechanism', 'Clinical Example']
    for i, header_text in enumerate(headers):
        set_cell_background(table13.rows[0].cells[i], 'FFE0B2')
        set_cell_text(table13.rows[0].cells[i], header_text, bold=True, size=12)

    # Data rows
    row_data = [
        ['Fimbriae/Pili', 'Attach to specific host molecules (tissue tropism)', 'E. coli attaches to intestinal epithelial cells via pili'],
        ['Capsule', 'Glycocalyx material aids adherence', 'S. aureus produces capsule for attachment'],
        ['Slime Layer', 'Creates biofilms on surfaces', 'Dental plaque; catheter-associated infections'],
    ]

    for row_idx, row_content in enumerate(row_data, start=1):
        cells = table13.rows[row_idx].cells
        set_cell_background(cells[0], 'FFF3E0')
        set_cell_text(cells[0], row_content[0], bold=True)
        for col_idx in range(1, 3):
            set_cell_background(cells[col_idx], 'FFFFFF')
            set_cell_text(cells[col_idx], row_content[col_idx])

    doc.add_paragraph()

    # TABLE 3: Invasins Detail
    doc.add_heading('TABLE 3: Invasins - Tissue Invasion Enzymes', 3)
    table14 = doc.add_table(rows=4, cols=2)
    table14.style = 'Table Grid'

    # Headers
    headers = ['Invasin Enzyme', 'Function']
    for i, header_text in enumerate(headers):
        set_cell_background(table14.rows[0].cells[i], 'FFCDD2')
        set_cell_text(table14.rows[0].cells[i], header_text, bold=True, size=12)

    # Data rows
    row_data = [
        ['Collagenase', 'Breaks down collagen in connective tissue'],
        ['Neuraminidase', 'Degrades neuraminic acid (intracellular "glue")'],
        ['Streptokinase/Staphylokinase', 'Fibrinolysins - break down blood clots'],
    ]

    for row_idx, row_content in enumerate(row_data, start=1):
        cells = table14.rows[row_idx].cells
        set_cell_background(cells[0], 'FFEBEE')
        set_cell_text(cells[0], row_content[0], bold=True)
        set_cell_background(cells[1], 'FFFFFF')
        set_cell_text(cells[1], row_content[1])

    doc.add_paragraph()

    # TABLE 4: Toxins Detail
    doc.add_heading('TABLE 4: Bacterial Toxins', 3)
    table15 = doc.add_table(rows=3, cols=3)
    table15.style = 'Table Grid'

    # Headers
    headers = ['Toxin Type', 'Characteristics', 'Examples']
    for i, header_text in enumerate(headers):
        set_cell_background(table15.rows[0].cells[i], 'FFCDD2')
        set_cell_text(table15.rows[0].cells[i], header_text, bold=True, size=12)

    # Data rows
    row_data = [
        ['Exotoxins/Enterotoxins', 'Secreted proteins with high specificity', 'Diphtheria toxin (C. diphtheriae), Botulinum toxin (C. botulinum), Cholera toxin (V. cholerae)'],
        ['Endotoxins', 'LPS in Gram-negative outer membrane; lipid A component causes toxicity', 'LPS from any Gram-negative bacteria (E. coli, Salmonella, etc.)'],
    ]

    for row_idx, row_content in enumerate(row_data, start=1):
        cells = table15.rows[row_idx].cells
        set_cell_background(cells[0], 'FFEBEE')
        set_cell_text(cells[0], row_content[0], bold=True)
        for col_idx in range(1, 3):
            set_cell_background(cells[col_idx], 'FFFFFF')
            set_cell_text(cells[col_idx], row_content[col_idx])

    doc.add_paragraph()

    # Clinical Pearls Box
    add_colored_box(doc, 'Clinical Pearls & High-Yield Points:', [
        'Virulence factors follow the pathogenesis sequence: Attach (adhesins) → Invade (invasins) → Damage (toxins) → Evade (antiphagocytic factors)',
        'Tissue tropism is mediated by adhesins - bacteria attach only where specific receptors exist (e.g., E. coli pili bind mannose on intestinal cells)',
        'Gram-negative septic shock is caused by endotoxin (LPS) release, specifically the lipid A component',
        'Exotoxins are typically more potent than endotoxins and are specific to certain bacterial species',
        'Invasins like streptokinase and staphylokinase break down clots, allowing bacteria to spread through tissues and bloodstream',
        'Encapsulated bacteria are especially dangerous in asplenic patients who lack proper opsonization and phagocytosis'
    ], (0, 77, 64), 'E0F2F1')

    doc.add_paragraph()

    # Memory Tricks Box
    add_colored_box(doc, 'Memory Tricks & Mnemonics:', [
        '"AIT - Attach, Invade, cause Toxic damage" - Sequence of virulence factors (Adhesins → Invasins → Toxins) [Researched concept]',
        '"Exotoxins are EXcreted; Endotoxins are ENclosed (in the membrane)"',
        '"LiPid A = Lethal Poison Always" - Lipid A component of LPS endotoxin causes toxic shock',
        '"Adhesins = Adhesive tape = Stick to surfaces"',
        '"Invasins = Invaders with scissors" - Enzymes that cut through tissues',
        '"Capsule = Slime coat = Slippery = Can\'t be grabbed by phagocytes"',
        '"Collagenase = Cuts Collagen; Streptokinase = Strips clots; Neuraminidase = Nukes the glue"'
    ], (230, 81, 0), 'FFF3E0')

    doc.add_paragraph()

    # Analogy Box
    add_colored_box(doc, 'Analogy:', [
        'Think of bacterial virulence factors like tools for a burglar breaking into a house. Adhesins are like grappling hooks that let the burglar attach to the building and climb up. Invasins are like bolt cutters and crowbars that break through locks and doors to get deeper inside. Toxins are like poison gas or explosives that damage the house and its occupants. Antiphagocytic factors are like wearing a slippery raincoat covered in oil - when security guards (phagocytes) try to grab the burglar, their hands slip right off and they can\'t catch him. The more virulence factors a bacteria has, the more successful the "break-in" (infection) will be.',
        '[Researched analogy concept]'
    ], (118, 75, 162), 'F3E5F5')

    doc.add_page_break()

    # ==========================================================================
    # SECTION 2: KEY COMPARISONS
    # ==========================================================================
    heading2 = doc.add_heading('Key Comparisons', 1)
    heading2.runs[0].font.color.rgb = RGBColor(118, 75, 162)

    # Comparison 1: Gram-Positive vs Gram-Negative
    doc.add_heading('Gram-Positive vs Gram-Negative Cell Envelopes', 2)
    table_comp1 = doc.add_table(rows=8, cols=3)
    table_comp1.style = 'Table Grid'

    headers = ['Feature', 'Gram-Positive', 'Gram-Negative']
    for i, header_text in enumerate(headers):
        set_cell_background(table_comp1.rows[0].cells[i], 'B3E5FC')
        set_cell_text(table_comp1.rows[0].cells[i], header_text, bold=True, size=12)

    row_data = [
        ['Peptidoglycan Layer', 'Thick (multiple layers)', 'Thin (single layer)'],
        ['Teichoic Acid', 'Present', 'Absent'],
        ['Outer Membrane', 'Absent', 'Present'],
        ['LPS (Endotoxin)', 'Absent', 'Present (lipid A causes toxicity)'],
        ['Gram Stain Color', 'Purple (retains crystal violet)', 'Red/Pink (takes up safranin)'],
        ['Antibiotic Susceptibility', 'More susceptible to lysozyme and penicillin', 'Outer membrane provides barrier to many antibiotics'],
        ['Examples', 'Staphylococcus, Streptococcus, Bacillus', 'E. coli, Salmonella, Pseudomonas, Neisseria'],
    ]

    for row_idx, row_content in enumerate(row_data, start=1):
        cells = table_comp1.rows[row_idx].cells
        set_cell_background(cells[0], 'E1F5FE')
        set_cell_text(cells[0], row_content[0], bold=True)
        set_cell_background(cells[1], 'FFFFFF')
        set_cell_text(cells[1], row_content[1])
        set_cell_background(cells[2], 'FFFFFF')
        set_cell_text(cells[2], row_content[2])

    doc.add_paragraph()

    # Comparison 2: Exotoxins vs Endotoxins
    doc.add_heading('Exotoxins vs Endotoxins', 2)
    table_comp2 = doc.add_table(rows=7, cols=3)
    table_comp2.style = 'Table Grid'

    headers = ['Feature', 'Exotoxins', 'Endotoxins']
    for i, header_text in enumerate(headers):
        set_cell_background(table_comp2.rows[0].cells[i], 'FFCDD2')
        set_cell_text(table_comp2.rows[0].cells[i], header_text, bold=True, size=12)

    row_data = [
        ['Nature', 'Secreted proteins', 'Lipopolysaccharide (LPS) in outer membrane'],
        ['Source', 'Mostly Gram-positive (some Gram-negative)', 'Gram-negative bacteria only'],
        ['Release', 'Actively secreted during growth', 'Released when bacteria die/lyse'],
        ['Toxicity', 'Highly toxic in small amounts', 'Moderately toxic; requires larger amounts'],
        ['Specificity', 'Highly specific mechanisms', 'Non-specific immune activation'],
        ['Examples', 'Diphtheria toxin, Botulinum toxin, Cholera toxin, Tetanus toxin', 'LPS from E. coli, Salmonella, Pseudomonas (lipid A component)'],
    ]

    for row_idx, row_content in enumerate(row_data, start=1):
        cells = table_comp2.rows[row_idx].cells
        set_cell_background(cells[0], 'FFEBEE')
        set_cell_text(cells[0], row_content[0], bold=True)
        set_cell_background(cells[1], 'FFFFFF')
        set_cell_text(cells[1], row_content[1])
        set_cell_background(cells[2], 'FFFFFF')
        set_cell_text(cells[2], row_content[2])

    doc.add_paragraph()

    # Comparison 3: Capsule vs Slime Layer
    doc.add_heading('Capsule vs Slime Layer (Glycocalyx)', 2)
    table_comp3 = doc.add_table(rows=5, cols=3)
    table_comp3.style = 'Table Grid'

    headers = ['Feature', 'Capsule', 'Slime Layer']
    for i, header_text in enumerate(headers):
        set_cell_background(table_comp3.rows[0].cells[i], 'C8E6C9')
        set_cell_text(table_comp3.rows[0].cells[i], header_text, bold=True, size=12)

    row_data = [
        ['Structure', 'Well-organized, firmly attached', 'Loosely organized, easily detached'],
        ['Composition', 'Polysaccharides', 'Polysaccharides and glycoproteins'],
        ['Main Function', 'Antiphagocytic (virulence factor)', 'Adherence and biofilm formation'],
        ['Clinical Significance', 'Prevents immune clearance', 'Biofilms on medical devices, dental plaque'],
    ]

    for row_idx, row_content in enumerate(row_data, start=1):
        cells = table_comp3.rows[row_idx].cells
        set_cell_background(cells[0], 'E8F5E9')
        set_cell_text(cells[0], row_content[0], bold=True)
        set_cell_background(cells[1], 'FFFFFF')
        set_cell_text(cells[1], row_content[1])
        set_cell_background(cells[2], 'FFFFFF')
        set_cell_text(cells[2], row_content[2])

    doc.add_page_break()

    # ==========================================================================
    # SECTION 3: MASTER CHART
    # ==========================================================================
    heading3 = doc.add_heading('Master Chart - Bacterial Structures & Virulence Factors', 1)
    heading3.runs[0].font.color.rgb = RGBColor(118, 75, 162)

    doc.add_heading('Comprehensive Reference Table', 2)

    # Master chart table
    table_master = doc.add_table(rows=26, cols=2)
    table_master.style = 'Table Grid'

    # Headers
    headers = ['Structure/Component/Factor', 'Key Characteristics']
    for i, header_text in enumerate(headers):
        set_cell_background(table_master.rows[0].cells[i], 'EDE7F6')
        set_cell_text(table_master.rows[0].cells[i], header_text, bold=True, size=12)

    # Master chart data
    master_data = [
        ('Coccus', 'Spherical/round bacterial shape'),
        ('Bacillus', 'Rod-shaped bacteria'),
        ('Spirochete', 'Flexible spiral rod (Treponema, Borrelia, Leptospira)'),
        ('Diplo- arrangement', 'Pairs of cells'),
        ('Strepto- arrangement', 'Chains of cells (one plane of division)'),
        ('Staphylo- arrangement', 'Grape-like clusters (two or more planes of division)'),
        ('Peptidoglycan', 'NAG-NAM backbone with tetrapeptide sidechains and pentapeptide cross-links; target of beta-lactams'),
        ('Gram-positive cell wall', 'Thick peptidoglycan + teichoic acid; stains purple'),
        ('Gram-negative cell wall', 'Thin peptidoglycan + outer membrane with LPS; stains red/pink'),
        ('Acid-fast cell wall', 'Thin peptidoglycan + mycolic acid layer; requires Ziehl-Neelsen stain'),
        ('LPS (Lipopolysaccharide)', 'Gram-negative endotoxin; lipid A causes toxic shock; O antigen is immunogenic'),
        ('Mycolic acid', 'Waxy outer layer of acid-fast bacteria; provides resistance to desiccation and antibiotics'),
        ('Capsule', 'Well-organized glycocalyx; inhibits phagocytosis (virulence factor)'),
        ('Flagella', 'Protein filaments for motility; arrangements: monotrichous, lophotrichous, amphitrichous, peritrichous, atrichous'),
        ('Pili/Fimbriae', 'Adhesins for attachment; sex pilus mediates conjugation in Gram-negatives'),
        ('Nucleoid', 'Single circular chromosome (supercoiled dsDNA)'),
        ('Plasmids', 'Small circular dsDNA; often carry antibiotic resistance genes'),
        ('70S Ribosomes', '30S + 50S subunits; target of antibiotics (vs 80S in eukaryotes)'),
        ('Endospores', 'Dormant cells (Bacillus, Clostridium, Clostridioides); resist heat, desiccation, chemicals'),
        ('Adhesins', 'Virulence factors that bind to host cell receptors (pili, capsule)'),
        ('Invasins', 'Enzymes that damage tissues (collagenase, neuraminidase, streptokinase)'),
        ('Exotoxins', 'Secreted toxins with high specificity (diphtheria, botulinum, cholera toxins)'),
        ('Endotoxins', 'LPS in Gram-negative outer membrane; released when bacteria lyse'),
        ('Antiphagocytic factors', 'Prevent phagocytosis (capsule, protein A, protein M)'),
        ('Biofilm', 'Slime layer-mediated bacterial communities on surfaces; resistant to antibiotics'),
    ]

    for row_idx, row_content in enumerate(master_data, start=1):
        cells = table_master.rows[row_idx].cells
        set_cell_background(cells[0], 'F3E5F5')
        set_cell_text(cells[0], row_content[0], bold=True)
        set_cell_background(cells[1], 'FFFFFF')
        set_cell_text(cells[1], row_content[1])

    doc.add_page_break()

    # ==========================================================================
    # SECTION 4: HIGH-YIELD SUMMARY
    # ==========================================================================
    heading4 = doc.add_heading('High-Yield Summary', 1)
    heading4.runs[0].font.color.rgb = RGBColor(118, 75, 162)

    # Bacterial Morphology Must-Knows
    add_colored_box(doc, 'MORPHOLOGY - Must Know:', [
        'Bacteria range 0.2-2.0 µm in size (visible under light microscopy)',
        'Three main shapes: cocci (spheres), bacilli (rods), spirochetes (spirals)',
        'Arrangements depend on plane of division: diplo- (pairs), strepto- (chains), staphylo- (clusters)',
        'Morphology and arrangement are key initial steps in bacterial identification'
    ], (118, 75, 162), 'F3E5F5')

    doc.add_paragraph()

    # Cell Wall Types
    add_colored_box(doc, 'CELL WALL TYPES - Must Know:', [
        'Gram-positive: Thick peptidoglycan + teichoic acid → stains PURPLE',
        'Gram-negative: Thin peptidoglycan + outer membrane (LPS) → stains RED/PINK',
        'Acid-fast: Thin peptidoglycan + mycolic acid → requires Ziehl-Neelsen stain (appears RED)',
        'Peptidoglycan = NAG-NAM backbone with peptide cross-links',
        'LPS (endotoxin) in Gram-negatives: lipid A causes toxic shock, O antigen is immunogenic'
    ], (118, 75, 162), 'F3E5F5')

    doc.add_paragraph()

    # Clinical Staining
    add_colored_box(doc, 'STAINING - Clinical Pearls:', [
        'Gram stain: Primary stain (crystal violet) + counterstain (safranin)',
        'Gram-positive = thick peptidoglycan RETAINS purple; Gram-negative = thin peptidoglycan LOSES purple, takes up pink',
        'Gram stain guides empiric antibiotic therapy before final bacterial ID',
        'Acid-fast stain for Mycobacterium species (TB, leprosy) - waxy mycolic acid prevents Gram staining'
    ], (0, 77, 64), 'E0F2F1')

    doc.add_paragraph()

    # Antibiotic Targets
    add_colored_box(doc, 'ANTIBIOTIC TARGETS - Key Concepts:', [
        'Beta-lactams (penicillin, vancomycin) target peptidoglycan synthesis → cell wall disruption',
        'Lysozyme (in tears, saliva, immune cells) breaks down peptidoglycan',
        '70S ribosomes (30S + 50S) are antibiotic targets - do NOT affect human 80S ribosomes',
        'Outer membrane in Gram-negatives provides barrier to many antibiotics (intrinsic resistance)'
    ], (27, 94, 32), 'E8F5E9')

    doc.add_paragraph()

    # Virulence Factors Summary
    add_colored_box(doc, 'VIRULENCE FACTORS - High-Yield:', [
        'Adhesins: Attach to host cells (pili, fimbriae, capsule) - enable colonization',
        'Invasins: Damage tissues (collagenase, neuraminidase, streptokinase) - allow spread',
        'Toxins: Exotoxins (secreted, highly specific) vs Endotoxins (LPS, non-specific)',
        'Antiphagocytic factors: Capsule (slippery), protein A (Staph), protein M (Strep)',
        'Sequence: Attach → Invade → Damage → Evade'
    ], (183, 28, 28), 'FFEBEE')

    doc.add_paragraph()

    # Critical Don't Miss
    add_colored_box(doc, 'CRITICAL - Never Miss:', [
        'Endospore-forming bacteria (Bacillus, Clostridium, Clostridioides) survive boiling, require autoclaving',
        'Gram-negative septic shock = LPS (endotoxin) release → lipid A causes toxic shock',
        'Encapsulated bacteria (S. pneumoniae, H. influenzae, N. meningitidis) especially dangerous in asplenic patients',
        'Plasmid-mediated antibiotic resistance spreads via conjugation (sex pilus)',
        'Biofilms on medical devices are extremely difficult to eradicate with antibiotics',
        'Acid-fast positive sputum = presumptive TB → start isolation and therapy immediately'
    ], (183, 28, 28), 'FFEBEE')

    # Save document
    doc.save(output_path)
    print(f"Word LO study guide created: {output_path}")

# =============================================================================
# MAIN
# =============================================================================

if __name__ == '__main__':
    output_path = '/Users/kimnguyen/Library/Mobile Documents/com~apple~CloudDocs/Documents/Midwestern/Quarter 3/Microbiology /Claude Study Tools/Intro_Bacterial_Morphology_LO_Guide.docx'
    create_bacterial_morphology_guide(output_path)
